#!/usr/bin/env python3
"""Runner de la golden (Llave 1) — multi-contrato (C1 + C4 + C3).

Dos pasos de validación (ver CONTRATOS_estado-validacion-ubicacion.md §2):
  1. TODOS los esquemas de contrato (packages/contracts/C*/*.schema.json) son JSON Schema
     bien formados.
  2. La golden los ejercita, DESPACHADA POR CONTRATO (Fase II·h1):
     - C1: compile real (caso.alto.json → IFC con engines/ifc) + recompute de métricas
       contra expected.json (costura cerrada en Fase I·h1).
     - C4: RECOMPUTE con el service real (services/federacion, costura cerrada en
       Fase II·h2 — misma costura que C1 en Fase 0 → Fase I): federar+validar sobre
       las entradas congeladas contra el MISMO expected.json, con tolerancias; MÁS
       los checks ANCLADOS del 2.1 (conformidad de esquemas, identidad por hash y
       coherencia interna), que se conservan íntegros (D10: más checks, nunca menos).

Regla de oro: un fallo se corrige en el código, NUNCA aflojando la golden.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import math
import sys
from pathlib import Path

# repo root = .../aqyra ; este fichero = aqyra/packages/golden/src/aqyra_golden/run_golden.py
_HERE = Path(__file__).resolve()
DEFAULT_GOLDEN_DIR = _HERE.parents[2]           # aqyra/packages/golden
DEFAULT_REPO = _HERE.parents[4]                 # aqyra
DEFAULT_CONTRACTS = DEFAULT_REPO / "packages" / "contracts"
DEFAULT_ENGINE = DEFAULT_REPO / "engines" / "ifc"   # engine C1 (compile narración→IFC)

# Familias de elemento que declaran las golden C1 y que deben llevar doble clasificación.
# (4 pilares → IfcColumn, 1 muro → IfcWall, 1 losa → IfcSlab, 1 ascensor → IfcTransportElement).
TARGET_CLASSES = {"IfcColumn", "IfcWall", "IfcSlab", "IfcTransportElement"}

GREEN = "\033[32m"; RED = "\033[31m"; DIM = "\033[2m"; RST = "\033[0m"


# --------------------------------------------------------------------------- #
# Esquemas de contrato                                                        #
# --------------------------------------------------------------------------- #
def load_c1_schemas(contracts_dir: Path) -> dict[str, dict]:
    base = contracts_dir / "C1-interoperabilidad"
    out = {}
    for key, fname in (("alto", "alto-spec.schema.json"),
                       ("neutro", "modelo-neutro.schema.json")):
        p = base / fname
        out[key] = json.loads(p.read_text(encoding="utf-8"))
    return out


def load_c4_schemas(contracts_dir: Path) -> dict[str, dict]:
    base = contracts_dir / "C4-federacion"
    out = {}
    for key, fname in (("reglas", "reglas-federacion.schema.json"),
                       ("manifiesto", "maestro-manifiesto.schema.json"),
                       ("informe", "informe-qa.schema.json")):
        out[key] = json.loads((base / fname).read_text(encoding="utf-8"))
    return out


def load_c3_schemas(contracts_dir: Path) -> dict[str, dict]:
    base = contracts_dir / "C3-cumplimiento"
    out = {}
    for key, fname in (("entrada", "entrada-cumplimiento.schema.json"),
                       ("veredicto", "veredicto-cumplimiento.schema.json")):
        out[key] = json.loads((base / fname).read_text(encoding="utf-8"))
    return out


def discover_schemas(contracts_dir: Path) -> dict[str, dict]:
    """TODOS los esquemas de contrato del registro (paso 1 es multi-contrato)."""
    out = {}
    for p in sorted(contracts_dir.glob("C*/*.schema.json")):
        out[f"{p.parent.name}/{p.name}"] = json.loads(p.read_text(encoding="utf-8"))
    return out


def check_schemas_wellformed(schemas: dict[str, dict]) -> list[tuple[str, bool, str]]:
    """Paso 1: los esquemas son JSON Schema bien formados (meta-validación)."""
    from jsonschema import Draft202012Validator
    results = []
    for key, schema in schemas.items():
        try:
            Draft202012Validator.check_schema(schema)
            results.append((f"esquema '{key}' bien formado", True, "JSON Schema 2020-12 válido"))
        except Exception as e:  # noqa: BLE001
            results.append((f"esquema '{key}' bien formado", False, str(e).splitlines()[0]))
    return results


def validate_against_schema(instance: dict, schema: dict) -> tuple[bool, str]:
    from jsonschema import Draft202012Validator
    v = Draft202012Validator(schema)
    errs = sorted(v.iter_errors(instance), key=lambda e: list(e.path))
    if not errs:
        return True, "conforma el esquema"
    e = errs[0]
    loc = "/".join(str(p) for p in e.path) or "(raíz)"
    return False, f"{loc}: {e.message}"


# --------------------------------------------------------------------------- #
# C1 · oráculo: compile real + métricas recomputadas                           #
# --------------------------------------------------------------------------- #
def _attr(seg, name, idx):
    try:
        return getattr(seg, name)
    except Exception:  # noqa: BLE001
        return seg[idx]


def compile_case(alto_path: Path, out_ifc: Path, engine_dir: Path = DEFAULT_ENGINE) -> None:
    """Cierra la costura: compila caso.alto.json → IFC con el engine (engines/ifc).

    El oráculo NO recibe ya un IFC congelado, sino el IFC que produce el compile real
    del canónico importado. Un fallo aquí se investiga en el compile, jamás en expected.json.
    """
    import json as _json
    p = str(engine_dir)
    if p not in sys.path:
        sys.path.insert(0, p)
    import compile_c1  # engines/ifc/compile_c1.py
    alto = _json.loads(alto_path.read_text(encoding="utf-8"))
    compile_c1.compilar_alto_a_ifc(alto, out_ifc)


def compute_metrics(ifc_path: Path) -> dict:
    import ifcopenshell  # lazy: no se necesita en --schema-only

    f = ifcopenshell.open(str(ifc_path))
    m: dict = {"esquema": f.schema}

    # huecos por tipo de elemento vaciado
    losa = muro = 0
    for r in f.by_type("IfcRelVoidsElement"):
        cls = r.RelatingBuildingElement.is_a()
        if cls == "IfcSlab":
            losa += 1
        elif cls == "IfcWall":
            muro += 1
    m["huecos"] = {"losa": losa, "muro": muro}

    # elementos de transporte (catálogo abierto)
    tes = f.by_type("IfcTransportElement")
    m["transporte"] = {
        "total": len(tes),
        "predefinidos": sorted({str(getattr(t, "PredefinedType", None)) for t in tes}),
    }

    # alineación horizontal
    segs = f.by_type("IfcAlignmentHorizontalSegment")
    tipos = [_attr(s, "PredefinedType", 8) for s in segs]
    longs = [float(_attr(s, "SegmentLength", 6)) for s in segs]
    clot_A = []
    for s in segs:
        if _attr(s, "PredefinedType", 8) == "CLOTHOID":
            r0 = float(_attr(s, "StartRadiusOfCurvature", 4) or 0.0)
            r1 = float(_attr(s, "EndRadiusOfCurvature", 5) or 0.0)
            R = r0 if r0 else r1
            L = float(_attr(s, "SegmentLength", 6))
            if R:
                clot_A.append(math.sqrt(L * R))
    m["alineacion"] = {
        "planta_secuencia": tipos,
        "longitud_total_m": round(sum(longs), 6),
        "clotoide_A": [round(a, 3) for a in clot_A],
    }

    # doble clasificación (bsDD + Uniclass) sobre las familias objetivo
    def system_of(ref):
        cur = ref
        for _ in range(8):
            rs = getattr(cur, "ReferencedSource", None)
            if rs is None:
                break
            cur = rs
        name = (getattr(cur, "Name", None) or getattr(cur, "Source", None) or "")
        n = name.lower()
        if "uniclass" in n:
            return "uniclass"
        if "bsdd" in n:
            return "bsdd"
        return "otro"

    sistemas: dict[int, set[str]] = {}
    targets = [o for o in f.by_type("IfcElement") if o.is_a() in TARGET_CLASSES]
    for o in targets:
        sistemas[o.id()] = set()
    for rel in f.by_type("IfcRelAssociatesClassification"):
        sys_tag = system_of(rel.RelatingClassification)
        for o in rel.RelatedObjects:
            if o.id() in sistemas:
                sistemas[o.id()].add(sys_tag)
    doble = sum(1 for s in sistemas.values() if {"uniclass", "bsdd"} <= s)
    sin_clasif = sum(1 for s in sistemas.values() if not s)
    m["doble_clasificacion"] = {
        "objetivo": len(targets), "elementos": doble, "sin_clasif": sin_clasif,
    }
    return m


# --------------------------------------------------------------------------- #
# C1 · comparación caso vs oráculo                                             #
# --------------------------------------------------------------------------- #
def compare(expected: dict, got: dict, tol: dict) -> list[tuple[str, bool, str]]:
    checks: list[tuple[str, bool, str]] = []

    def chk(name, ok, detail):
        checks.append((name, bool(ok), detail))

    chk("esquema IFC", got["esquema"] == expected["esquema"],
        f"{got['esquema']} (esperado {expected['esquema']})")

    eh, gh = expected["huecos"], got["huecos"]
    chk("huecos losa", gh["losa"] == eh["losa"], f"{gh['losa']} (esperado {eh['losa']})")
    chk("huecos muro", gh["muro"] == eh["muro"], f"{gh['muro']} (esperado {eh['muro']})")

    et, gt = expected["transporte"], got["transporte"]
    chk("ascensor total", gt["total"] == et["total"], f"{gt['total']} (esperado {et['total']})")
    chk("ascensor predefinido", gt["predefinidos"] == et["predefinidos"],
        f"{gt['predefinidos']} (esperado {et['predefinidos']})")

    ea, ga = expected["alineacion"], got["alineacion"]
    chk("planta secuencia", ga["planta_secuencia"] == ea["planta_secuencia"],
        f"{ga['planta_secuencia']}")
    tolL = float(tol.get("longitud_total_m", 1e-6))
    okL = abs(ga["longitud_total_m"] - ea["longitud_total_m"]) <= tolL
    chk("longitud total (m)", okL,
        f"{ga['longitud_total_m']} (esperado {ea['longitud_total_m']} ±{tolL})")
    tolA = float(tol.get("clotoide_A", 0.1))
    A_ok = bool(ga["clotoide_A"]) and all(abs(a - ea["clotoide_A"]) <= tolA for a in ga["clotoide_A"])
    chk("clotoide A", A_ok, f"{ga['clotoide_A']} (esperado {ea['clotoide_A']} ±{tolA})")

    ec, gc = expected["doble_clasificacion"], got["doble_clasificacion"]
    chk("doble clasificación", gc["elementos"] == ec["elementos"],
        f"{gc['elementos']}/{gc['objetivo']} (esperado {ec['elementos']})")
    chk("sin clasificar", gc["sin_clasif"] == ec["sin_clasif"],
        f"{gc['sin_clasif']} (esperado {ec['sin_clasif']})")
    return checks


# --------------------------------------------------------------------------- #
# Casos por contrato                                                          #
# --------------------------------------------------------------------------- #
def run_case_c1(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict) -> bool:
    """C1: conformidad del caso + compile real + recompute de métricas (Fase I·h1)."""
    schemas = load_c1_schemas(contracts_dir)
    case_ok = True

    alto_path = case_dir / "caso.alto.json"
    if alto_path.exists():
        alto = json.loads(alto_path.read_text(encoding="utf-8"))
        ok, detail = validate_against_schema(alto, schemas["alto"])
        print_checks([("caso conforma alto-spec", ok, detail)])
        case_ok &= ok
    else:
        print_checks([("caso.alto.json presente", False, "no encontrado (necesario para compilar)")])
        return False

    import tempfile
    try:
        with tempfile.TemporaryDirectory() as td:
            compiled = Path(td) / "compilado.ifc"
            compile_case(alto_path, compiled)
            print_checks([("compile narración→IFC", True, f"engine → {compiled.name}")])
            got = compute_metrics(compiled)
            checks = compare(expected, got, tol)
            print_checks(checks)
            case_ok &= all(ok for _, ok, _ in checks)
    except Exception as e:  # noqa: BLE001
        print_checks([("compile+oráculo", False, f"{type(e).__name__}: {e}")])
        case_ok = False
    return case_ok


def _md5(path: Path) -> str:
    return hashlib.md5(path.read_bytes()).hexdigest()


def _ids_identifiers(ids_path: Path) -> set[str]:
    """Identificadores de las specifications del .ids (IDS 1.0, buildingSMART)."""
    import xml.etree.ElementTree as ET
    ns = "{http://standards.buildingsmart.org/IDS}"
    root = ET.parse(str(ids_path)).getroot()
    out = set()
    for spec in root.iter(f"{ns}specification"):
        out.add(spec.get("identifier") or spec.get("name") or "")
    return out - {""}


def _lock_packs_ids(repo: Path) -> dict:
    """Sección [packs.ids] de versions.lock (ancla del pack IDS)."""
    try:  # Python 3.11+
        import tomllib
    except ModuleNotFoundError:  # 3.10
        import tomli as tomllib
    with open(repo / "versions.lock", "rb") as f:
        lock = tomllib.load(f)
    return lock.get("packs", {}).get("ids", {})


# --------------------------------------------------------------------------- #
# C4 · recompute con el service real (Fase II·h2 — costura cerrada)            #
# --------------------------------------------------------------------------- #
DEFAULT_SERVICE_SRC = DEFAULT_REPO / "services" / "federacion" / "src"

# Campos de TEXTO LIBRE del contrato (presentación humana): se comprueban por
# esquema (presencia/tipo) pero NO se comparan literalmente en el recompute —
# la semántica del contrato son ids, resultados, conteos, GUIDs, estados,
# severidades y veredicto. Ídem claves '_*' (comentario, convención del repo).
_CAMPOS_LIBRES = {"titulo", "detalle"}


def _norm_contrato(x):
    if isinstance(x, dict):
        return {k: _norm_contrato(v) for k, v in x.items()
                if not k.startswith("_") and k not in _CAMPOS_LIBRES}
    if isinstance(x, list):
        return [_norm_contrato(v) for v in x]
    return x


def _norm_manifiesto(m: dict) -> dict:
    """Además de la normalización general, la procedencia es metadato de
    generación: se compara reglas_md5 (exacto) y la PRESENCIA de generado_por,
    no su literal (el expected lo generó el oráculo manual; el recompute, el service)."""
    m = _norm_contrato(m)
    p = m.get("procedencia", {}) or {}
    m["procedencia"] = {"reglas_md5": p.get("reglas_md5"),
                        "tiene_generado_por": bool(p.get("generado_por"))}
    return m


def _diffs(exp, got, tol_tras: float, tol_rot: float,
           path: str = "", clave=None, out=None) -> list[str]:
    """Comparación profunda expected vs recompute con las tolerancias del caso."""
    if out is None:
        out = []
    if isinstance(exp, dict) and isinstance(got, dict):
        for k in sorted(set(exp) | set(got)):
            if k not in exp:
                out.append(f"{path}/{k}: clave no esperada")
            elif k not in got:
                out.append(f"{path}/{k}: falta en el recompute")
            else:
                _diffs(exp[k], got[k], tol_tras, tol_rot, f"{path}/{k}", k, out)
    elif isinstance(exp, list) and isinstance(got, list):
        if len(exp) != len(got):
            out.append(f"{path}: longitud {len(exp)} vs {len(got)}")
        else:
            for i, (a, b) in enumerate(zip(exp, got)):
                _diffs(a, b, tol_tras, tol_rot, f"{path}[{i}]", clave, out)
    elif isinstance(exp, float) or isinstance(got, float):
        tol = tol_tras if clave == "traslacion" else (tol_rot if clave == "rotacion_deg" else 0.0)
        if abs(float(exp) - float(got)) > tol:
            out.append(f"{path}: {exp} vs {got} (tol ±{tol})")
    elif exp != got:
        out.append(f"{path}: {exp!r} vs {got!r}")
    return out


def _recompute_c4(case_dir: Path, reglas: dict, repo: Path) -> tuple[dict, dict]:
    """Cierra la costura: federar+validar con services/federacion (import de path,
    mismo patrón que engines/ifc). Un fallo aquí se investiga en el SERVICE,
    jamás en expected.json (regla de oro)."""
    p = str(repo / "services" / "federacion" / "src")
    if p not in sys.path:
        sys.path.insert(0, p)
    from aqyra_federacion import federar_fichero, validar
    manifiesto = federar_fichero(case_dir / "reglas.json")
    pack = reglas.get("ids_pack", {})
    pack_dir = (repo / "data" / "packs" / str(pack.get("familia", "ids"))
                / str(pack.get("id")) / str(pack.get("version")))
    informe = validar(manifiesto, pack_dir, case_dir)
    return manifiesto, informe


def _derivar_c4(case_dir: Path, expected: dict, man_rec: dict, checks: list,
                tmpdir: Path):
    """Fase II·h6 (D26/D30): paso de DERIVACIÓN, activado por el expected
    (presencia de maestro_manifiesto.ifc_derivado — patrón D14: 01–05 ni pasan
    por aquí). Deriva con el service en un directorio temporal, ancla el .ifc
    por md5 BYTE A BYTE (D26, cabecera determinista con time_stamp inyectado de
    derivado_generacion) y devuelve el manifiesto actualizado (con ifc_derivado,
    que el diff compara contra el expected — doble anclaje) + la ruta del
    derivado (la CÁMARA del BCF lo consume — D29)."""
    from aqyra_federacion import derivar
    exp_der = expected["maestro_manifiesto"]["ifc_derivado"]
    gen = expected.get("derivado_generacion", {})
    salida = tmpdir / exp_der.get("fichero", "federado.ifc")
    man_rec = derivar(man_rec, case_dir, salida, fecha=gen.get("fecha"))
    got = man_rec["ifc_derivado"]["md5"]
    exp_md5 = exp_der.get("md5")
    checks.append(("derivación IFC federado ejecuta", True,
                   f"'{salida.name}' escrito (cabecera SPF determinista)"))
    checks.append(("derivado md5 == expected (byte a byte)", got == exp_md5,
                   f"md5 {got[:12]}… (esperado {str(exp_md5)[:12]}…)"))
    return man_rec, salida


def _emitir_bcf_c4(inf_rec: dict, expected: dict, checks: list,
                   derivado: Path | None = None) -> dict:
    """Fase II·h3 (tarea 1.2, D14): paso de EMISIÓN, activado por el expected
    (informe_qa.bcf.emitido == true). Emite con el service en un directorio
    temporal, ancla el ÁRBOL BCF 3.0 por md5 de fichero (D12) y devuelve el
    informe actualizado (emitido=true) para el diff contra el expected.
    C4-FED-01 (emitido=false) ni pasa por aquí — su expected no se toca.
    Fase II·h6 (D29): si el caso derivó (paso anterior), la emisión recibe el
    derivado y cada viewpoint gana su CÁMARA determinista — 02/04 (sin derivado)
    quedan intactos POR CONSTRUCCIÓN."""
    import tempfile
    from aqyra_federacion import emitir_bcf
    gen = expected.get("bcf_generacion", {})
    nombre = expected.get("informe_qa", {}).get("bcf", {}).get("carpeta", "bcf")
    with tempfile.TemporaryDirectory() as td:
        carpeta = Path(td) / nombre
        inf_rec = emitir_bcf(inf_rec, carpeta, caso=expected.get("caso"),
                             autor=gen.get("autor"), fecha=gen.get("fecha"),
                             derivado=derivado)
        got = {q.relative_to(carpeta).as_posix(): _md5(q)
               for q in sorted(carpeta.rglob("*")) if q.is_file()}
    checks.append(("emisión BCF 3.0 ejecuta", True,
                   f"{len(got)} fichero/s en '{nombre}/' (un topic por incidencia)"))
    exp = {k: v for k, v in expected.get("bcf_md5", {}).items()
           if not k.startswith("_")}
    dif = sorted(set(exp.items()) ^ set(got.items()))
    checks.append(("árbol BCF == expected (md5/fichero)", got == exp,
                   f"{len(dif)} diff/s — {dif[0][0]}" if dif
                   else f"{len(exp)} fichero/s reproducidos byte a byte"))
    return inf_rec


def run_case_c4(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                repo: Path = DEFAULT_REPO) -> bool:
    """C4: recompute con el service (Fase II·h2) + checks anclados del 2.1 — ver ficha.

    0. RECOMPUTE (costura cerrada): federar+validar con services/federacion sobre
       las entradas congeladas → manifiesto e informe comparados contra el MISMO
       expected.json (tolerancias del caso; texto libre y procedencia normalizados).
       Fase II·h6: si el expected declara ifc_derivado, se ANTEPONE la derivación
       (md5 byte a byte, D26) y la emisión gana la cámara (D29).
    1. Conformidad de esquemas (reglas + manifiesto esperado + informe esperado).
    2. Identidad por hash de las entradas congeladas (triple coherencia de md5).
    3. Coherencia interna (modelos, pack IDS anclado, requisitos ⊆ .ids, veredicto).
    """
    checks: list[tuple[str, bool, str]] = []
    schemas = load_c4_schemas(contracts_dir)
    manifiesto = expected.get("maestro_manifiesto", {})
    informe = expected.get("informe_qa", {})

    reglas_path = case_dir / "reglas.json"
    if not reglas_path.exists():
        print_checks([("reglas.json presente", False, "no encontrado")])
        return False
    reglas = json.loads(reglas_path.read_text(encoding="utf-8"))

    # 0 · RECOMPUTE con el service real (Fase II·h2, D10) — se ANTEPONE al anclado
    tol_tras = float(tol.get("traslacion_m", 0.0))
    tol_rot = float(tol.get("rotacion_deg", 0.0))
    try:
        import tempfile
        with tempfile.TemporaryDirectory() as td6:
            man_rec, inf_rec = _recompute_c4(case_dir, reglas, repo)
            checks.append(("service federar()+validar() ejecuta", True,
                           "services/federacion (import de path)"))
            derivado = None
            if manifiesto.get("ifc_derivado"):
                man_rec, derivado = _derivar_c4(case_dir, expected, man_rec,
                                                checks, Path(td6))   # Fase II·h6 (D26)
            if informe.get("bcf", {}).get("emitido"):
                inf_rec = _emitir_bcf_c4(inf_rec, expected, checks,
                                         derivado)   # Fase II·h3 (D14) + cámara (D29)
            for name, inst, key in (("manifiesto recomputado conforma", man_rec, "manifiesto"),
                                    ("informe recomputado conforma", inf_rec, "informe")):
                ok, detail = validate_against_schema(inst, schemas[key])
                checks.append((name, ok, detail))
            d_man = _diffs(_norm_manifiesto(manifiesto), _norm_manifiesto(man_rec),
                           tol_tras, tol_rot)
            checks.append(("recompute manifiesto == expected", not d_man,
                           f"{len(d_man)} diff/s — {d_man[0]}" if d_man
                           else f"reproducido (±{tol_tras} m, ±{tol_rot}°)"))
            d_inf = _diffs(_norm_contrato(informe), _norm_contrato(inf_rec),
                           tol_tras, tol_rot)
            checks.append(("recompute informe == expected", not d_inf,
                           f"{len(d_inf)} diff/s — {d_inf[0]}" if d_inf
                           else "reproducido (pass/fail, conteos, GUIDs, estados, veredicto)"))
    except Exception as e:  # noqa: BLE001
        checks.append(("recompute con el service", False, f"{type(e).__name__}: {e}"))

    # 1 · conformidad de esquemas
    for name, inst, key in (("reglas conforman esquema", reglas, "reglas"),
                            ("manifiesto esperado conforma", manifiesto, "manifiesto"),
                            ("informe esperado conforma", informe, "informe")):
        ok, detail = validate_against_schema(inst, schemas[key])
        checks.append((name, ok, detail))

    # 2 · identidad de las entradas congeladas (hash) + triple coherencia
    entradas = expected.get("entradas_md5", {})
    checks.append(("entradas declaradas", bool(entradas), f"{len(entradas)} fichero/s"))
    for rel, exp_md5 in sorted(entradas.items()):
        p = case_dir / rel
        if not p.exists():
            checks.append((f"identidad {rel}", False, "fichero no encontrado"))
            continue
        got = _md5(p)
        checks.append((f"identidad {rel}", got == exp_md5,
                       f"md5 {got[:12]}… (esperado {exp_md5[:12]}…)"))
    md5_reglas = {m["fichero"]: m.get("md5") for m in reglas.get("modelos", [])}
    md5_man = {m["fichero_origen"]: m.get("md5") for m in manifiesto.get("modelos", [])}
    checks.append(("md5 reglas == entradas", md5_reglas == entradas,
                   "los md5 de reglas.json anclan los mismos ficheros"))
    checks.append(("md5 manifiesto == entradas", md5_man == entradas,
                   "los md5 del manifiesto anclan los mismos ficheros"))

    # 3 · coherencia interna
    ids_reglas = {m["id"] for m in reglas.get("modelos", [])}
    ids_man = {m["id"] for m in manifiesto.get("modelos", [])}
    ids_estados = set(informe.get("estados", {}).get("por_modelo", {}))
    checks.append(("modelos coherentes", ids_reglas == ids_man == ids_estados,
                   f"reglas={sorted(ids_reglas)} manifiesto={sorted(ids_man)} informe={sorted(ids_estados)}"))

    pack_reglas = reglas.get("ids_pack", {})
    pack_informe = informe.get("ids", {})
    mismo_pack = (pack_reglas.get("id") == pack_informe.get("pack_id")
                  and pack_reglas.get("version") == pack_informe.get("pack_version"))
    checks.append(("pack IDS reglas == informe", mismo_pack,
                   f"{pack_reglas.get('id')}/{pack_reglas.get('version')}"))

    lock_ids = _lock_packs_ids(repo)
    anclado = (lock_ids.get("id") == pack_reglas.get("id")
               and lock_ids.get("version") == pack_reglas.get("version"))
    checks.append(("pack IDS anclado en versions.lock", anclado,
                   f"lock={lock_ids.get('id')}/{lock_ids.get('version')}"))

    pack_dir = repo / "data" / "packs" / "ids" / str(pack_reglas.get("id")) / str(pack_reglas.get("version"))
    ids_ok, ids_detail = False, "pack.json no encontrado"
    if (pack_dir / "pack.json").exists():
        pack = json.loads((pack_dir / "pack.json").read_text(encoding="utf-8"))
        ids_file = pack_dir / pack["contenido"]["fichero"]
        declarados = _ids_identifiers(ids_file)
        exigidos = {r["id"] for r in informe.get("requisitos", []) if r.get("origen", "ids") == "ids"}
        faltan = exigidos - declarados
        ids_ok = not faltan
        ids_detail = (f"{sorted(exigidos)} ⊆ .ids" if ids_ok
                      else f"requisitos SIN respaldo en el .ids: {sorted(faltan)}")
    checks.append(("requisitos (origen=ids) ⊆ .ids del pack", ids_ok, ids_detail))

    resultados = [r.get("resultado") for r in informe.get("requisitos", [])]
    esperado_veredicto = "no-conforme" if "fail" in resultados else "conforme"
    checks.append(("veredicto coherente", informe.get("veredicto") == esperado_veredicto,
                   f"{informe.get('veredicto')} (requisitos: "
                   f"{resultados.count('pass')} pass / {resultados.count('fail')} fail)"))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


# --------------------------------------------------------------------------- #
# C3 · cumplimiento — modo ANCLADO (contract-first, Fase III·h2 — sin engine)  #
# --------------------------------------------------------------------------- #
RESULTADOS_C3 = {"cumple", "no-cumple", "no-aplica", "no-verificable"}


def _lock_packs(repo: Path, clave: str) -> dict:
    """Sección [packs.<clave>] de versions.lock (ancla del pack)."""
    try:  # Python 3.11+
        import tomllib
    except ModuleNotFoundError:  # 3.10
        import tomli as tomllib
    with open(repo / "versions.lock", "rb") as f:
        lock = tomllib.load(f)
    return lock.get("packs", {}).get(clave, {})


def _banco_anclado_en_lock(repo: Path, ref: dict) -> tuple[bool, str]:
    """El banco de coste se ancla por su PROPIA clave [packs.banco*] (banco=AQ-DEMO,
    banco_bc3=AQ-BC3-DEMO, banco_bcca=BCCA). Busca la clave que ancla id+version del `banco_ref`
    del caso; el pointer [packs.banco] (=AQ-DEMO/v1) NO se re-mueve al anadir un banco real
    (D50/E5.1). Generaliza la comprobacion para servir AQ-DEMO y BCCA sin tocar las golden ancladas."""
    for clave in ("banco", "banco_bc3", "banco_bcca", "banco_bcca_nativo"):
        sec = _lock_packs(repo, clave)
        if sec.get("id") == ref.get("id") and sec.get("version") == ref.get("version"):
            return True, f"lock[{clave}]={sec.get('id')}/{sec.get('version')}"
    return False, f"sin ancla [packs.banco*] para {ref.get('id')}/{ref.get('version')}"


def _criterio_anclado_por_sha(repo: Path, ref: dict) -> tuple[bool, str]:
    """Un criterio que NO es la version anclada en [packs.criterio] (p. ej. AQ/v3 en GOL-PRE-05) se
    ancla por su PROPIO content_sha256 (patron AQ/v2 en la ruta de carbono, E3.3): el pointer
    [packs.criterio] NO se mueve (GOL-PRE-01..04 intactas). Compara content_hash(pack) contra el
    golden de pack del criterio."""
    pp = str(repo / "packages" / "packs" / "src")
    if pp not in sys.path:
        sys.path.insert(0, pp)
    try:
        import aqyra_packs as _packs
        man = _packs.load_pack(repo / "data" / "packs", "criterio", ref.get("id"), ref.get("version"))
        got = _packs.content_hash(man)
        exp = json.loads((repo / "data" / "packs" / "criterio" / str(ref.get("id"))
                          / str(ref.get("version")) / "golden" / "expected.json")
                         .read_text(encoding="utf-8"))["content_sha256"]
        return got == exp, f"content_sha256 {got[:12]}… (esperado {exp[:12]}…)"
    except Exception as e:  # noqa: BLE001
        return False, f"{type(e).__name__}: {e}"


# Casa del engine C3 (Fase III·h3, D6): engines/cumplimiento — importado por path (patrón
# engines/ifc y _recompute_c4). El runner regenera el Maestro y el engine da el veredicto (D7).
DEFAULT_ENGINE_C3 = DEFAULT_REPO / "engines" / "cumplimiento" / "src"

# Texto libre del veredicto C3 (presentación humana): se comprueba por esquema (presencia/tipo)
# pero NO se compara literalmente en el recompute (D9) — la semántica del contrato son ids,
# resultados, referencias, conteos y veredicto. Ídem claves '_*'.
_LIBRES_C3 = {"evidencia", "motivo_no_verificable", "exigencia", "detalle"}


def _norm_c3(x):
    if isinstance(x, dict):
        return {k: _norm_c3(v) for k, v in x.items()
                if not k.startswith("_") and k not in _LIBRES_C3}
    if isinstance(x, list):
        return [_norm_c3(v) for v in x]
    return x


def _recompute_c3(case_dir: Path, entrada: dict, repo: Path, tmpdir: Path) -> dict:
    """Cierra la costura C3 (Fase III·h3, D9): regenera el Maestro con services/federacion
    (federar+derivar — el engine NO federa, D7) y emite el veredicto con engines/cumplimiento
    (import de path). Un fallo aquí se investiga en el ENGINE, jamás en expected.json."""
    ps = str(repo / "services" / "federacion" / "src")
    if ps not in sys.path:
        sys.path.insert(0, ps)
    pe = str(DEFAULT_ENGINE_C3)
    if pe not in sys.path:
        sys.path.insert(0, pe)
    from aqyra_federacion import federar_fichero, derivar
    from aqyra_cumplimiento import verificar
    manifiesto = federar_fichero(case_dir / "reglas.json")
    derivado = Path(tmpdir) / "federado.ifc"
    manifiesto = derivar(manifiesto, case_dir, derivado)
    pack = entrada.get("pack_normativo", {})
    pack_dir = (repo / "data" / "packs" / "normativa"
                / str(pack.get("id")) / str(pack.get("version")))
    maestro = {"manifiesto": manifiesto, "base_dir": case_dir,
               "ifc_derivado": derivado, "proyecto": entrada.get("proyecto")}
    return verificar(maestro, entrada.get("uso", {}),
                     entrada.get("localizacion", {}), pack_dir)


# --------------------------------------------------------------------------- #
# C3 · write-back 6D — el cumplimiento vuelve al modelo (vertical visor-cumplimiento-6d)        #
# --------------------------------------------------------------------------- #
# Severidad del peor caso (D-6D-2) — REIMPLEMENTADA aquí para NO depender del engine: el runner
# reproyecta el veredicto a elementos y lo casa con el Pset del IFC 6D (dos implementaciones que
# deben coincidir, como _checks_5d casa el cost schedule con el presupuesto).
_SEVERIDAD_6D = {"no-cumple": 3, "no-verificable": 2, "cumple": 1, "no-aplica": 0}


def _proyectar_peor_caso_6d(veredicto: dict, guid2mod: dict) -> dict:
    """{GlobalId → (resultado, exigencia_id)} proyectando `por_modelo` a cada elemento con la regla
    del peor caso (D-6D-2): no-cumple≻no-verificable≻cumple; no-aplica neutro; empate → primera."""
    exigencias = veredicto.get("exigencias", [])
    por_modelo: dict[str, tuple] = {}
    for m in set(guid2mod.values()):
        best = None  # (severidad, -orden, exig_id, resultado)
        for orden, e in enumerate(exigencias):
            pm = e.get("por_modelo")
            if pm is not None:
                if m not in pm:
                    continue
                r = pm[m].get("resultado")
            else:
                r = e.get("resultado")
            if r is None:
                continue
            cand = (_SEVERIDAD_6D.get(r, 0), -orden, e.get("id"), r)
            if best is None or cand[:2] > best[:2]:
                best = cand
        if best is not None:
            por_modelo[m] = (best[3], best[2])
    return {g: por_modelo[m] for g, m in guid2mod.items() if m in por_modelo}


def _checks_6d(f, veredicto: dict, guid2mod: dict, exp6d: dict, r1: dict) -> list[tuple[str, bool, str]]:
    """SEMÁNTICA 6D (D-6D-2): el `Pset_Aqyra_Cumplimiento` del IFC casa con el veredicto proyectado a
    elementos por peor caso (proyección INDEPENDIENTE del runner). Ancla conteos contra el expected."""
    from collections import Counter

    import ifcopenshell.util.element as _ue

    out: list[tuple[str, bool, str]] = []
    pset_name = exp6d.get("pset", "Pset_Aqyra_Cumplimiento")
    proy = _proyectar_peor_caso_6d(veredicto, guid2mod)  # {guid: (resultado, exig)}

    elems = {el.GlobalId: el for el in f.by_type("IfcElement")}
    mapeados = set(elems) & set(guid2mod)  # elementos del derivado con procedencia (D1)
    leido: dict[str, dict] = {}
    for g in mapeados:
        try:
            ps = _ue.get_psets(elems[g]) or {}
        except Exception:  # noqa: BLE001
            ps = {}
        if pset_name in ps:
            leido[g] = ps[pset_name]

    # 1 · cobertura total (invariante D-6D-2: cada elemento del Maestro recibe un Resultado)
    sin_pset = mapeados - set(leido)
    out.append((f"{pset_name} en cada elemento del Maestro", not sin_pset,
                f"{len(leido)}/{len(mapeados)} elementos con Pset" if not sin_pset
                else f"sin Pset: {len(sin_pset)}"))
    out.append(("n_elementos (write-back) == elementos del derivado",
                r1["n_elementos"] == len(mapeados),
                f"{r1['n_elementos']} escritos (== {len(mapeados)} mapeados)"))

    # 2 · taxonomía cerrada del Resultado (D4)
    resultados = [str(v.get("Resultado")) for v in leido.values()]
    out.append(("taxonomía de Resultado cerrada (D4)",
                bool(resultados) and all(r in RESULTADOS_C3 for r in resultados),
                f"{sorted(set(resultados))}"))

    # 3 · SEMÁNTICA — Resultado (y Exigencia dominante) del Pset == peor caso proyectado
    malos = [g for g in leido if g in proy and str(leido[g].get("Resultado")) != proy[g][0]]
    out.append(("Resultado del Pset == veredicto por peor caso", not malos,
                "todos los elementos casan" if not malos else f"descuadran: {len(malos)}"))
    mal_exig = [g for g in leido if g in proy and str(leido[g].get("Exigencia")) != str(proy[g][1])]
    out.append(("Exigencia dominante del Pset == proyección", not mal_exig,
                "exigencia dominante coherente" if not mal_exig else f"descuadran: {len(mal_exig)}"))

    # 4 · conteos == expected (ancla) y == proyección independiente
    # La proyección se restringe a los elementos ESCRITOS (IfcElement): guid2mod mapea todo
    # IfcProduct con procedencia (incl. estructura espacial), pero el engine escribe el Pset solo
    # en IfcElement — el conteo relevante del write-back es el de esos elementos.
    cnt = Counter(resultados)
    cnt_proy = Counter(proy[g][0] for g in leido if g in proy)
    exp_por = exp6d.get("por_resultado", {})
    conteo_ok = (all(cnt.get(k, 0) == exp_por.get(k, 0) for k in RESULTADOS_C3)
                 and all(cnt.get(k, 0) == cnt_proy.get(k, 0) for k in RESULTADOS_C3))
    out.append(("por_resultado (IFC) == expected == proyección", conteo_ok, f"{dict(cnt)}"))
    if "n_elementos" in exp6d:
        out.append(("nº de elementos con Pset == esperado", len(leido) == exp6d["n_elementos"],
                    f"{len(leido)} (esperado {exp6d['n_elementos']})"))

    # 5 · veredicto agregado (cross-check con el veredicto del engine)
    if "veredicto_agregado" in exp6d:
        out.append(("veredicto agregado == esperado",
                    veredicto.get("veredicto") == exp6d["veredicto_agregado"],
                    f"{veredicto.get('veredicto')}"))
    return out


def _run_c3_6d(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
               repo: Path = DEFAULT_REPO) -> bool:
    """C3 write-back 6D (vertical visor-cumplimiento-6d, D-6D-1..4): el cumplimiento vuelve al modelo.

    El runner federa+deriva el Maestro (fixtures de GOL-CTE-01/C4-FED-06), VERIFICA con el engine C3
    y ESCRIBE `Pset_Aqyra_Cumplimiento` por elemento (engines/cumplimiento.escribir_cumplimiento).
    ANCLA (opción b, patrón 5D D14 — sin md5 hardcodeado, EOL-frágil):
      - DETERMINISMO: escribir el 6D 2× → bytes idénticos (cabecera SPF fija + GUIDs uuid5).
      - SEMÁNTICA: el Pset casa con el veredicto proyectado a elementos por peor caso (proyección
        INDEPENDIENTE del runner) + conteos anclados en el expected.
    Un fallo se investiga en el ENGINE (escritura.py), jamás aflojando el check. GOL-CTE-01 (el
    oráculo anclado del veredicto) queda intacto — este caso ancla el WRITE-BACK."""
    import tempfile

    checks: list[tuple[str, bool, str]] = []
    schemas = load_c3_schemas(contracts_dir)
    entrada = json.loads((case_dir / "entrada.json").read_text(encoding="utf-8"))
    exp6d = expected.get("cumplimiento_6d", {})

    # 1 · identidad por hash de las fixtures (reutilizadas de GOL-CTE-01, intactas)
    entradas = expected.get("entradas_md5", {})
    checks.append(("entradas declaradas", bool(entradas), f"{len(entradas)} fichero/s"))
    for rel, exp_md5 in sorted(entradas.items()):
        p = case_dir / rel
        got = _md5(p) if p.exists() else None
        checks.append((f"identidad {rel}", got == exp_md5,
                       f"md5 {str(got)[:12]}… (esperado {exp_md5[:12]}…)"))

    ps = str(repo / "services" / "federacion" / "src")
    if ps not in sys.path:
        sys.path.insert(0, ps)
    pe = str(DEFAULT_ENGINE_C3)
    if pe not in sys.path:
        sys.path.insert(0, pe)
    try:
        import ifcopenshell

        from aqyra_cumplimiento import escribir_cumplimiento, verificar
        from aqyra_cumplimiento import modelo as _CM
        from aqyra_federacion import derivar, federar_fichero
    except Exception as e:  # noqa: BLE001
        checks.append(("import engine/servicio/ifcopenshell", False, f"{type(e).__name__}: {e}"))
        print_checks(checks)
        return False

    with tempfile.TemporaryDirectory() as td:
        tdp = Path(td)
        # 2 · federar + derivar el Maestro (el engine ABRE el derivado, no federa — D7)
        manifiesto = federar_fichero(case_dir / "reglas.json")
        derivado = tdp / "federado.ifc"
        derivar(manifiesto, case_dir, derivado)
        checks.append(("federar+derivar el Maestro", derivado.exists(),
                       "derivado federado de las fixtures de GOL-CTE-01"))

        # 3 · verificar (engine C3) → veredicto ; conforma esquema
        pack = entrada.get("pack_normativo", {})
        pack_dir = (repo / "data" / "packs" / "normativa"
                    / str(pack.get("id")) / str(pack.get("version")))
        maestro = {"manifiesto": manifiesto, "base_dir": case_dir,
                   "ifc_derivado": derivado, "proyecto": entrada.get("proyecto")}
        veredicto = verificar(maestro, entrada.get("uso", {}),
                              entrada.get("localizacion", {}), pack_dir)
        ok, detail = validate_against_schema(veredicto, schemas["veredicto"])
        checks.append(("veredicto (para el 6D) conforma esquema", ok, detail))

        # 4 · escribir_cumplimiento 2× → DETERMINISMO (D-6D-2, opción b D14)
        guid2mod = _CM.guid_a_modelo(manifiesto, case_dir)
        r1 = escribir_cumplimiento(veredicto, maestro, tdp / "a" / "federado_6d.ifc", guid2mod=guid2mod)
        r2 = escribir_cumplimiento(veredicto, maestro, tdp / "b" / "federado_6d.ifc", guid2mod=guid2mod)
        checks.append(("6D determinista (escribir 2× = bytes idénticos)", r1["md5"] == r2["md5"],
                       f"md5 {r1['md5'][:12]}… (== {r2['md5'][:12]}…); {r1['n_elementos']} elementos"))

        # 5 · SEMÁNTICA — el Pset casa con el veredicto proyectado por peor caso
        f = ifcopenshell.open(str(tdp / "a" / "federado_6d.ifc"))
        checks.extend(_checks_6d(f, veredicto, guid2mod, exp6d, r1))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


def run_case_c3(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                repo: Path = DEFAULT_REPO) -> bool:
    """C3: RECOMPUTE con el engine (Fase III·h3, D9) + checks anclados del 3.2 — ver ficha.

    El engine `engines/cumplimiento` (D6) regenera el Maestro (federar+derivar con el service, D7)
    y emite el veredicto POR EXIGENCIA; el runner lo ANTEPONE contra el MISMO expected (el checklist
    FIRMADO que era el oráculo del modo anclado), normalizando el texto libre (D9). Conserva íntegros
    los checks anclados del 2.1/3.2 (D10: más checks, nunca menos):

      0. RECOMPUTE (costura cerrada): engine.verificar() == expected["veredicto_cumplimiento"].
      1. Conformidad de los 2 esquemas (entrada + veredicto).
      2. Identidad por hash (entradas del C4-FED-06 reutilizadas + derivado anclado; reglas
         regeneran el Maestro).
      3. Coherencia del pack (entrada == veredicto == versions.lock; exigencias ⊆ pack CTE).
      4. Taxonomía CERRADA de resultado (D4) + motivo en cada no-verificable.
      5. resumen == recuento real.
      6. Veredicto agregado según la regla D4.
      7. uso/localización coherentes entrada↔veredicto.
      8. por_modelo ⊆ modelos de la entrada.
    """
    # Dispatch por modo (vertical 6D): GOL-CTE-6D-01 (write-back del cumplimiento) tiene su rama.
    if expected.get("modo") == "6d" or "cumplimiento_6d" in expected:
        return _run_c3_6d(case_dir, contracts_dir, expected, tol, repo)

    from collections import Counter

    checks: list[tuple[str, bool, str]] = []
    schemas = load_c3_schemas(contracts_dir)

    entrada_path = case_dir / "entrada.json"
    if not entrada_path.exists():
        print_checks([("entrada.json presente", False, "no encontrado")])
        return False
    entrada = json.loads(entrada_path.read_text(encoding="utf-8"))
    vc = expected.get("veredicto_cumplimiento", {})

    # 0 · RECOMPUTE con el engine real (Fase III·h3, D9) — se ANTEPONE al anclado.
    # Regenera el Maestro (federar+derivar) y emite el veredicto con engines/cumplimiento
    # contra el MISMO expected (D10: más checks, nunca menos). Un fallo se investiga en el engine.
    try:
        import tempfile
        with tempfile.TemporaryDirectory() as td:
            got = _recompute_c3(case_dir, entrada, repo, Path(td))
            checks.append(("engine cumplimiento verificar() ejecuta", True,
                           "engines/cumplimiento sobre el Maestro regenerado (federar+derivar)"))
            ok, detail = validate_against_schema(got, schemas["veredicto"])
            checks.append(("veredicto recomputado conforma", ok, detail))
            d_vc = _diffs(_norm_c3(vc), _norm_c3(got), 0.0, 0.0)
            checks.append(("recompute veredicto == expected", not d_vc,
                           f"{len(d_vc)} diff/s — {d_vc[0]}" if d_vc
                           else "reproducido (ids, resultados, referencias, por_modelo, resumen, veredicto)"))
    except Exception as e:  # noqa: BLE001
        checks.append(("recompute con el engine", False, f"{type(e).__name__}: {e}"))

    # 1 · conformidad de esquemas
    for name, inst, key in (("entrada conforma esquema", entrada, "entrada"),
                            ("veredicto conforma esquema", vc, "veredicto")):
        ok, detail = validate_against_schema(inst, schemas[key])
        checks.append((name, ok, detail))

    # 2 · identidad por hash (entradas del 06 + derivado anclado)
    entradas = expected.get("entradas_md5", {})
    checks.append(("entradas declaradas", bool(entradas), f"{len(entradas)} fichero/s"))
    for rel, exp_md5 in sorted(entradas.items()):
        p = case_dir / rel
        if not p.exists():
            checks.append((f"identidad {rel}", False, "fichero no encontrado"))
            continue
        got = _md5(p)
        checks.append((f"identidad {rel}", got == exp_md5,
                       f"md5 {got[:12]}… (esperado {exp_md5[:12]}…)"))
    md5_entrada = {m["fichero"]: m.get("md5")
                   for m in entrada.get("modelo", {}).get("modelos", [])}
    checks.append(("md5 entrada == entradas", md5_entrada == entradas,
                   "los md5 de entrada.json anclan los mismos ficheros"))
    der_ent = entrada.get("modelo", {}).get("ifc_derivado_md5")
    der_exp = expected.get("maestro", {}).get("ifc_derivado_md5")
    checks.append(("derivado (vista del Maestro) coherente", bool(der_ent) and der_ent == der_exp,
                   f"{str(der_ent)[:12]}… (entrada == expected)"))
    reglas_rel = entrada.get("modelo", {}).get("reglas")
    if reglas_rel:
        rp = case_dir / reglas_rel
        if not rp.exists():
            checks.append(("reglas.json presente", False, f"{reglas_rel} no encontrado"))
        else:
            reglas = json.loads(rp.read_text(encoding="utf-8"))
            md5_reglas = {m["fichero"]: m.get("md5") for m in reglas.get("modelos", [])}
            checks.append(("reglas regeneran el Maestro (md5 == entradas)", md5_reglas == entradas,
                           "las reglas de federación anclan los mismos IFC"))

    # 3 · coherencia del pack
    pack_ent = entrada.get("pack_normativo", {})
    pack_ver = vc.get("pack", {})
    mismo_pack = (pack_ent.get("id") == pack_ver.get("id")
                  and pack_ent.get("version") == pack_ver.get("version"))
    checks.append(("pack entrada == veredicto", mismo_pack,
                   f"{pack_ent.get('id')}/{pack_ent.get('version')}"))
    lock_norm = _lock_packs(repo, "normativa")
    anclado = (lock_norm.get("id") == pack_ent.get("id")
               and lock_norm.get("version") == pack_ent.get("version"))
    checks.append(("pack anclado en versions.lock", anclado,
                   f"lock={lock_norm.get('id')}/{lock_norm.get('version')}"))
    exigidos = [e.get("id") for e in vc.get("exigencias", [])]
    pack_dir = (repo / "data" / "packs" / "normativa"
                / str(pack_ent.get("id")) / str(pack_ent.get("version")))
    exig_ok, exig_detail = False, "pack.json no encontrado"
    if (pack_dir / "pack.json").exists():
        pack = json.loads((pack_dir / "pack.json").read_text(encoding="utf-8"))
        exig_file = pack_dir / pack["contenido"]["fichero"]
        declarados = {e["id"] for e in
                      json.loads(exig_file.read_text(encoding="utf-8")).get("exigencias", [])}
        faltan = set(exigidos) - declarados
        exig_ok = bool(exigidos) and not faltan
        exig_detail = (f"{sorted(exigidos)} ⊆ pack" if exig_ok
                       else f"exigencias SIN respaldo en el pack: {sorted(faltan)}")
    checks.append(("exigencias ⊆ pack CTE", exig_ok, exig_detail))

    # 4 · taxonomía CERRADA (D4) + motivo en no-verificable
    resultados = [e.get("resultado") for e in vc.get("exigencias", [])]
    checks.append(("taxonomía de resultado cerrada", all(r in RESULTADOS_C3 for r in resultados),
                   f"{sorted(set(resultados))}"))
    nv = [e for e in vc.get("exigencias", []) if e.get("resultado") == "no-verificable"]
    checks.append(("no-verificable declara motivo", all(e.get("motivo_no_verificable") for e in nv),
                   f"{len(nv)} no-verificable/s con motivo"))

    # 5 · resumen == recuento real
    cnt = Counter(resultados)
    resumen = vc.get("resumen", {})
    por = resumen.get("por_resultado", {})
    conteo_ok = (resumen.get("total") == len(resultados)
                 and all(por.get(k, 0) == cnt.get(k, 0) for k in RESULTADOS_C3)
                 and sum(por.get(k, 0) for k in RESULTADOS_C3) == len(resultados))
    checks.append(("resumen == recuento real", conteo_ok,
                   f"total {resumen.get('total')} · {dict(cnt)}"))

    # 6 · veredicto agregado (regla D4)
    if "no-cumple" in resultados:
        esperado = "no-conforme"
    elif "no-verificable" in resultados:
        esperado = "conforme-con-reservas"
    else:
        esperado = "conforme"
    checks.append(("veredicto agregado coherente", vc.get("veredicto") == esperado,
                   f"{vc.get('veredicto')} (esperado {esperado})"))

    # 7 · uso/localización coherentes entrada↔veredicto
    checks.append(("uso coherente entrada↔veredicto",
                   entrada.get("uso", {}).get("principal") == vc.get("uso", {}).get("principal"),
                   f"{vc.get('uso', {}).get('principal')}"))
    loc_e, loc_v = entrada.get("localizacion", {}), vc.get("localizacion", {})
    loc_keys = ("pais", "provincia", "municipio", "zona_climatica_he")
    checks.append(("localización coherente entrada↔veredicto",
                   all(loc_e.get(k) == loc_v.get(k) for k in loc_keys),
                   f"{loc_v.get('municipio')} ({loc_v.get('zona_climatica_he')})"))

    # 8 · por_modelo ⊆ modelos de la entrada
    ids_entrada = {m["id"] for m in entrada.get("modelo", {}).get("modelos", [])}
    pm_ids: set[str] = set()
    for e in vc.get("exigencias", []):
        pm_ids |= set((e.get("por_modelo") or {}).keys())
    checks.append(("por_modelo ⊆ modelos entrada", pm_ids <= ids_entrada,
                   f"por_modelo={sorted(pm_ids)} ⊆ {sorted(ids_entrada)}"))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


# --------------------------------------------------------------------------- #
# C5 · presupuesto — modo ANCLADO (contract-first, Fase IV·h1 — sin engine)     #
# --------------------------------------------------------------------------- #
ORIGEN_C5 = {"modelo", "regla", "manual"}

# Casa del engine C5 (hilo posterior): engines/presupuesto — se importará por path (patrón
# engines/ifc y _recompute_c4/_recompute_c3). El runner regenerará la medición desde los Qto
# (parser) y el motor dará el presupuesto; aquí, modo ANCLADO (oráculo congelado, sin engine).
DEFAULT_ENGINE_C5 = DEFAULT_REPO / "engines" / "presupuesto" / "src"

# Casa del compositor de documentos (capa de documentos): documentos/presupuesto — importado por
# path (patrón engines/*). Reproduce el .docx del despacho desde el `presupuesto` ANCLADO de
# GOL-PRE-01 (D4); el runner ancla ESTRUCTURA + CONTENIDO (D3), no bytes ni píxeles.
DEFAULT_DOCUMENTO_C5 = DEFAULT_REPO / "documentos" / "presupuesto" / "src"

# Casa del compositor de pliego (capa de documentos, Ola 3): documentos/pliego — importado por
# path (patron engines/*). Reproduce el .docx del pliego desde el `presupuesto` ANCLADO de
# GOL-PRE-01 (D6) + criterio + pack de textos; ancla ESTRUCTURA + CONTENIDO (patron GOL-DOC-01/D3).
DEFAULT_PLIEGO_C5 = DEFAULT_REPO / "documentos" / "pliego" / "src"
DEFAULT_EXPORT_C5 = DEFAULT_REPO / "documentos" / "export" / "src"      # E6.2 · export firmable
DEFAULT_COMUN_DOC = DEFAULT_REPO / "documentos" / "comun" / "src"       # formato compartido (D-EX-1)
DEFAULT_EXPORT_ESQUEMAS = DEFAULT_REPO / "documentos" / "export" / "esquemas"


def load_c5_schemas(contracts_dir: Path) -> dict[str, dict]:
    base = contracts_dir / "C5-presupuesto"
    out = {}
    for key, fname in (("entrada", "entrada-presupuesto.schema.json"),
                       ("salida", "salida-presupuesto.schema.json")):
        out[key] = json.loads((base / fname).read_text(encoding="utf-8"))
    return out


def _pack_c5(repo: Path, familia: str, ref: dict) -> tuple[dict, Path]:
    """Carga el pack.json de un pack C5 (criterio/banco) y devuelve (manifiesto, dir)."""
    pack_dir = (repo / "data" / "packs" / familia
                / str(ref.get("id")) / str(ref.get("version")))
    return json.loads((pack_dir / "pack.json").read_text(encoding="utf-8")), pack_dir


def _pack_contenido_c5(repo: Path, familia: str, ref: dict) -> dict:
    """Contenido (criterio.json / banco.json) de un pack C5, por su manifiesto."""
    man, pdir = _pack_c5(repo, familia, ref)
    return json.loads((pdir / man["contenido"]["fichero"]).read_text(encoding="utf-8"))


# Texto libre del presupuesto C5 (presentación humana): se comprueba por esquema (presencia/tipo)
# pero NO se compara literalmente en el recompute (D9). La semántica del contrato son códigos,
# cantidades, precios, importes, capítulos y trazabilidad; descripciones y justificaciones (incluida
# la letra del nº1) son presentación.
_LIBRES_C5 = {"descripcion", "criterio_aplicado", "precio_en_letra", "nota", "proyecto"}


def _recompute_c5(case_dir: Path, entrada: dict, repo: Path) -> tuple[dict, dict]:
    """Cierra la costura C5 (Fase IV·h2, D9): el engine engines/presupuesto MIDE desde las fixtures
    con Qto (parser, módulo 1 — la medición NACE del modelo, D7) y PRESUPUESTA (motor). El runner lo
    ANTEPONE contra el MISMO expected. Un fallo aquí se investiga en el ENGINE, jamás en expected."""
    pe = str(DEFAULT_ENGINE_C5)
    if pe not in sys.path:
        sys.path.insert(0, pe)
    from aqyra_presupuesto import medir, presupuestar
    fuentes = [
        {"id": m.get("id"), "disciplina": m.get("disciplina"),
         "path": case_dir / m["fichero"], "fichero": m["fichero"]}
        for m in entrada.get("modelo", {}).get("fuente_maestro", {}).get("modelos", [])
    ]
    modelo = medir(fuentes)
    modelo["proyecto"] = entrada.get("proyecto")
    criterio = _pack_contenido_c5(repo, "criterio", entrada.get("criterio_ref", {}))
    banco = _pack_contenido_c5(repo, "banco", entrada.get("banco_ref", {}))
    return modelo, presupuestar(modelo, criterio, banco, entrada.get("parametros", {}))


def _diff_medicion_c5(medido: dict, entrada_modelo: dict, tol: dict) -> list[str]:
    """El parser reproduce las cantidades que el modelo neutro declara (la medición nace del Qto)."""
    rel = float(tol.get("cantidad_rel", 0.005))
    got = {o.get("guid"): o for o in medido.get("objetos", [])}
    diffs: list[str] = []
    for obj in entrada_modelo.get("objetos", []):
        gd = obj.get("guid")
        if gd not in got:
            diffs.append(f"objeto {gd} no medido por el parser")
            continue
        gmag = {c.get("magnitud"): c.get("valor") for c in got[gd].get("cantidades", [])}
        for c in obj.get("cantidades", []):
            mag = c.get("magnitud")
            if mag == "conteo":
                continue
            if mag not in gmag:
                diffs.append(f"{gd}: falta magnitud {mag}")
            elif abs(float(gmag[mag]) - float(c.get("valor", 0))) > rel * max(1.0, abs(float(c.get("valor", 0)))):
                diffs.append(f"{gd}.{mag} medido {gmag[mag]} != {c.get('valor')}")
    return diffs


def _diff_presupuesto_c5(got: dict, exp: dict, tol: dict) -> list[str]:
    """Compara el presupuesto recomputado contra el expected (D3): cantidades ±0,5%; mapeo, precios,
    importes y PEM…PEC EXACTOS (±0,01). Normaliza el texto libre (D9)."""
    rel = float(tol.get("cantidad_rel", 0.005))
    ia = float(tol.get("importe_abs", 0.01))
    diffs: list[str] = []

    def approx(a, b, t) -> bool:
        try:
            return abs(float(a) - float(b)) <= t
        except (TypeError, ValueError):
            return False

    g = {p["codigo"]: p for p in got.get("estado_mediciones", [])}
    e = {p["codigo"]: p for p in exp.get("estado_mediciones", [])}
    if set(g) != set(e):
        diffs.append(f"partidas: faltan {sorted(set(e) - set(g))} / sobran {sorted(set(g) - set(e))}")
    for cod in sorted(set(e) & set(g)):
        a, b = g[cod], e[cod]
        for k in ("capitulo", "unidad", "origen"):
            if a.get(k) != b.get(k):
                diffs.append(f"{cod}.{k}: {a.get(k)} != {b.get(k)}")
        if not approx(a.get("cantidad"), b.get("cantidad"), rel * max(1.0, abs(float(b.get("cantidad", 0))))):
            diffs.append(f"{cod}.cantidad: {a.get('cantidad')} != {b.get('cantidad')}")
        if not approx(a.get("precio_unitario"), b.get("precio_unitario"), ia):
            diffs.append(f"{cod}.precio_unitario: {a.get('precio_unitario')} != {b.get('precio_unitario')}")
        if not approx(a.get("importe"), b.get("importe"), ia):
            diffs.append(f"{cod}.importe: {a.get('importe')} != {b.get('importe')}")
        if set(a.get("trazabilidad", [])) != set(b.get("trazabilidad", [])):
            diffs.append(f"{cod}.trazabilidad != expected")

    g1 = {c["codigo"]: c for c in got.get("cuadro_precios_1", [])}
    for cod, c in {c["codigo"]: c for c in exp.get("cuadro_precios_1", [])}.items():
        if cod not in g1:
            diffs.append(f"nº1 falta {cod}")
        elif not approx(g1[cod].get("precio"), c.get("precio"), ia) or g1[cod].get("unidad") != c.get("unidad"):
            diffs.append(f"nº1 {cod}: precio/unidad")

    g2 = {c["codigo"]: c for c in got.get("cuadro_precios_2", [])}
    for cod, c in {c["codigo"]: c for c in exp.get("cuadro_precios_2", [])}.items():
        if cod not in g2:
            diffs.append(f"nº2 falta {cod}")
            continue
        if not approx(g2[cod].get("precio_total"), c.get("precio_total"), ia):
            diffs.append(f"nº2 {cod}.precio_total")
        if not approx(g2[cod].get("costes_indirectos"), c.get("costes_indirectos"), ia):
            diffs.append(f"nº2 {cod}.costes_indirectos")
        gc, ec = g2[cod].get("componentes", []), c.get("componentes", [])
        if len(gc) != len(ec):
            diffs.append(f"nº2 {cod}.componentes: {len(gc)} != {len(ec)}")
        else:
            for i, (x, y) in enumerate(zip(gc, ec)):
                if x.get("tipo") != y.get("tipo") or not approx(x.get("subtotal"), y.get("subtotal"), ia):
                    diffs.append(f"nº2 {cod}.comp[{i}]")

    gr, er = got.get("resumen", {}), exp.get("resumen", {})
    for k in ("PEM", "gg", "bi", "base", "iva", "PEC", "gg_pct", "bi_pct", "iva_pct"):
        if not approx(gr.get(k), er.get(k), ia):
            diffs.append(f"resumen.{k}: {gr.get(k)} != {er.get(k)}")
    if gr.get("moneda") != er.get("moneda"):
        diffs.append("resumen.moneda")
    gcp = {c["codigo"]: c for c in gr.get("capitulos", [])}
    for cod, c in {c["codigo"]: c for c in er.get("capitulos", [])}.items():
        if cod not in gcp:
            diffs.append(f"capítulo falta {cod}")
            continue
        if not approx(gcp[cod].get("importe"), c.get("importe"), ia):
            diffs.append(f"capítulo {cod}.importe")
        if set(gcp[cod].get("partidas", [])) != set(c.get("partidas", [])):
            diffs.append(f"capítulo {cod}.partidas")
    return diffs


def _checks_5d(f, pres: dict, cost: dict, ia: float) -> list[tuple[str, bool, str]]:
    """Comprueba que el cost schedule del IFC 5D CASA con el presupuesto (D14, semántica)."""
    out: list[tuple[str, bool, str]] = []

    def approx(a, b) -> bool:
        try:
            return abs(float(a) - float(b)) <= ia
        except (TypeError, ValueError):
            return False

    def _val(cv):
        av = cv.AppliedValue
        return float(getattr(av, "wrappedValue", av))

    scheds = f.by_type("IfcCostSchedule")
    out.append(("IfcCostSchedule (BUDGET) presente",
                len(scheds) == 1 and getattr(scheds[0], "PredefinedType", None) == cost.get("predefined_type", "BUDGET"),
                f"{len(scheds)} schedule/s"))
    munits = f.by_type("IfcMonetaryUnit")
    out.append(("IfcMonetaryUnit == moneda",
                any(getattr(u, "Currency", None) == cost.get("moneda", "EUR") for u in munits),
                f"{[getattr(u, 'Currency', None) for u in munits]}"))

    items = {ci.Identification: ci for ci in f.by_type("IfcCostItem")
             if getattr(ci, "Identification", None)}
    em = pres.get("estado_mediciones", [])

    malos = []
    for l in em:
        ci = items.get(l["codigo"])
        if ci is None or not ci.CostValues or not approx(_val(ci.CostValues[0]), l["importe"]):
            malos.append(l["codigo"])
    out.append(("IfcCostItem por partida con IfcCostValue == importe", not malos,
                "todas las partidas" if not malos else f"descuadran: {malos}"))

    asg: dict[str, set] = {}
    for rel in f.by_type("IfcRelAssignsToControl"):
        cod = getattr(rel.RelatingControl, "Identification", None)
        if not cod:
            continue
        guids = set()
        for o in rel.RelatedObjects or []:
            try:
                guids.add(o.GlobalId)
            except Exception:  # noqa: BLE001
                pass
        asg[cod] = guids
    mal_asg, n_asg = [], 0
    for l in em:
        tz = set(l.get("trazabilidad", []))
        if not tz:
            continue
        n_asg += 1
        if asg.get(l["codigo"], set()) != tz:
            mal_asg.append(l["codigo"])
    out.append(("IfcRelAssignsToControl == trazabilidad por partida (por GUID)", not mal_asg,
                f"{n_asg} partidas asignadas a sus elementos" if not mal_asg else f"descuadran: {mal_asg}"))
    if "n_partidas_con_asignacion" in cost:
        out.append(("nº de partidas con asignación", n_asg == cost["n_partidas_con_asignacion"],
                    f"{n_asg} (esperado {cost['n_partidas_con_asignacion']})"))

    suma = sum(_val(items[l["codigo"]].CostValues[0]) for l in em
               if items.get(l["codigo"]) and items[l["codigo"]].CostValues)
    PEM = float(pres.get("resumen", {}).get("PEM", 0))
    out.append(("Σ IfcCostValue de partida == PEM", approx(suma, PEM),
                f"Σ {round(suma, 2)} vs PEM {PEM}"))

    res = pres.get("resumen", {})
    res_item = items.get("RESUMEN")
    cat = {}
    if res_item:
        for cv in res_item.CostValues or []:
            cat[str(getattr(cv, "Category", None)).upper()] = _val(cv)
    claves = {"PEM": "PEM", "GG": "gg", "BI": "bi", "BASE": "base", "IVA": "iva", "PEC": "PEC"}
    mal_res = [k for k, rk in claves.items() if rk in res and not approx(cat.get(k), res[rk])]
    out.append(("IfcCostItem resumen == PEM/GG/BI/base/IVA/PEC",
                res_item is not None and not mal_res,
                "coherente" if res_item is not None and not mal_res else f"descuadran: {mal_res}"))

    n_items = len(f.by_type("IfcCostItem"))
    if "n_partidas" in cost and "n_capitulos" in cost:
        esperado = cost["n_partidas"] + cost["n_capitulos"] + 1
        out.append(("nº de IfcCostItem == partidas+capítulos+resumen", n_items == esperado,
                    f"{n_items} (esperado {esperado})"))
    return out


def _run_c5_5d(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict, repo: Path) -> bool:
    """C5 write-back 5D (Fase IV·h3, D11–D14): el coste vuelve al modelo. Ancla por DETERMINISMO
    (escribir 2× = bytes idénticos) + SEMÁNTICA (el cost schedule casa con el presupuesto). Opción b:
    sin md5 hardcodeado — el gate genera y verifica el 5D con ifcopenshell. Un fallo se investiga en
    el ENGINE (escritura.py), jamás aflojando el check."""
    import tempfile

    checks: list[tuple[str, bool, str]] = []
    schemas = load_c5_schemas(contracts_dir)
    entrada = json.loads((case_dir / "entrada.json").read_text(encoding="utf-8"))
    ia = float(tol.get("importe_abs", 0.01))
    cost = expected.get("cost_5d", {})

    # 1 · identidad por hash de las fixtures con Qto (reutilizadas del 4.1; intactas)
    for rel, exp_md5 in sorted(expected.get("entradas_md5", {}).items()):
        p = case_dir / rel
        got = _md5(p) if p.exists() else None
        checks.append((f"identidad {rel}", got == exp_md5,
                       f"md5 {str(got)[:12]}… (esperado {exp_md5[:12]}…)"))

    pe = str(DEFAULT_ENGINE_C5)
    if pe not in sys.path:
        sys.path.insert(0, pe)
    ps = str(repo / "services" / "federacion" / "src")
    if ps not in sys.path:
        sys.path.insert(0, ps)
    try:
        import ifcopenshell

        from aqyra_federacion import derivar, federar_fichero
        from aqyra_presupuesto import escribir_coste, medir, presupuestar
    except Exception as e:  # noqa: BLE001
        checks.append(("import engine/servicio/ifcopenshell", False, f"{type(e).__name__}: {e}"))
        print_checks(checks)
        return False

    # 2 · presupuesto (reproduce GOL-PRE-01) desde la medición de las fixtures con Qto
    fuentes = [{"id": m.get("id"), "disciplina": m.get("disciplina"),
                "path": case_dir / m["fichero"], "fichero": m["fichero"]}
               for m in entrada.get("modelo", {}).get("fuente_maestro", {}).get("modelos", [])]
    modelo = medir(fuentes)
    modelo["proyecto"] = entrada.get("proyecto")
    criterio = _pack_contenido_c5(repo, "criterio", entrada.get("criterio_ref", {}))
    banco = _pack_contenido_c5(repo, "banco", entrada.get("banco_ref", {}))
    pres = presupuestar(modelo, criterio, banco, entrada.get("parametros", {}))
    ok, detail = validate_against_schema(pres, schemas["salida"])
    checks.append(("presupuesto (para el 5D) conforma esquema", ok, detail))
    checks.append(("presupuesto reproduce PEM…PEC del oráculo",
                   abs(float(pres.get("resumen", {}).get("PEM", 0)) - float(cost.get("PEM", 0))) <= ia
                   and abs(float(pres.get("resumen", {}).get("PEC", 0)) - float(cost.get("PEC", 0))) <= ia,
                   f"PEM {pres.get('resumen', {}).get('PEM')} PEC {pres.get('resumen', {}).get('PEC')}"))

    with tempfile.TemporaryDirectory() as td:
        tdp = Path(td)
        # 3 · federar + derivar el Maestro (el engine ABRE el derivado, no federa — D7)
        manifiesto = federar_fichero(case_dir / "reglas.json", base_dir=case_dir)
        derivado = tdp / "federado.ifc"
        derivar(manifiesto, case_dir, derivado)
        checks.append(("federar+derivar el Maestro", derivado.exists(),
                       "derivado federado de las fixtures con Qto"))
        # 4 · escribir_coste 2× → DETERMINISMO (D14); mismo basename, dirs distintas
        r1 = escribir_coste(pres, derivado, tdp / "a" / "federado_5d.ifc")
        r2 = escribir_coste(pres, derivado, tdp / "b" / "federado_5d.ifc")
        checks.append(("5D determinista (escribir 2× = bytes idénticos)", r1["md5"] == r2["md5"],
                       f"md5 {r1['md5'][:12]}… (== {r2['md5'][:12]}…); {r1['n_cost_items']} cost items, "
                       f"{r1['n_asignaciones']} asignaciones"))
        # 5 · SEMÁNTICA — el cost schedule casa con el presupuesto
        f = ifcopenshell.open(str(tdp / "a" / "federado_5d.ifc"))
        checks.extend(_checks_5d(f, pres, cost, ia))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


# --------------------------------------------------------------------------- #
# C5 · Documento de Presupuesto — capa de documentos (compositor DETERMINISTA)  #
# --------------------------------------------------------------------------- #
def _num_es(s: str):
    """'1.234,56 EUR' → 1234.56 ; texto sin dígitos → None. Formato español del despacho."""
    s = (s or "").replace("EUR", "").replace("€", "").strip()
    if not any(c.isdigit() for c in s):
        return None
    s = s.replace(".", "").replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


def _extraer_docx(path: Path) -> tuple[list[str], list]:
    """(párrafos, tablas) del .docx con python-docx — NO OCR (D3). Cada tabla es lista de filas
    de celdas (texto). Es la evidencia sobre la que se ancla la conformidad."""
    from docx import Document
    d = Document(str(path))
    paras = [p.text for p in d.paragraphs]
    tablas = [[[c.text for c in row.cells] for row in t.rows] for t in d.tables]
    return paras, tablas


def _fmt_es(x, dec: int = 2) -> str:
    """1234.5 -> '1.234,50' (miles '.', decimales ',') — identico a formato.fmt_num del despacho."""
    s = f"{float(x):,.{dec}f}"
    return s.replace(",", "\x00").replace(".", ",").replace("\x00", ".")


def _xlsx_numeros(path: Path) -> list:
    """Todos los valores numericos de todas las hojas (para comprobar cifras)."""
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for c in row:
                if isinstance(c, (int, float)):
                    out.append(float(c))
    return out


def _xlsx_celdas(path: Path):
    """Contenido (por hoja) del xlsx para comparar determinismo."""
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    return {ws.title: [list(r) for r in ws.iter_rows(values_only=True)] for ws in wb.worksheets}


def _pdf_texto(path: Path) -> str:
    """Texto extraido del PDF (pypdf, NO OCR) — la evidencia del sellado (sin pixeles)."""
    from pypdf import PdfReader
    return "\n".join((pg.extract_text() or "") for pg in PdfReader(str(path)).pages)


def _run_c5_documento(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                      repo: Path) -> bool:
    """C5 · Documento de Presupuesto (capa de documentos, D1–D5): el compositor DETERMINISTA
    documentos/presupuesto reproduce el .docx del despacho desde el `presupuesto` ANCLADO de la
    fuente (D4, GOL-PRE-01 — se LEE, no se re-ancla). Ancla ESTRUCTURA + CONTENIDO (D3), no bytes
    ni píxeles. Un fallo se investiga en el COMPOSITOR, jamás aflojando la golden."""
    import tempfile

    checks: list[tuple[str, bool, str]] = []
    schemas = load_c5_schemas(contracts_dir)
    ia = float(tol.get("importe_abs", 0.01))
    doc_exp = expected.get("documento", {})
    fecha = doc_exp.get("fecha", "-")

    # 1 · cargar el `presupuesto` de la fuente anclada (D4) y comprobar que conforma el esquema
    fuente = expected.get("fuente_presupuesto")
    fuente_path = case_dir.parent / str(fuente) / "expected.json"
    if not fuente_path.exists():
        print_checks([("fuente de presupuesto presente", False, f"{fuente} no encontrada")])
        return False
    presupuesto = json.loads(fuente_path.read_text(encoding="utf-8")).get("presupuesto", {})
    ok, detail = validate_against_schema(presupuesto, schemas["salida"])
    checks.append((f"presupuesto fuente ({fuente}) conforma salida-presupuesto", ok, detail))

    # 2 · componer el .docx con el compositor (import de path, patrón engines/*)
    pd = str(DEFAULT_DOCUMENTO_C5)
    if pd not in sys.path:
        sys.path.insert(0, pd)
    try:
        import aqyra_documento_presupuesto as adp
    except Exception as e:  # noqa: BLE001
        checks.append(("import del compositor documentos/presupuesto", False,
                       f"{type(e).__name__}: {e}"))
        print_checks(checks)
        return False

    with tempfile.TemporaryDirectory() as td:
        tdp = Path(td)
        salida_a = adp.componer_documento(presupuesto, {"salida": tdp / "a.docx", "fecha": fecha})
        checks.append(("compositor genera el .docx",
                       salida_a.exists() and salida_a.stat().st_size > 0,
                       f"{salida_a.name} ({salida_a.stat().st_size} bytes)"))
        paras, tablas = _extraer_docx(salida_a)
        texto = "\n".join(paras)

        # 3 · las 5 secciones presentes
        marcas = [("carátula", adp.SEC_CARATULA), ("estado_mediciones", adp.SEC_MEDICIONES),
                  ("cuadro_precios_1", adp.SEC_CP1), ("cuadro_precios_2", adp.SEC_CP2),
                  ("resumen", adp.SEC_RESUMEN)]
        faltan_sec = [n for n, m in marcas if m not in texto]
        checks.append(("las 5 secciones presentes", not faltan_sec,
                       "carátula · mediciones · nº1 · nº2 · resumen" if not faltan_sec
                       else f"faltan: {faltan_sec}"))

        # 4 · las N partidas en la tabla de mediciones, importe == JSON (±ia)
        em = {p["codigo"]: p for p in presupuesto.get("estado_mediciones", [])}
        imp_doc: dict[str, float] = {}
        for t in tablas:
            for fila in t:
                cod = (fila[0] or "").strip()
                val = _num_es(fila[-1]) if fila else None
                if cod in em and val is not None and cod not in imp_doc:
                    imp_doc[cod] = val
        faltan_part = [c for c in em if c not in imp_doc]
        checks.append(("todas las partidas en la tabla de mediciones", not faltan_part,
                       f"{len(em)} partidas" if not faltan_part else f"faltan: {faltan_part}"))
        if "n_partidas" in doc_exp:
            checks.append(("nº de partidas == esperado", len(em) == doc_exp["n_partidas"],
                           f"{len(em)} (esperado {doc_exp['n_partidas']})"))
        malos_imp = [c for c in em
                     if c in imp_doc and abs(imp_doc[c] - float(em[c]["importe"])) > ia]
        checks.append(("importe de partida en el documento == JSON (±0,01)", not malos_imp,
                       "todos casan" if not malos_imp else f"descuadran: {malos_imp}"))

        # 5 · Σ capítulos == PEM y cada importe de capítulo == JSON (tabla resumen por capítulos)
        resumen = presupuesto.get("resumen", {})
        caps = {c["codigo"]: float(c.get("importe", 0)) for c in resumen.get("capitulos", [])}
        cap_doc: dict[str, float] = {}
        for t in tablas:
            cab = [c.strip() for c in t[0]] if t else []
            if cab[:1] == ["Capítulo"]:
                for fila in t[1:]:
                    cod = (fila[0] or "").strip()
                    val = _num_es(fila[-1])
                    if cod in caps and val is not None:
                        cap_doc[cod] = val
        malos_cap = [c for c in caps if c not in cap_doc or abs(cap_doc[c] - caps[c]) > ia]
        checks.append(("importe de capítulo en el documento == JSON", not malos_cap,
                       f"{len(caps)} capítulos" if not malos_cap else f"descuadran: {malos_cap}"))
        if "n_capitulos" in doc_exp:
            checks.append(("nº de capítulos == esperado", len(caps) == doc_exp["n_capitulos"],
                           f"{len(caps)} (esperado {doc_exp['n_capitulos']})"))
        PEM = float(resumen.get("PEM", 0))
        suma_cap = round(sum(cap_doc.values()), 2)
        checks.append(("Σ capítulos (del documento) == PEM", abs(suma_cap - PEM) <= ia,
                       f"Σ {suma_cap} vs PEM {PEM}"))

        # 6 · cadena PEM→GG→BI→base→IVA→PEC presente y == JSON (tabla 'Concepto')
        concep: dict[str, float] = {}
        for t in tablas:
            cab = [c.strip() for c in t[0]] if t else []
            if cab[:1] == ["Concepto"]:
                for fila in t[1:]:
                    concep[(fila[0] or "").strip()] = _num_es(fila[-1])
        cadena = {"PEM": "(PEM)", "gg": "Gastos generales", "bi": "Beneficio industrial",
                  "base": "Base imponible", "iva": "IVA", "PEC": "(PEC)"}
        malos_cad = []
        for k, marca in cadena.items():
            val = next((v for kk, v in concep.items() if marca in kk), None)
            if val is None or abs(val - float(resumen.get(k, 0))) > ia:
                malos_cad.append(k)
        checks.append(("cadena PEM→GG→BI→base→IVA→PEC == JSON", not malos_cad,
                       "coherente" if not malos_cad else f"descuadran: {malos_cad}"))

        # 7 · precio EN LETRA por partida (cuadro nº1) presente
        letras: dict[str, str] = {}
        for t in tablas:
            cab = [c.strip() for c in t[0]] if t else []
            if "Precio en letra" in cab:
                for fila in t[1:]:
                    letras[(fila[0] or "").strip()] = (fila[-1] or "").strip()
        cp1 = [c["codigo"] for c in presupuesto.get("cuadro_precios_1", [])]
        sin_letra = [c for c in cp1 if not letras.get(c)]
        checks.append(("precio en letra por partida (cuadro nº1)", not sin_letra,
                       f"{len(cp1)} partidas en letra" if not sin_letra
                       else f"sin letra: {sin_letra}"))

        # 8 · DETERMINISMO por contenido (componer 2× = texto/tablas idénticos)
        salida_b = adp.componer_documento(presupuesto, {"salida": tdp / "b.docx", "fecha": fecha})
        contenido_b = _extraer_docx(salida_b)
        checks.append(("documento determinista (2× = mismo contenido extraído)",
                       (paras, tablas) == contenido_b,
                       "texto y tablas idénticos" if (paras, tablas) == contenido_b
                       else "el contenido difiere entre dos composiciones"))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


def _cargar_pack_criterio(repo: Path, ref: str) -> dict:
    """`AQ/v2` -> data/packs/criterio/AQ/v2/criterio.json (dict) o {} si no esta."""
    ident, _, version = (ref or "").partition("/")
    p = repo / "data" / "packs" / "criterio" / ident / version / "criterio.json"
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else {}


def _cargar_pack_textos(repo: Path, ref: str) -> dict:
    """`pliego-textos/AQ-DEMO/v1` -> data/packs/pliego-textos/AQ-DEMO/v1/textos.json (dict) o {}."""
    p = repo / "data" / "packs" / str(ref) / "textos.json"
    return json.loads(p.read_text(encoding="utf-8")) if p.exists() else {}


def _run_c5_pliego(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                   repo: Path = DEFAULT_REPO) -> bool:
    """C5 · Pliego de Condiciones Tecnicas (capa de documentos, E4.1): el compositor DETERMINISTA
    documentos/pliego reproduce el .docx del despacho desde el `presupuesto` ANCLADO de la fuente
    (D6, GOL-PRE-01 — se LEE, no se re-ancla), el criterio (mapeo partida->sistema) y el pack de
    textos de prescripcion. Ancla ESTRUCTURA + CONTENIDO (patron GOL-DOC-01/D3), no bytes ni pixeles.
    Cierra el trio coste + carbono + prescripcion. Un fallo se investiga en el COMPOSITOR, jamas
    aflojando la golden."""
    import tempfile

    checks: list[tuple[str, bool, str]] = []
    schemas = load_c5_schemas(contracts_dir)
    ia = float(tol.get("importe_abs", 0.01))
    ca = float(tol.get("cantidad_abs", 0.001))
    doc_exp = expected.get("documento", {})
    fecha = doc_exp.get("fecha", "-")

    # 1 · presupuesto de la fuente anclada (D6) + conformidad de esquema
    fuente = expected.get("fuente_presupuesto")
    fuente_path = case_dir.parent / str(fuente) / "expected.json"
    if not fuente_path.exists():
        print_checks([("fuente de presupuesto presente", False, f"{fuente} no encontrada")])
        return False
    presupuesto = json.loads(fuente_path.read_text(encoding="utf-8")).get("presupuesto", {})
    ok, detail = validate_against_schema(presupuesto, schemas["salida"])
    checks.append((f"presupuesto fuente ({fuente}) conforma salida-presupuesto", ok, detail))

    # criterio + pack de textos declarados por el caso
    criterio = _cargar_pack_criterio(repo, expected.get("criterio", ""))
    textos = _cargar_pack_textos(repo, expected.get("pack_textos", ""))
    checks.append(("criterio + pack de textos cargados",
                   bool(criterio) and bool(textos.get("por_partida")),
                   f"criterio={expected.get('criterio')} textos={expected.get('pack_textos')}"))

    # 2 · componer el .docx (import por path, patron engines/*)
    pp = str(DEFAULT_PLIEGO_C5)
    if pp not in sys.path:
        sys.path.insert(0, pp)
    try:
        import aqyra_documento_pliego as apl
    except Exception as e:  # noqa: BLE001
        checks.append(("import del compositor documentos/pliego", False, f"{type(e).__name__}: {e}"))
        print_checks(checks)
        return False

    with tempfile.TemporaryDirectory() as td:
        tdp = Path(td)
        par = {"salida": tdp / "a.docx", "fecha": fecha, "pack_textos": textos}
        salida_a = apl.componer_pliego(presupuesto, criterio, par)
        checks.append(("compositor genera el .docx",
                       salida_a.exists() and salida_a.stat().st_size > 0,
                       f"{salida_a.name} ({salida_a.stat().st_size} bytes)"))
        paras, tablas = _extraer_docx(salida_a)
        texto = "\n".join(paras)

        # 3 · secciones presentes
        marcas = [("caratula", apl.SEC_CARATULA), ("condiciones_generales", apl.SEC_GENERALES),
                  ("prescripciones_particulares", apl.SEC_PARTICULARES), ("trazabilidad", apl.SEC_TRAZA)]
        faltan_sec = [n for n, m in marcas if m not in texto]
        checks.append(("las secciones presentes", not faltan_sec,
                       "caratula · generales · particulares · trazabilidad" if not faltan_sec
                       else f"faltan: {faltan_sec}"))

        # fichas por partida (tablas Concepto/Valor con fila 'Partida')
        em = {p["codigo"]: p for p in presupuesto.get("estado_mediciones", [])}
        fichas: dict[str, dict] = {}
        for t in tablas:
            cab = [c.strip() for c in t[0]] if t else []
            if cab[:2] != ["Concepto", "Valor"]:
                continue
            d = {(f[0] or "").strip(): (f[1] or "").strip() for f in t[1:]}
            cod = d.get("Partida")
            if cod:
                fichas[cod] = d

        # 4 · todas las partidas presentes con prescripcion (sin fallback)
        faltan_part = [c for c in em if c not in fichas]
        checks.append(("todas las partidas en el pliego", not faltan_part,
                       f"{len(em)} partidas" if not faltan_part else f"faltan: {faltan_part}"))
        if "n_partidas" in doc_exp:
            checks.append(("nº de partidas == esperado", len(em) == doc_exp["n_partidas"],
                           f"{len(em)} (esperado {doc_exp['n_partidas']})"))
        pres_lineas = [x for x in paras if x.startswith("Prescripcion.")]
        con_fallback = [x for x in pres_lineas if "pendiente de completar" in x]
        checks.append(("prescripcion por partida (texto base, sin fallback)",
                       len(pres_lineas) >= len(em) and not con_fallback,
                       f"{len(pres_lineas)} prescripciones" if not con_fallback
                       else f"{len(con_fallback)} cayeron al fallback"))

        # 5 · medicion (cantidad) y 6 · coste (importe) por partida
        malos_cant, malos_imp = [], []
        for c, p in em.items():
            f = fichas.get(c, {})
            # 'Cantidad medida' = "0,128 m3": tomar el token numerico inicial (la unidad, p.ej.
            # 'm3', lleva digito y no se puede limpiar con _num_es); el importe acaba en EUR (lo pela).
            cant_tok = (f.get("Cantidad medida", "") or "").split()
            vc = _num_es(cant_tok[0]) if cant_tok else None
            vi = _num_es(f.get("Importe (coste)", ""))
            if vc is None or abs(vc - float(p["cantidad"])) > ca:
                malos_cant.append(c)
            if vi is None or abs(vi - float(p["importe"])) > ia:
                malos_imp.append(c)
        checks.append(("cantidad de partida en el pliego == JSON (±0,001)", not malos_cant,
                       "todas casan" if not malos_cant else f"descuadran: {malos_cant}"))
        checks.append(("importe de partida en el pliego == JSON (±0,01)", not malos_imp,
                       "todos casan" if not malos_imp else f"descuadran: {malos_imp}"))

        # 7 · trazabilidad: GUIDs de cada partida origen=modelo presentes
        blob = texto + "\n" + "\n".join(" ".join(c for fila in t for c in fila) for t in tablas)
        sin_guid = []
        for c, p in em.items():
            for g in p.get("trazabilidad", []):
                if g not in blob:
                    sin_guid.append((c, g))
        checks.append(("trazabilidad (GUIDs) al modelo por partida", not sin_guid,
                       "GUIDs presentes" if not sin_guid else f"faltan: {sin_guid[:3]}"))

        # 8 · carbono forward-open: coherente con lo que traiga el presupuesto fuente
        trae_carbono = any((p.get("valores") or {}).get("carbono") for p in em.values())
        hay_huella = "Huella de carbono" in blob
        checks.append(("carbono forward-open (presente sii el presupuesto lo trae)",
                       trae_carbono == hay_huella,
                       f"presupuesto={'con' if trae_carbono else 'sin'} carbono · pliego={'con' if hay_huella else 'sin'} huella"))

        # 9 · DETERMINISMO por contenido (componer 2× = idem)
        par_b = {"salida": tdp / "b.docx", "fecha": fecha, "pack_textos": textos}
        contenido_b = _extraer_docx(apl.componer_pliego(presupuesto, criterio, par_b))
        checks.append(("pliego determinista (2× = mismo contenido extraido)",
                       (paras, tablas) == contenido_b,
                       "texto y tablas identicos" if (paras, tablas) == contenido_b
                       else "el contenido difiere entre dos composiciones"))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


def _run_c5_proyeccion(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                       repo: Path = DEFAULT_REPO) -> bool:
    """C5 · proyección (E2.2, D24–D29): la vista `proyectar(eje, corte)` — GOL-PRE-03.

    El corte es CONSULTA, no cálculo: se mide y presupuesta como GOL-PRE-01 (con `criterio/AQ/v2`
    para el *fallback* funcional) y se PROYECTA por eje/corte. ANCLA (patrón GOL-PRE-02/D14, sin md5
    de salida hardcodeado):
      1. Identidad por md5 de las fixtures aumentadas (entradas_md5).
      2. `criterio/AQ/v2` anclado por su content_sha256 (golden de pack de E2.1); banco == v1.
      3. DETERMINISMO: proyectar 2× por vista → misma salida.
      4. INVARIANTE: Σ valor_total == Σ estado_mediciones del eje (== PEM en coste), ±importe_abs.
      5. SEMÁNTICA: las CINCO vistas del expected (grupos, valor_total, fuente) casan con la proyección.
    Un fallo se investiga en el ENGINE (proyeccion.py/presupuesto.py), jamás aflojando el check."""
    checks: list[tuple[str, bool, str]] = []
    schemas = load_c5_schemas(contracts_dir)
    entrada = json.loads((case_dir / "entrada.json").read_text(encoding="utf-8"))
    ia = float(tol.get("importe_abs", 0.01))
    cost = expected.get("cost", {})
    vistas = expected.get("vistas", [])

    # 1 · identidad por hash de las fixtures aumentadas (nuevas; GOL-PRE-01 intacta)
    for rel, exp_md5 in sorted(expected.get("entradas_md5", {}).items()):
        p = case_dir / rel
        got = _md5(p) if p.exists() else None
        checks.append((f"identidad {rel}", got == exp_md5,
                       f"md5 {str(got)[:12]}… (esperado {str(exp_md5)[:12]}…)"))

    # 2 · criterio/AQ/v2 anclado por su content_sha256 (E2.1); banco AQ-DEMO/v1
    crit_ref = entrada.get("criterio_ref", {})
    checks.append(("criterio_ref == AQ/v2", crit_ref.get("id") == "AQ" and crit_ref.get("version") == "v2",
                   f"{crit_ref.get('id')}/{crit_ref.get('version')}"))
    pp = str(repo / "packages" / "packs" / "src")
    if pp not in sys.path:
        sys.path.insert(0, pp)
    try:
        import aqyra_packs as _packs
        man_v2 = _packs.load_pack(repo / "data" / "packs", "criterio", "AQ", "v2")
        got_hash = _packs.content_hash(man_v2)
        exp_hash = json.loads((repo / "data" / "packs" / "criterio" / "AQ" / "v2"
                               / "golden" / "expected.json").read_text(encoding="utf-8"))["content_sha256"]
        checks.append(("criterio/AQ/v2 anclado (content_sha256)", got_hash == exp_hash,
                       f"{got_hash[:12]}… (esperado {exp_hash[:12]}…)"))
    except Exception as e:  # noqa: BLE001
        checks.append(("criterio/AQ/v2 anclado (content_sha256)", False, f"{type(e).__name__}: {e}"))

    pe = str(DEFAULT_ENGINE_C5)
    if pe not in sys.path:
        sys.path.insert(0, pe)
    try:
        from aqyra_presupuesto import medir, presupuestar, proyectar, suma_proyeccion
    except Exception as e:  # noqa: BLE001
        checks.append(("import engine (medir/presupuestar/proyectar)", False, f"{type(e).__name__}: {e}"))
        print_checks(checks)
        return False

    # 3 · medir (con criterio para el fallback funcional) + presupuestar
    fuentes = [{"id": m.get("id"), "disciplina": m.get("disciplina"),
                "path": case_dir / m["fichero"], "fichero": m["fichero"]}
               for m in entrada.get("modelo", {}).get("fuente_maestro", {}).get("modelos", [])]
    criterio = _pack_contenido_c5(repo, "criterio", crit_ref)
    banco = _pack_contenido_c5(repo, "banco", entrada.get("banco_ref", {}))
    modelo = medir(fuentes, criterio)
    modelo["proyecto"] = entrada.get("proyecto")
    pres = presupuestar(modelo, criterio, banco, entrada.get("parametros", {}))
    ok, detail = validate_against_schema(pres, schemas["salida"])
    checks.append(("presupuesto (para la vista) conforma esquema", ok, detail))
    pem = float(pres.get("resumen", {}).get("PEM", 0))
    pec = float(pres.get("resumen", {}).get("PEC", 0))
    checks.append(("presupuesto reproduce PEM…PEC de referencia",
                   abs(pem - float(cost.get("PEM", 0))) <= ia and abs(pec - float(cost.get("PEC", 0))) <= ia,
                   f"PEM {pem} PEC {pec}"))

    # 4+5 · por cada vista: DETERMINISMO + INVARIANTE + SEMÁNTICA
    checks.append(("vistas declaradas", len(vistas) >= 3,
                   f"{len(vistas)} vista/s (≥3 proyecciones distintas: espacial + uniclass + funcional; "
                   f"funcional engloba IfcSystem + IfcZone 50/50 + fallback)"))
    for v in vistas:
        vid = v.get("id", f"{v.get('eje')}×{v.get('corte')}")
        eje, corte = v.get("eje", "coste"), v.get("corte")
        f1 = proyectar(pres, modelo, eje, corte)
        f2 = proyectar(pres, modelo, eje, corte)
        checks.append((f"[{vid}] determinista", f1 == f2, "proyectar 2× = misma salida"))
        suma_exp = float(v.get("suma", cost.get("PEM", 0)))
        checks.append((f"[{vid}] invariante Σ == total", abs(suma_proyeccion(f1) - suma_exp) <= ia,
                       f"Σ {suma_proyeccion(f1)} (esperado {suma_exp})"))
        got_g = {r["grupo"]: r for r in f1}
        exp_g = {g["grupo"]: g for g in v.get("grupos", [])}
        set_ok = set(got_g) == set(exp_g)
        checks.append((f"[{vid}] grupos == expected", set_ok,
                       f"{sorted(got_g)}" if set_ok
                       else f"faltan {sorted(set(exp_g) - set(got_g))} / sobran {sorted(set(got_g) - set(exp_g))}"))
        malos = []
        for grupo, ge in exp_g.items():
            gg = got_g.get(grupo)
            if gg is None:
                continue
            if abs(float(gg["valor_total"]) - float(ge["valor_total"])) > ia:
                malos.append(f"{grupo}:{gg['valor_total']}≠{ge['valor_total']}")
            elif "fuente" in ge and gg.get("fuente") != ge["fuente"]:
                malos.append(f"{grupo}:fuente {gg.get('fuente')}≠{ge['fuente']}")
        checks.append((f"[{vid}] valor_total + fuente por grupo", not malos,
                       "todos casan" if not malos else f"{len(malos)} — {malos[0]}"))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


def _run_c5_export(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                   repo: Path = DEFAULT_REPO) -> bool:
    """C5 · Export firmable (capa transversal documentos/export, E6.2): despacha por CONSUMIDOR. El
    PRIMARIO es CONTRACTUAL (presupuesto-obra); proyeccion-valor es export de gestion (secundario)."""
    consumidor = expected.get("consumidor") or ("presupuesto-obra" if expected.get("fuente_presupuesto")
                                                else "proyeccion-valor")
    if consumidor == "presupuesto-obra":
        return _run_export_presupuesto(case_dir, contracts_dir, expected, tol, repo)
    return _run_export_proyeccion(case_dir, contracts_dir, expected, tol, repo)


def _run_export_presupuesto(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                            repo: Path = DEFAULT_REPO) -> bool:
    """C5 · Export firmable del PRESUPUESTO CONTRACTUAL (consumidor primario, E6.2). El nucleo envuelve
    los compositores que ya existen: Word=componer_documento, BC3=emitir_bc3, y anade el PDF sellado +
    manifiesto. Artefacto autoritativo = el `salida-presupuesto` ANCLADO de GOL-PRE-01 (se LEE). Ancla
    por CONTENIDO (patron GOL-DOC-01/D3), no bytes ni pixeles. Corre SIN ifcopenshell. Justificacion de
    medicion v0 = cantidad + criterio + origen + GUIDs. Un fallo se investiga en el NUCLEO/compositor."""
    import tempfile

    checks: list[tuple[str, bool, str]] = []
    schemas = load_c5_schemas(contracts_dir)
    ia = float(tol.get("importe_abs", 0.01))
    exp_export = expected.get("export", {})

    # 1 · salida-presupuesto anclado (se LEE) + conforma esquema
    fuente = expected.get("fuente_presupuesto")
    fuente_path = case_dir.parent / str(fuente) / "expected.json"
    if not fuente_path.exists():
        print_checks([("fuente de presupuesto presente", False, f"{fuente} no encontrada")])
        return False
    presupuesto = json.loads(fuente_path.read_text(encoding="utf-8")).get("presupuesto", {})
    ok, detail = validate_against_schema(presupuesto, schemas["salida"])
    checks.append((f"presupuesto fuente ({fuente}) conforma salida-presupuesto", ok, detail))
    descriptor = dict(expected.get("descriptor") or {})

    # 2 · import del nucleo (comun + export + presupuesto + bc3)
    for pth in (DEFAULT_COMUN_DOC, DEFAULT_EXPORT_C5, DEFAULT_DOCUMENTO_C5,
                repo / "engines" / "bc3" / "src"):
        if str(pth) not in sys.path:
            sys.path.insert(0, str(pth))
    try:
        import aqyra_documento_export as adx
        from aqyra_documento_export import manifiesto as _man
        from aqyra_documento_export import firma as _firma
    except Exception as e:  # noqa: BLE001
        checks.append(("import del nucleo documentos/export", False, f"{type(e).__name__}: {e}"))
        print_checks(checks)
        return False

    # 3 · contract-first: descriptor conforma su esquema
    man_schema = None
    try:
        import jsonschema
        desc_schema = json.loads((DEFAULT_EXPORT_ESQUEMAS / "descriptor-export.schema.json").read_text(encoding="utf-8"))
        man_schema = json.loads((DEFAULT_EXPORT_ESQUEMAS / "manifiesto-export.schema.json").read_text(encoding="utf-8"))
        jsonschema.validate(descriptor, desc_schema)
        checks.append(("descriptor conforma descriptor-export.schema.json", True, "forward-open"))
    except Exception as e:  # noqa: BLE001
        checks.append(("descriptor conforma descriptor-export.schema.json", False, f"{type(e).__name__}: {e}"))

    with tempfile.TemporaryDirectory() as td:
        tdp = Path(td)
        bundle_a = adx.componer_export(presupuesto, descriptor, {"salida": tdp / "a"})
        man = json.loads((bundle_a / adx.NOMBRE_MANIFIESTO).read_text(encoding="utf-8"))
        if man_schema is not None:
            try:
                import jsonschema
                jsonschema.validate(man, man_schema)
                checks.append(("manifiesto conforma manifiesto-export.schema.json", True, "forward-open"))
            except Exception as e:  # noqa: BLE001
                checks.append(("manifiesto conforma manifiesto-export.schema.json", False, f"{type(e).__name__}: {e}"))

        fmap = adx.consumidor_de("presupuesto-obra") or {}
        formatos = list(descriptor.get("formatos") or list(fmap))
        # 4 · bundle: formatos declarados + manifiesto
        faltan = [f for f in formatos if f in fmap and not (bundle_a / fmap[f][0]).exists()]
        faltan_man = not (bundle_a / adx.NOMBRE_MANIFIESTO).exists()
        checks.append(("bundle con los formatos declarados + manifiesto", not faltan and not faltan_man,
                       f"{sorted(q.name for q in bundle_a.iterdir())}" if not faltan and not faltan_man
                       else f"faltan: {faltan}{' + manifiesto' if faltan_man else ''}"))

        em = {p["codigo"]: p for p in presupuesto.get("estado_mediciones", [])}
        resumen = presupuesto.get("resumen", {})
        pec_txt = _fmt_es(resumen.get("PEC", 0))

        # 5 · Word CONTRACTUAL (envuelve componer_documento): partidas + PEC
        if "word" in fmap:
            paras, tablas = _extraer_docx(bundle_a / fmap["word"][0])
            blob_w = "\n".join(paras) + "\n" + "\n".join(" ".join(c for f in t for c in f) for t in tablas)
            falt_pw = [c for c in em if c not in blob_w]
            checks.append(("Word contractual: partidas + PEC", not falt_pw and pec_txt in blob_w,
                           f"{len(em)} partidas, PEC {pec_txt}" if not falt_pw and pec_txt in blob_w
                           else f"faltan partidas {falt_pw[:3]} / PEC {'ok' if pec_txt in blob_w else 'no'}"))
            if "n_partidas" in exp_export:
                checks.append(("n de partidas == esperado", len(em) == exp_export["n_partidas"],
                               f"{len(em)} (esperado {exp_export['n_partidas']})"))

        # 6 · PDF firmable: partidas + PEC + JUSTIFICACION (criterio + GUID) + content_sha256
        if "pdf" in fmap:
            pdf_txt = _pdf_texto(bundle_a / fmap["pdf"][0])
            sha = (man.get("artefacto") or {}).get("content_sha256", "")
            falt_pp = [c for c in em if c not in pdf_txt]
            criterios = {p.get("criterio_aplicado", "") for p in em.values() if p.get("criterio_aplicado")}
            crit_ok = all(cr.split()[0] in pdf_txt for cr in criterios) if criterios else True
            guids = [g for p in em.values() for g in p.get("trazabilidad", [])]
            guid_ok = all(g in pdf_txt for g in guids) if guids else True
            ok_pdf = not falt_pp and pec_txt in pdf_txt and crit_ok and guid_ok and bool(sha) and sha in pdf_txt
            checks.append(("PDF firmable: partidas + PEC + justificacion (criterio+GUIDs) + sha256", ok_pdf,
                           "presentes" if ok_pdf else
                           f"partidas {'ok' if not falt_pp else falt_pp[:2]} PEC {pec_txt in pdf_txt} crit {crit_ok} guid {guid_ok} sha {sha in pdf_txt}"))

        # 7 · BC3 (licitacion): codigos de partida presentes
        if "bc3" in fmap:
            bc3_txt = (bundle_a / fmap["bc3"][0]).read_text(encoding="utf-8")
            falt_bc3 = [c for c in em if c not in bc3_txt]
            checks.append(("BC3 (FIEBDC-3) con los codigos de partida", not falt_bc3,
                           f"{len(em)} codigos + ~V/~C/~M" if not falt_bc3 else f"faltan: {falt_bc3[:3]}"))

        # 8 · XLSX mediciones: cantidades presentes
        if "xlsx" in fmap:
            nums = _xlsx_numeros(bundle_a / fmap["xlsx"][0])
            cants = [round(float(p.get("cantidad", 0)), 3) for p in em.values()]
            falt_x = [c for c in cants if not any(abs(c - n) <= 0.001 for n in nums)]
            checks.append(("XLSX mediciones: cantidades por partida", not falt_x,
                           f"{len(nums)} celdas" if not falt_x else f"faltan {len(falt_x)}"))

        # 9 · manifiesto INTEGRO (Llave 1)
        ok_int, det_int = _man.integridad(man, presupuesto)
        checks.append(("manifiesto integro (content_sha256 + versiones ancladas)", ok_int, det_int))

        # 10 · isCertified: bundle SIN firmar NO es verified-signed; hook Llave 2 presente
        cert = adx.es_certificado(bundle_a)
        estado = _firma.estado_firmable(bundle_a)
        checks.append(("isCertified: bundle sin firmar NO es verified-signed",
                       (not cert) and estado == "computed", f"estado={estado}, certificado={cert}"))
        checks.append(("hook de firma Llave 2 presente (release-time, NO ejecutado en el gate)",
                       callable(getattr(_firma, "firmar_detached", None)), "firmar_detached disponible"))

        # 11 · DETERMINISMO por contenido (componer 2x): word/pdf/bc3/xlsx/manifiesto
        bundle_b = adx.componer_export(presupuesto, descriptor, {"salida": tdp / "b"})
        det = True
        for f in formatos:
            if f not in fmap:
                continue
            fa, fb = bundle_a / fmap[f][0], bundle_b / fmap[f][0]
            if f == "word":
                det = det and _extraer_docx(fa) == _extraer_docx(fb)
            elif f == "xlsx":
                det = det and _xlsx_celdas(fa) == _xlsx_celdas(fb)
            elif f == "pdf":
                det = det and _pdf_texto(fa) == _pdf_texto(fb)
            else:
                det = det and fa.read_bytes() == fb.read_bytes()
        det = det and (bundle_a / adx.NOMBRE_MANIFIESTO).read_bytes() == (bundle_b / adx.NOMBRE_MANIFIESTO).read_bytes()
        checks.append(("export determinista (2x = Word/PDF/BC3/XLSX/manifiesto identicos)", det,
                       "contenido identico" if det else "difiere entre dos composiciones"))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


def _run_export_proyeccion(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                           repo: Path = DEFAULT_REPO) -> bool:
    """C5 · Export firmable de la proyeccion (SECUNDARIO/gestion, documentos/export). El nucleo
    DETERMINISTA componer_export toma el artefacto de proyeccion ANCLADO (GOL-PRE-03 — se LEE, no se
    re-ancla) + un descriptor y produce el bundle firmable (Word + XLSX + PDF sellado + manifiesto).
    Ancla ESTRUCTURA + CONTENIDO (patron GOL-DOC-01/GOL-PLI-01/D3), no bytes ni pixeles. Dos llaves:
    INTEGRIDAD en el gate (Llave 1); firma GPG de JM en el release (Llave 2, hook NO ejecutado aqui).
    Corre SIN ifcopenshell (lee la proyeccion anclada, no re-mide). Un fallo se investiga en el NUCLEO."""
    import tempfile

    checks: list[tuple[str, bool, str]] = []
    ia = float(tol.get("importe_abs", 0.01))
    exp_export = expected.get("export", {})

    # 1 · artefacto de proyeccion anclado (se LEE)
    fuente = expected.get("fuente_proyeccion")
    fuente_path = case_dir.parent / str(fuente) / "expected.json"
    if not fuente_path.exists():
        print_checks([("fuente de proyeccion presente", False, f"{fuente} no encontrada")])
        return False
    artefacto = json.loads(fuente_path.read_text(encoding="utf-8"))
    vistas = artefacto.get("vistas") or []
    checks.append((f"artefacto de proyeccion ({fuente}) con vistas", len(vistas) >= 1,
                   f"{len(vistas)} vista/s"))
    descriptor = dict(expected.get("descriptor") or {})

    # 2 · import del nucleo (comun + export), patron engines/*
    for pth in (DEFAULT_COMUN_DOC, DEFAULT_EXPORT_C5):
        if str(pth) not in sys.path:
            sys.path.insert(0, str(pth))
    try:
        import aqyra_documento_export as adx
        from aqyra_documento_export import manifiesto as _man
        from aqyra_documento_export import firma as _firma
    except Exception as e:  # noqa: BLE001
        checks.append(("import del nucleo documentos/export", False, f"{type(e).__name__}: {e}"))
        print_checks(checks)
        return False

    # 3 · contract-first: descriptor conforma su esquema (el manifiesto, tras componer)
    man_schema = None
    try:
        import jsonschema
        desc_schema = json.loads((DEFAULT_EXPORT_ESQUEMAS / "descriptor-export.schema.json").read_text(encoding="utf-8"))
        man_schema = json.loads((DEFAULT_EXPORT_ESQUEMAS / "manifiesto-export.schema.json").read_text(encoding="utf-8"))
        jsonschema.validate(descriptor, desc_schema)
        checks.append(("descriptor conforma descriptor-export.schema.json", True, "forward-open"))
    except Exception as e:  # noqa: BLE001
        checks.append(("descriptor conforma descriptor-export.schema.json", False, f"{type(e).__name__}: {e}"))

    with tempfile.TemporaryDirectory() as td:
        tdp = Path(td)
        bundle_a = adx.componer_export(artefacto, descriptor, {"salida": tdp / "a"})
        man = json.loads((bundle_a / adx.NOMBRE_MANIFIESTO).read_text(encoding="utf-8"))
        if man_schema is not None:
            try:
                import jsonschema
                jsonschema.validate(man, man_schema)
                checks.append(("manifiesto conforma manifiesto-export.schema.json", True, "forward-open"))
            except Exception as e:  # noqa: BLE001
                checks.append(("manifiesto conforma manifiesto-export.schema.json", False, f"{type(e).__name__}: {e}"))

        # 4 · bundle: formatos declarados + manifiesto.json
        formatos = list(descriptor.get("formatos") or [])
        nombre = {f: v[0] for f, v in (adx.consumidor_de("proyeccion-valor") or {}).items()}
        faltan = [f for f in formatos if f in nombre and not (bundle_a / nombre[f]).exists()]
        faltan_man = not (bundle_a / adx.NOMBRE_MANIFIESTO).exists()
        checks.append(("bundle con los formatos declarados + manifiesto", not faltan and not faltan_man,
                       f"{sorted(q.name for q in bundle_a.iterdir())}" if not faltan and not faltan_man
                       else f"faltan: {faltan}{' + manifiesto' if faltan_man else ''}"))

        # cifras esperadas: grupos + Sigma de todas las vistas
        cifras = []
        for v in vistas:
            cifras.append(round(float(v.get("suma", 0)), 2))
            for g in v.get("grupos", []):
                cifras.append(round(float(g.get("valor_total", 0)), 2))

        # 5 · Word: cifras presentes (formato ES)
        paras, tablas = _extraer_docx(bundle_a / nombre["word"])
        blob_w = "\n".join(paras) + "\n" + "\n".join(" ".join(c for f in t for c in f) for t in tablas)
        faltan_w = [c for c in cifras if _fmt_es(c) not in blob_w]
        checks.append(("cifras de la proyeccion en el Word (grupos + Sigma)", not faltan_w,
                       f"{len(cifras)} cifras" if not faltan_w else f"faltan {len(faltan_w)}: {faltan_w[:3]}"))
        if "n_vistas" in exp_export:
            checks.append(("n de vistas == esperado", len(vistas) == exp_export["n_vistas"],
                           f"{len(vistas)} (esperado {exp_export['n_vistas']})"))

        # 6 · XLSX: cifras numericas presentes
        nums = _xlsx_numeros(bundle_a / nombre["xlsx"])
        faltan_x = [c for c in cifras if not any(abs(c - n) <= ia for n in nums)]
        checks.append(("cifras de la proyeccion en el XLSX", not faltan_x,
                       f"{len(nums)} celdas numericas" if not faltan_x else f"faltan {len(faltan_x)}: {faltan_x[:3]}"))

        # 7 · PDF: Sigma de cada vista + content_sha256 en el texto extraido
        pdf_txt = _pdf_texto(bundle_a / nombre["pdf"])
        sha = (man.get("artefacto") or {}).get("content_sha256", "")
        sumas = [round(float(v.get("suma", 0)), 2) for v in vistas]
        faltan_p = [x for x in sumas if _fmt_es(x) not in pdf_txt]
        checks.append(("PDF sellado: Sigma de cada vista + content_sha256 en el texto",
                       not faltan_p and bool(sha) and sha in pdf_txt,
                       "presentes" if not faltan_p and sha in pdf_txt
                       else f"faltan sumas {faltan_p} / sha {'ok' if sha in pdf_txt else 'no'}"))

        # 8 · manifiesto INTEGRO (Llave 1)
        ok_int, det_int = _man.integridad(man, artefacto)
        checks.append(("manifiesto integro (content_sha256 + modelo_md5 + versiones ancladas)",
                       ok_int, det_int))
        if "suma" in exp_export:
            inv = (man.get("invariante") or {}).get("suma_declarada")
            checks.append(("invariante suma declarada == esperado",
                           inv is not None and abs(float(inv) - float(exp_export["suma"])) <= ia,
                           f"{inv} (esperado {exp_export['suma']})"))

        # 9 · isCertified (D-EX-5): bundle SIN firmar NO es verified-signed; hook Llave 2 presente
        cert = adx.es_certificado(bundle_a)
        estado = _firma.estado_firmable(bundle_a)
        checks.append(("isCertified: bundle sin firmar NO es verified-signed",
                       (not cert) and estado == "computed", f"estado={estado}, certificado={cert}"))
        checks.append(("hook de firma Llave 2 presente (release-time, NO ejecutado en el gate)",
                       callable(getattr(_firma, "firmar_detached", None)), "firmar_detached disponible"))

        # 10 · DETERMINISMO por contenido (componer 2x)
        bundle_b = adx.componer_export(artefacto, descriptor, {"salida": tdp / "b"})
        det = ((paras, tablas) == _extraer_docx(bundle_b / nombre["word"])
               and _xlsx_celdas(bundle_a / nombre["xlsx"]) == _xlsx_celdas(bundle_b / nombre["xlsx"])
               and pdf_txt == _pdf_texto(bundle_b / nombre["pdf"])
               and (bundle_a / adx.NOMBRE_MANIFIESTO).read_bytes() == (bundle_b / adx.NOMBRE_MANIFIESTO).read_bytes())
        checks.append(("export determinista (2x = Word/XLSX/PDF/manifiesto identicos)", det,
                       "contenido identico" if det else "difiere entre dos composiciones"))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


def run_case_c5(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                repo: Path = DEFAULT_REPO) -> bool:
    """C5: presupuesto trazable — modo ANCLADO (contract-first, Fase IV·h1 — ver ficha/DECISIONES).

    El presupuesto esperado (`expected["presupuesto"]`) es el ORÁCULO (medición manual congelada,
    verificada x2). El engine `engines/presupuesto` (hilo posterior) ANTEPONDRÁ el recompute (parser
    de Qto + motor) contra este MISMO expected (costura C1/C3/C4). Modo anclado — más checks que hoy:

      1. Conformidad de los 2 esquemas (entrada + salida).
      2. Identidad por hash de las fixtures con Qto (entradas_md5 == ficheros == modelo neutro).
      3. Packs criterio + banco anclados en versions.lock; refs de la entrada == lock.
      4. Partidas origen=modelo ⊆ banco (precio == banco) y ⊆ criterio; origen=regla no exige banco.
      5. importe == cantidad × precio_unitario (±importe_abs).
      6. precio_unitario == cuadro nº1; nº1 == nº2 (Σ componentes + indirectos).
      7. PEM == Σ importes == Σ capítulos; GG/BI/base/IVA/PEC según parametros.
      8. trazabilidad ⊆ GUIDs del modelo (origen=modelo) / vacía (origen=regla); taxonomía cerrada.
      9. S&S (origen=regla) == ratio del criterio × PEM medible.
    """
    # Dispatch por modo (Fase IV·h3): GOL-PRE-02 (write-back 5D) tiene su propia rama.
    if expected.get("modo") == "5d" or "cost_5d" in expected:
        return _run_c5_5d(case_dir, contracts_dir, expected, tol, repo)
    # Capa de documentos: GOL-DOC-01 (Documento de Presupuesto) — compositor determinista (D1–D5).
    if expected.get("modo") == "documento":
        return _run_c5_documento(case_dir, contracts_dir, expected, tol, repo)
    # Proyección (E2.2, D24–D29): GOL-PRE-03 (vista por eje/corte) — group-by determinista.
    if expected.get("modo") == "proyeccion":
        return _run_c5_proyeccion(case_dir, contracts_dir, expected, tol, repo)
    # Eje CARBONO (E3.3, D39–D44): GOL-CAR-01 (valoración carbono + etapas + proyección de carbono).
    if expected.get("modo") == "carbono":
        return _run_c5_carbono(case_dir, contracts_dir, expected, tol, repo)
    # Capa de documentos: GOL-PLI-01 (Pliego de Condiciones Tecnicas) — compositor determinista (E4.1).
    if expected.get("modo") == "pliego":
        return _run_c5_pliego(case_dir, contracts_dir, expected, tol, repo)
    # Capa transversal de export firmable: GOL-EXP-01 (bundle firmable de la proyeccion) — E6.2.
    if expected.get("modo") == "export":
        return _run_c5_export(case_dir, contracts_dir, expected, tol, repo)

    checks: list[tuple[str, bool, str]] = []
    schemas = load_c5_schemas(contracts_dir)

    entrada_path = case_dir / "entrada.json"
    if not entrada_path.exists():
        print_checks([("entrada.json presente", False, "no encontrado")])
        return False
    entrada = json.loads(entrada_path.read_text(encoding="utf-8"))
    pres = expected.get("presupuesto", {})
    ia = float(tol.get("importe_abs", 0.01))

    def approx(a, b) -> bool:
        try:
            return abs(float(a) - float(b)) <= ia
        except (TypeError, ValueError):
            return False

    # 0 · RECOMPUTE con el engine real (Fase IV·h2, D9) — se ANTEPONE al anclado.
    # engines/presupuesto MIDE desde las fixtures con Qto (la medición nace del modelo, D7) y
    # PRESUPUESTA; se compara contra el MISMO expected (D10 del C4: más checks, nunca menos).
    # Un fallo se investiga en el ENGINE (parser/motor), jamás en el expected.
    try:
        modelo_medido, got_pres = _recompute_c5(case_dir, entrada, repo)
        checks.append(("engine presupuesto medir()+presupuestar() ejecuta", True,
                       "engines/presupuesto sobre las fixtures con Qto"))
        ok, detail = validate_against_schema(got_pres, schemas["salida"])
        checks.append(("presupuesto recomputado conforma", ok, detail))
        d_med = _diff_medicion_c5(modelo_medido, entrada.get("modelo", {}), tol)
        checks.append(("medición nace del modelo (parser reproduce cantidades)", not d_med,
                       f"{len(d_med)} diff/s — {d_med[0]}" if d_med
                       else "cantidades leídas de los Qto del IFC"))
        d_pre = _diff_presupuesto_c5(got_pres, pres, tol)
        checks.append(("recompute presupuesto == expected", not d_pre,
                       f"{len(d_pre)} diff/s — {d_pre[0]}" if d_pre
                       else "reproducido (cantidad/precio/importe/capítulos/PEM…PEC/trazabilidad)"))
    except Exception as e:  # noqa: BLE001
        checks.append(("recompute con el engine", False, f"{type(e).__name__}: {e}"))

    # 1 · conformidad de esquemas
    for name, inst, key in (("entrada conforma esquema", entrada, "entrada"),
                            ("presupuesto conforma esquema", pres, "salida")):
        ok, detail = validate_against_schema(inst, schemas[key])
        checks.append((name, ok, detail))

    # 2 · identidad por hash de las fixtures con Qto
    entradas = expected.get("entradas_md5", {})
    checks.append(("entradas declaradas", bool(entradas), f"{len(entradas)} fichero/s"))
    for rel, exp_md5 in sorted(entradas.items()):
        p = case_dir / rel
        if not p.exists():
            checks.append((f"identidad {rel}", False, "fichero no encontrado"))
            continue
        got = _md5(p)
        checks.append((f"identidad {rel}", got == exp_md5,
                       f"md5 {got[:12]}… (esperado {exp_md5[:12]}…)"))
    md5_modelo = {m["fichero"]: m.get("md5")
                  for m in entrada.get("modelo", {}).get("fuente_maestro", {}).get("modelos", [])}
    checks.append(("md5 modelo neutro == entradas", md5_modelo == entradas,
                   "las fixtures del modelo neutro anclan los mismos IFC con Qto"))

    # 3 · packs criterio + banco anclados en versions.lock
    crit_ref = entrada.get("criterio_ref", {})
    banco_ref = entrada.get("banco_ref", {})
    lock_crit = _lock_packs(repo, "criterio")
    if (lock_crit.get("id") == crit_ref.get("id")
            and lock_crit.get("version") == crit_ref.get("version")):
        # criterio anclado por el POINTER [packs.criterio] (AQ/v1 en GOL-PRE-01..04)
        checks.append(("criterio anclado en versions.lock", True,
                       f"lock={lock_crit.get('id')}/{lock_crit.get('version')}"))
    else:
        # criterio que NO es el pointer (AQ/v3 en GOL-PRE-05): anclado por su content_sha256, como
        # AQ/v2 en la ruta de carbono. [packs.criterio] NO se mueve (GOL-PRE-01..04 intactas, D54).
        crit_ok, crit_detalle = _criterio_anclado_por_sha(repo, crit_ref)
        checks.append(("criterio anclado (content_sha256)", crit_ok, crit_detalle))
    banco_ok, banco_detalle = _banco_anclado_en_lock(repo, banco_ref)
    checks.append(("banco anclado en versions.lock", banco_ok, banco_detalle))

    # cargar contenido de los packs (banco.json + criterio.json)
    banco = criterio = None
    try:
        bman, bdir = _pack_c5(repo, "banco", banco_ref)
        banco = json.loads((bdir / bman["contenido"]["fichero"]).read_text(encoding="utf-8"))
        cman, cdir = _pack_c5(repo, "criterio", crit_ref)
        criterio = json.loads((cdir / cman["contenido"]["fichero"]).read_text(encoding="utf-8"))
        checks.append(("packs criterio + banco cargados", True,
                       f"{len(banco.get('partidas', []))} partidas de banco"))
    except Exception as e:  # noqa: BLE001
        checks.append(("packs criterio + banco cargados", False, f"{type(e).__name__}: {e}"))

    precio_banco = {p["codigo"]: p["precio"] for p in (banco or {}).get("partidas", [])}
    codigos_criterio = set()
    for regla in (criterio or {}).get("reglas_por_clase", []):
        for p in regla.get("partidas", []):
            codigos_criterio.add(p["codigo"])
    codigos_sin_geo = {p["codigo"] for p in (criterio or {}).get("reglas_sin_geometria", [])}

    em = pres.get("estado_mediciones", [])
    cp1 = {c["codigo"]: c for c in pres.get("cuadro_precios_1", [])}
    cp2 = {c["codigo"]: c for c in pres.get("cuadro_precios_2", [])}

    # 4 · partidas origen=modelo ⊆ banco (precio) y ⊆ criterio; origen=regla no exige banco
    faltan_banco = [l["codigo"] for l in em
                    if l.get("origen") == "modelo" and l["codigo"] not in precio_banco]
    checks.append(("partidas origen=modelo ⊆ banco", not faltan_banco,
                   "todas con precio en el banco" if not faltan_banco
                   else f"sin respaldo en banco: {faltan_banco}"))
    faltan_crit = [l["codigo"] for l in em
                   if l.get("origen") == "modelo" and l["codigo"] not in codigos_criterio]
    checks.append(("partidas origen=modelo ⊆ criterio", not faltan_crit,
                   "todas mapeadas por el criterio" if not faltan_crit
                   else f"sin regla de medición: {faltan_crit}"))

    # 5 · importe == cantidad × precio_unitario ; 6 · precio == cuadro nº1 == cuadro nº2
    malos_importe, malos_p1, malos_p2 = [], [], []
    for l in em:
        cod = l["codigo"]
        if not approx(l.get("importe"), float(l.get("cantidad", 0)) * float(l.get("precio_unitario", 0))):
            malos_importe.append(cod)
        if cod in cp1 and not approx(l.get("precio_unitario"), cp1[cod].get("precio")):
            malos_p1.append(cod)
        # precio del banco (solo origen=modelo)
        if l.get("origen") == "modelo" and cod in precio_banco and not approx(l.get("precio_unitario"), precio_banco[cod]):
            malos_p1.append(cod + "(banco)")
    checks.append(("importe == cantidad × precio", not malos_importe,
                   "todos" if not malos_importe else f"descuadran: {malos_importe}"))
    checks.append(("precio_unitario == cuadro nº1 (y banco)", not malos_p1,
                   "coherentes" if not malos_p1 else f"descuadran: {malos_p1}"))
    for cod, c2 in cp2.items():
        suma = sum(float(x.get("subtotal", 0)) for x in c2.get("componentes", []))
        suma += float(c2.get("costes_indirectos", 0))
        if not approx(c2.get("precio_total"), suma):
            malos_p2.append(cod)
        if cod in cp1 and not approx(c2.get("precio_total"), cp1[cod].get("precio")):
            malos_p2.append(cod + "(!=nº1)")
    checks.append(("cuadro nº2 == Σ componentes + indirectos == nº1", not malos_p2,
                   "descompuesto cuadra" if not malos_p2 else f"descuadran: {malos_p2}"))
    # todas las partidas medidas tienen cuadro nº1; las origen=modelo, cuadro nº2
    sin_p1 = [l["codigo"] for l in em if l["codigo"] not in cp1]
    sin_p2 = [l["codigo"] for l in em if l.get("origen") == "modelo" and l["codigo"] not in cp2]
    checks.append(("cuadros cubren las partidas", not sin_p1 and not sin_p2,
                   "nº1 todas, nº2 las de modelo" if not sin_p1 and not sin_p2
                   else f"sin nº1: {sin_p1} · sin nº2: {sin_p2}"))

    # 7 · PEM == Σ importes == Σ capítulos ; GG/BI/base/IVA/PEC
    res = pres.get("resumen", {})
    par = entrada.get("parametros", {})
    suma_imp = sum(float(l.get("importe", 0)) for l in em)
    PEM = float(res.get("PEM", 0))
    checks.append(("PEM == Σ importes", approx(PEM, suma_imp),
                   f"PEM {PEM} vs Σ {round(suma_imp, 2)}"))
    caps = res.get("capitulos", [])
    imp_por_cod = {l["codigo"]: float(l.get("importe", 0)) for l in em}
    cap_ok = True
    for c in caps:
        suma_c = sum(imp_por_cod.get(cod, 0) for cod in c.get("partidas", []))
        if not approx(c.get("importe"), suma_c):
            cap_ok = False
    checks.append(("capítulos == Σ partidas y Σ capítulos == PEM",
                   cap_ok and approx(sum(float(c.get("importe", 0)) for c in caps), PEM),
                   f"{len(caps)} capítulos"))
    # Espeja el redondeo POR PASOS del engine (mul2/r2, HALF_UP 2 dec): el presupuesto es
    # traduccion determinista, no aritmetica en crudo. Un recompute en crudo (sin redondeo
    # intermedio) puede divergir del engine >0,01 por acumulacion (GOL-PRE-05: PEC 14389.50 del
    # engine vs 14389.51 en crudo) -> falso negativo. El engine es la fuente de verdad; este check
    # valida SU aritmetica (D9: se investiga el emisor, no se afloja el oraculo).
    from decimal import Decimal as _Dec, ROUND_HALF_UP as _HU
    def _r2c(x): return float(_Dec(str(x)).quantize(_Dec("0.01"), rounding=_HU))
    def _m2c(a, b): return float((_Dec(str(a)) * _Dec(str(b))).quantize(_Dec("0.01"), rounding=_HU))
    gg = _m2c(par.get("gg_pct", 0), PEM)
    bi = _m2c(par.get("bi_pct", 0), PEM)
    base = _r2c(_Dec(str(PEM)) + _Dec(str(gg)) + _Dec(str(bi)))
    iva = _m2c(par.get("iva_pct", 0), base)
    pec = _r2c(_Dec(str(base)) + _Dec(str(iva)))
    econ_ok = (approx(res.get("gg"), gg) and approx(res.get("bi"), bi)
               and approx(res.get("base"), base) and approx(res.get("iva"), iva)
               and approx(res.get("PEC"), pec))
    checks.append(("GG/BI/base/IVA/PEC según parametros", econ_ok,
                   f"PEC {res.get('PEC')} (esperado {round(pec, 2)})"))

    # 8 · trazabilidad ⊆ GUIDs del modelo (origen=modelo) / vacía (origen=regla); taxonomía cerrada
    guids = {o["guid"] for o in entrada.get("modelo", {}).get("objetos", [])}
    traza_ok = True
    origen_ok = True
    for l in em:
        org = l.get("origen")
        if org not in ORIGEN_C5:
            origen_ok = False
        tz = set(l.get("trazabilidad", []))
        if org == "modelo" and not (tz and tz <= guids):
            traza_ok = False
        if org == "regla" and tz:
            traza_ok = False
    checks.append(("origen taxonomía cerrada (modelo/regla/manual)", origen_ok,
                   f"{sorted({l.get('origen') for l in em})}"))
    checks.append(("trazabilidad ⊆ GUIDs del modelo", traza_ok,
                   f"{len(guids)} objetos en el modelo neutro"))

    # 9 · S&S (origen=regla) == ratio del criterio × PEM medible
    pem_medible = sum(float(l.get("importe", 0)) for l in em if l.get("origen") == "modelo")
    reglas_sg = {p["codigo"]: p for p in (criterio or {}).get("reglas_sin_geometria", [])}
    sys_ok, sys_detail = True, "sin partidas de regla"
    for l in em:
        if l.get("origen") == "regla" and l["codigo"] in reglas_sg:
            ratio = float(reglas_sg[l["codigo"]].get("ratio", 0))
            esperado = ratio * pem_medible
            sys_ok = approx(l.get("importe"), esperado)
            sys_detail = f"{l['codigo']} {l.get('importe')} (esperado {round(esperado, 2)} = {ratio}×{round(pem_medible, 2)})"
    checks.append(("partida por regla == ratio × PEM medible", sys_ok, sys_detail))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


def _run_c5_carbono(case_dir: Path, contracts_dir: Path, expected: dict, tol: dict,
                    repo: Path = DEFAULT_REPO) -> bool:
    """C5 · eje CARBONO (E3.3, D39-D44): valora la medicion de GOL-PRE-01 en carbono — GOL-CAR-01.

    Reusa las fixtures AUMENTADAS de GOL-PRE-03 (mismos md5) + criterio/AQ/v2 (D44). El carbono es
    traduccion determinista: el kgCO2e sale del factor del banco-carbono anclado (convencion, no verdad
    fisica). ANCLA (patron D14/D28, sin md5 de salida):
      1. Identidad por md5 de las fixtures aumentadas (entradas_md5).
      2. criterio/AQ/v2 + banco-carbono/generico/<version del caso> anclados por su content_sha256.
      3. Conformidad del esquema de salida (el carbono conforma SIN tocar el esquema).
      4. RECOMPUTE: medir(fixtures, v2) + presupuestar(..., banco-carbono, eje="carbono").
      5. SEMANTICA carbono: por partida valores.carbono {unitario x cantidad = total, unidad kgCO2e,
         banco, etapas A1A3/A4A5, Sigma etapas = total}; resumen del eje == Sigma totales.
      6. DETERMINISMO + INVARIANTE de proyeccion: proyectar 2x igual; Sigma == PEM eje; grupos/valores.
    Un fallo se investiga en el ENGINE (presupuesto.py), jamas aflojando el check."""
    checks: list[tuple[str, bool, str]] = []
    schemas = load_c5_schemas(contracts_dir)
    entrada = json.loads((case_dir / "entrada.json").read_text(encoding="utf-8"))
    ia = float(tol.get("importe_abs", 0.01))
    eje_exp = expected.get("eje", {})
    estado_exp = {e["codigo"]: e for e in expected.get("estado_carbono", [])}
    vistas = expected.get("vistas", [])

    # 1 · identidad de las fixtures aumentadas (las de GOL-PRE-03; GOL-PRE-01 intacta)
    for rel, exp_md5 in sorted(expected.get("entradas_md5", {}).items()):
        p = case_dir / rel
        got = _md5(p) if p.exists() else None
        checks.append((f"identidad {rel}", got == exp_md5,
                       f"md5 {str(got)[:12]}… (esperado {str(exp_md5)[:12]}…)"))

    # 2 · packs anclados (criterio/AQ/v2 + banco-carbono/generico/v1) por content_sha256
    crit_ref = entrada.get("criterio_ref", {})
    banco_ref = entrada.get("banco_ref", {})
    checks.append(("criterio_ref == AQ/v2", crit_ref.get("id") == "AQ" and crit_ref.get("version") == "v2",
                   f"{crit_ref.get('id')}/{crit_ref.get('version')}"))
    banco_ref_str = f"{banco_ref.get('familia')}/{banco_ref.get('id')}/{banco_ref.get('version')}"
    checks.append((f"banco_ref == banco-carbono/generico/{banco_ref.get('version')}",
                   banco_ref.get("familia") == "banco-carbono" and banco_ref.get("id") == "generico"
                   and banco_ref.get("version") in ("v1", "v2"),
                   banco_ref_str))
    pp = str(repo / "packages" / "packs" / "src")
    if pp not in sys.path:
        sys.path.insert(0, pp)
    try:
        import aqyra_packs as _packs
        for fam, rid, rver, clave in (("criterio", "AQ", "v2", "criterio/AQ/v2"),
                                      (banco_ref.get("familia"), banco_ref.get("id"),
                                       banco_ref.get("version"), banco_ref_str)):
            man = _packs.load_pack(repo / "data" / "packs", fam, rid, rver)
            got_h = _packs.content_hash(man)
            exp_h = json.loads((repo / "data" / "packs" / fam / rid / rver / "golden"
                                / "expected.json").read_text(encoding="utf-8"))["content_sha256"]
            checks.append((f"{clave} anclado (content_sha256)", got_h == exp_h,
                           f"{got_h[:12]}… (esperado {exp_h[:12]}…)"))
    except Exception as e:  # noqa: BLE001
        checks.append(("packs anclados (content_sha256)", False, f"{type(e).__name__}: {e}"))

    pe = str(DEFAULT_ENGINE_C5)
    if pe not in sys.path:
        sys.path.insert(0, pe)
    try:
        from aqyra_presupuesto import medir, presupuestar, proyectar, suma_proyeccion
    except Exception as e:  # noqa: BLE001
        checks.append(("import engine (medir/presupuestar/proyectar)", False, f"{type(e).__name__}: {e}"))
        print_checks(checks)
        return False

    # 3 · recompute del eje carbono
    fuentes = [{"id": m.get("id"), "disciplina": m.get("disciplina"),
                "path": case_dir / m["fichero"], "fichero": m["fichero"]}
               for m in entrada.get("modelo", {}).get("fuente_maestro", {}).get("modelos", [])]
    criterio = _pack_contenido_c5(repo, crit_ref.get("familia", "criterio"), crit_ref)
    banco = _pack_contenido_c5(repo, banco_ref.get("familia", "banco-carbono"), banco_ref)
    modelo = medir(fuentes, criterio)
    modelo["proyecto"] = entrada.get("proyecto")
    pres = presupuestar(modelo, criterio, banco, entrada.get("parametros", {}))
    ok, detail = validate_against_schema(pres, schemas["salida"])
    checks.append(("presupuesto de carbono conforma esquema (sin tocarlo)", ok, detail))

    # 4 · semantica del eje por partida (valores.carbono + etapas + Sigma etapas = total)
    got = {p["codigo"]: p for p in pres.get("estado_mediciones", [])}
    malos: list[str] = []
    for cod, e in estado_exp.items():
        p = got.get(cod)
        if p is None:
            malos.append(f"{cod}: no recomputada")
            continue
        v = (p.get("valores") or {}).get("carbono")
        if not v:
            malos.append(f"{cod}: sin valores.carbono")
            continue
        if v.get("unidad") != "kgCO2e":
            malos.append(f"{cod}: unidad {v.get('unidad')}")
        if abs(float(v.get("total", 0)) - float(e["total"])) > ia:
            malos.append(f"{cod}: total {v.get('total')} != {e['total']}")
        if abs(float(v.get("unitario", 0)) - float(e["unitario"])) > ia:
            malos.append(f"{cod}: unitario {v.get('unitario')} != {e['unitario']}")
        if abs(float(v.get("total", 0)) - float(p.get("importe", 0))) > ia:
            malos.append(f"{cod}: espejo total != importe")
        et_e = e.get("etapas")
        if et_e:
            et = v.get("etapas") or {}
            if set(et) != set(et_e):
                malos.append(f"{cod}: etapas {sorted(et)} != {sorted(et_e)}")
            else:
                if abs(sum(float(et[k]) for k in et) - float(v.get("total", 0))) > ia:
                    malos.append(f"{cod}: Sigma etapas != total")
                for k in et_e:
                    if abs(float(et[k]) - float(et_e[k])) > ia:
                        malos.append(f"{cod}.etapas.{k}")
            if v.get("banco") != banco_ref_str:
                malos.append(f"{cod}: banco {v.get('banco')}")
        elif "etapas" in v:
            malos.append(f"{cod}: no deberia tener etapas (origen=regla)")
    checks.append(("valores.carbono por partida (total/unitario/etapas, Sigma etapas = total)", not malos,
                   "todas casan" if not malos else f"{len(malos)} — {malos[0]}"))
    pem_car = float(pres.get("resumen", {}).get("PEM", 0))
    checks.append(("resumen PEM del eje carbono", abs(pem_car - float(eje_exp.get("PEM", 0))) <= ia,
                   f"PEM {pem_car} kgCO2e (esperado {eje_exp.get('PEM')})"))

    # 5 · proyeccion de carbono: DETERMINISMO + INVARIANTE + SEMANTICA
    for v in vistas:
        vid = v.get("id", f"{v.get('eje')}×{v.get('corte')}")
        eje, corte = v.get("eje", "carbono"), v.get("corte")
        f1 = proyectar(pres, modelo, eje, corte)
        f2 = proyectar(pres, modelo, eje, corte)
        checks.append((f"[{vid}] determinista", f1 == f2, "proyectar 2× = misma salida"))
        suma_exp = float(v.get("suma", eje_exp.get("PEM", 0)))
        checks.append((f"[{vid}] invariante Σ == PEM eje", abs(suma_proyeccion(f1) - suma_exp) <= ia,
                       f"Σ {suma_proyeccion(f1)} (esperado {suma_exp})"))
        got_g = {r["grupo"]: r for r in f1}
        exp_g = {g["grupo"]: g for g in v.get("grupos", [])}
        set_ok = set(got_g) == set(exp_g)
        checks.append((f"[{vid}] grupos == expected", set_ok,
                       f"{sorted(got_g)}" if set_ok
                       else f"faltan {sorted(set(exp_g) - set(got_g))} / sobran {sorted(set(got_g) - set(exp_g))}"))
        mg = []
        for grupo, ge in exp_g.items():
            gg = got_g.get(grupo)
            if gg is None:
                continue
            if abs(float(gg["valor_total"]) - float(ge["valor_total"])) > ia:
                mg.append(f"{grupo}:{gg['valor_total']}≠{ge['valor_total']}")
            elif "fuente" in ge and gg.get("fuente") != ge["fuente"]:
                mg.append(f"{grupo}:fuente {gg.get('fuente')}≠{ge['fuente']}")
        checks.append((f"[{vid}] valor_total + fuente por grupo", not mg,
                       "todos casan" if not mg else f"{len(mg)} — {mg[0]}"))

    print_checks(checks)
    return all(ok for _, ok, _ in checks)


CASE_RUNNERS = {"C1": run_case_c1, "C4": run_case_c4, "C3": run_case_c3, "C5": run_case_c5}


# --------------------------------------------------------------------------- #
# Orquestación                                                                #
# --------------------------------------------------------------------------- #
def discover_cases(golden_dir: Path) -> list[Path]:
    cases = []
    for contract_dir in sorted(golden_dir.glob("C*")):
        if not contract_dir.is_dir():
            continue
        for case_dir in sorted(contract_dir.iterdir()):
            if case_dir.is_dir() and (case_dir / "expected.json").exists():
                cases.append(case_dir)
    return cases


def print_checks(checks, indent="    "):
    for name, ok, detail in checks:
        mark = f"{GREEN}[OK]{RST}" if ok else f"{RED}[FALLA]{RST}"
        print(f"{indent}{mark} {name:<38} {DIM}{detail}{RST}")


def run(golden_dir: Path, contracts_dir: Path, schema_only: bool) -> int:
    print(f"{DIM}contracts: {contracts_dir}{RST}")
    print(f"{DIM}golden:    {golden_dir}{RST}\n")

    schemas = discover_schemas(contracts_dir)
    schema_checks = check_schemas_wellformed(schemas)
    print("Paso 1 · esquemas de contrato (JSON Schema válido):")
    print_checks(schema_checks)
    all_ok = all(ok for _, ok, _ in schema_checks)

    if schema_only:
        verdict = "VERDE" if all_ok else "ROJO"
        color = GREEN if all_ok else RED
        print(f"\n{color}SCHEMA-CHECK: {verdict}{RST}")
        return 0 if all_ok else 1

    cases = discover_cases(golden_dir)
    if not cases:
        print(f"{RED}No hay casos golden con expected.json bajo {golden_dir}{RST}")
        return 1

    print(f"\nPaso 2 · golden ({len(cases)} caso/s):")
    for case_dir in cases:
        contract = case_dir.parent.name
        expected = json.loads((case_dir / "expected.json").read_text(encoding="utf-8"))
        tol_path = case_dir / "tolerancias.json"
        tol = json.loads(tol_path.read_text(encoding="utf-8")) if tol_path.exists() else {}
        name = expected.get("caso", case_dir.name)
        runner = CASE_RUNNERS.get(contract)
        print(f"\n  ▶ {name} {DIM}[{contract}]{RST}")
        if runner is None:
            print_checks([("contrato con runner", False,
                           f"'{contract}' sin modo de ejecución en CASE_RUNNERS")])
            all_ok = False
            continue
        all_ok &= runner(case_dir, contracts_dir, expected, tol)

    verdict = "VERDE" if all_ok else "ROJO"
    color = GREEN if all_ok else RED
    print(f"\n{color}VEREDICTO GOLDEN (Llave 1): {verdict}{RST}")
    print(f"{DIM}La firma (Llave 2) es humana; el CI nunca certifica.{RST}")
    return 0 if all_ok else 1


def main() -> int:
    ap = argparse.ArgumentParser(description="Runner de la golden (Llave 1) — multi-contrato.")
    ap.add_argument("--golden-dir", type=Path, default=DEFAULT_GOLDEN_DIR)
    ap.add_argument("--contracts-dir", type=Path, default=None,
                    help="por defecto: <repo>/packages/contracts")
    ap.add_argument("--schema-only", action="store_true",
                    help="solo el paso 1 (esquemas bien formados)")
    args = ap.parse_args()

    golden_dir = args.golden_dir.resolve()
    contracts_dir = (args.contracts_dir or (golden_dir.parent / "contracts")).resolve()
    return run(golden_dir, contracts_dir, args.schema_only)


if __name__ == "__main__":
    raise SystemExit(main())
# Fase I · hilo 1: costura cerrada (compile real narración→IFC). Ver engines/ifc/compile_c1.py.
# Fase II · hilo 1: dispatch por contrato (CASE_RUNNERS); C4 en modo ANCLADO hasta el service (1.1).
# Fase II · hilo 2: costura C4 CERRADA — recompute federar+validar con services/federacion (D10).
# Fase II · hilo 3: emisión BCF 3.0 (tarea 1.2) — paso activado por el expected (C4-FED-02, D14).
# Fase II · hilo 6: derivación IFC (D26/D30) — paso activado por el expected (C4-FED-06); cámara BCF (D29).
# Fase III · hilo 2: contrato C3 (cumplimiento) contract-first — run_case_c3 en modo ANCLADO (D5); el
#                    engine engines/cumplimiento (3.3) antepondrá el recompute contra el MISMO expected.
