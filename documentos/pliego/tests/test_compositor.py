# -*- coding: utf-8 -*-
"""Tests del compositor del Pliego de Condiciones Tecnicas (unidad).

La entrada es el `presupuesto` ANCLADO de GOL-PRE-01 (D4/D6): no se re-ancla, se LEE. La conformidad
profunda vive en el golden GOL-PLI-01 (run_golden.py); aqui, cobertura de unidad: genera, secciones,
partidas con prescripcion + medicion + coste, trazabilidad (GUIDs), carbono FORWARD-OPEN y
DETERMINISMO por contenido (D6).
"""
from __future__ import annotations

import copy
import json
import re
from pathlib import Path

import pytest
from docx import Document

from aqyra_documento_pliego import (SEC_CARATULA, SEC_GENERALES, SEC_PARTICULARES, SEC_TRAZA,
                                    componer_pliego)

_REPO = Path(__file__).resolve().parents[3]
_EXPECTED = _REPO / "packages" / "golden" / "C5" / "GOL-PRE-01" / "expected.json"
_CRITERIO = _REPO / "data" / "packs" / "criterio" / "AQ" / "v2" / "criterio.json"
_TEXTOS = _REPO / "data" / "packs" / "pliego-textos" / "AQ-DEMO" / "v1" / "textos.json"
FECHA = "2026-07-11"


def _cargar(p: Path, clave=None):
    if not p.exists():
        pytest.skip(f"no esta {p}")
    d = json.loads(p.read_text(encoding="utf-8"))
    return d[clave] if clave else d


def _presupuesto() -> dict:
    return _cargar(_EXPECTED, "presupuesto")


def _criterio() -> dict:
    return _cargar(_CRITERIO)


def _textos() -> dict:
    return _cargar(_TEXTOS)


def _extraer(path: Path):
    d = Document(str(path))
    paras = [p.text for p in d.paragraphs]
    tablas = [[[c.text for c in row.cells] for row in t.rows] for t in d.tables]
    return paras, tablas


def _num(s: str):
    s = (s or "").replace("EUR", "").replace("kgCO2e", "").strip()
    if not re.search(r"\d", s):
        return None
    s = re.sub(r"[^0-9.,-]", "", s).replace(".", "").replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


def _fichas_por_partida(tablas):
    """Tablas Concepto/Valor por partida -> {codigo: {concepto: valor_texto}}."""
    fichas = {}
    for t in tablas:
        cab = [c.strip() for c in t[0]] if t else []
        if cab[:2] != ["Concepto", "Valor"]:
            continue
        d = {(fila[0] or "").strip(): (fila[1] or "").strip() for fila in t[1:]}
        cod = d.get("Partida")
        if cod:
            fichas[cod] = d
    return fichas


def _componer(tmp_path, presupuesto=None, nombre="pliego.docx") -> Path:
    salida = tmp_path / nombre
    return componer_pliego(presupuesto or _presupuesto(), _criterio(),
                           {"salida": salida, "fecha": FECHA, "pack_textos": _textos()})


def test_genera_docx(tmp_path):
    ruta = _componer(tmp_path)
    assert ruta.exists() and ruta.stat().st_size > 0


def test_secciones(tmp_path):
    paras, _ = _extraer(_componer(tmp_path))
    texto = "\n".join(paras)
    for marca in (SEC_CARATULA, SEC_GENERALES, SEC_PARTICULARES, SEC_TRAZA):
        assert marca in texto, f"falta la seccion: {marca}"


def test_todas_las_partidas_con_prescripcion(tmp_path):
    pres = _presupuesto()
    paras, tablas = _extraer(_componer(tmp_path))
    texto = "\n".join(paras)
    fichas = _fichas_por_partida(tablas)
    for p in pres["estado_mediciones"]:
        cod = p["codigo"]
        assert cod in fichas, f"partida ausente en el pliego: {cod}"
        # prescripcion: hay una linea "Prescripcion. ..." no vacia y sin el aviso de fallback
        pres_lineas = [x for x in paras if x.startswith("Prescripcion.")]
        assert any(len(x) > len("Prescripcion. ") + 20 for x in pres_lineas), cod
        assert not any("pendiente de completar" in x for x in pres_lineas), \
            f"la partida {cod} cayo al fallback (sin texto base)"


def test_medicion_y_coste_por_partida(tmp_path):
    pres = _presupuesto()
    _, tablas = _extraer(_componer(tmp_path))
    fichas = _fichas_por_partida(tablas)
    for p in pres["estado_mediciones"]:
        f = fichas[p["codigo"]]
        assert abs(_num(f["Cantidad medida"]) - p["cantidad"]) <= 0.001, p["codigo"]
        assert abs(_num(f["Importe (coste)"]) - p["importe"]) <= 0.01, p["codigo"]


def test_trazabilidad_guids(tmp_path):
    pres = _presupuesto()
    paras, tablas = _extraer(_componer(tmp_path))
    blob = "\n".join(paras) + "\n" + "\n".join(
        " ".join(c for fila in t for c in fila) for t in tablas)
    for p in pres["estado_mediciones"]:
        for guid in p.get("trazabilidad", []):
            assert guid in blob, f"GUID ausente ({p['codigo']}): {guid}"


def test_carbono_forward_open(tmp_path):
    # sin carbono: no aparece "Huella de carbono"
    paras0, tablas0 = _extraer(_componer(tmp_path, nombre="sin.docx"))
    blob0 = "\n".join(c for t in tablas0 for fila in t for c in fila)
    assert "Huella de carbono" not in blob0
    # con carbono inyectado en una partida: aparece la huella y una etapa
    pres = copy.deepcopy(_presupuesto())
    pres["estado_mediciones"][0]["valores"] = {
        "carbono": {"unitario": 10.0, "total": 33.3, "unidad": "kgCO2e",
                    "etapas": {"A1A3": 30.0, "A4A5": 3.3}}}
    _, tablas1 = _extraer(_componer(tmp_path, presupuesto=pres, nombre="con.docx"))
    blob1 = "\n".join(c for t in tablas1 for fila in t for c in fila)
    assert "Huella de carbono" in blob1 and "Carbono A1A3" in blob1


def test_determinismo_contenido(tmp_path):
    a = _extraer(_componer(tmp_path, nombre="a.docx"))
    b = _extraer(_componer(tmp_path, nombre="b.docx"))
    assert a == b
