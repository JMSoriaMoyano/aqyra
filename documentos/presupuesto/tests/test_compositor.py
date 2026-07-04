# -*- coding: utf-8 -*-
"""Tests del compositor del Documento de Presupuesto (unidad).

La entrada es el `presupuesto` ANCLADO de GOL-PRE-01 (D4): no se re-ancla, se LEE. La conformidad
profunda vive en el golden GOL-DOC-01 (run_golden.py); aquí, cobertura de unidad: genera, secciones,
partidas/importes, totales y DETERMINISMO por contenido (D3).
"""
from __future__ import annotations

import json
import re
from pathlib import Path

import pytest
from docx import Document

from aqyra_documento_presupuesto import (SEC_CP1, SEC_CP2, SEC_MEDICIONES, SEC_RESUMEN,
                                         componer_documento)

_REPO = Path(__file__).resolve().parents[3]
_EXPECTED = _REPO / "packages" / "golden" / "C5" / "GOL-PRE-01" / "expected.json"
FECHA = "2026-07-04"


def _presupuesto() -> dict:
    if not _EXPECTED.exists():
        pytest.skip(f"no está GOL-PRE-01/expected.json en {_EXPECTED}")
    return json.loads(_EXPECTED.read_text(encoding="utf-8"))["presupuesto"]


def _extraer(path: Path):
    """(párrafos, tablas) del .docx — el mismo tipo de extracción que el golden (no OCR)."""
    d = Document(str(path))
    paras = [p.text for p in d.paragraphs]
    tablas = [[[c.text for c in row.cells] for row in t.rows] for t in d.tables]
    return paras, tablas


def _num(s: str):
    """'1.234,56' -> 1234.56 ; '' o texto -> None."""
    s = (s or "").replace("EUR", "").replace("€", "").strip()
    if not re.search(r"\d", s):
        return None
    s = s.replace(".", "").replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


def _componer(tmp_path, nombre="doc.docx") -> Path:
    salida = tmp_path / nombre
    return componer_documento(_presupuesto(), {"salida": salida, "fecha": FECHA})


def test_genera_docx(tmp_path):
    ruta = _componer(tmp_path)
    assert ruta.exists() and ruta.stat().st_size > 0


def test_cinco_secciones(tmp_path):
    ruta = _componer(tmp_path)
    paras, _ = _extraer(ruta)
    texto = "\n".join(paras)
    for marca in ("EJECUCION POR CONTRATA", SEC_MEDICIONES, SEC_CP1, SEC_CP2, SEC_RESUMEN):
        assert marca in texto, f"falta la sección: {marca}"


def test_partidas_e_importes(tmp_path):
    pres = _presupuesto()
    ruta = _componer(tmp_path)
    _, tablas = _extraer(ruta)
    # importe por partida leído de las tablas de mediciones (primera celda = código)
    imp_doc = {}
    for t in tablas:
        for fila in t:
            cod = (fila[0] or "").strip()
            val = _num(fila[-1])
            if cod and val is not None and not cod.startswith(("PARCIAL", "C0")):
                imp_doc[cod] = val
    for p in pres["estado_mediciones"]:
        assert p["codigo"] in imp_doc, f"partida ausente: {p['codigo']}"
        assert abs(imp_doc[p["codigo"]] - p["importe"]) <= 0.01, p["codigo"]


def test_precio_en_letra_presente(tmp_path):
    pres = _presupuesto()
    ruta = _componer(tmp_path)
    _, tablas = _extraer(ruta)
    letras = {}
    for t in tablas:
        cab = [c.strip() for c in t[0]] if t else []
        if "Precio en letra" in cab:
            for fila in t[1:]:
                letras[(fila[0] or "").strip()] = (fila[-1] or "").strip()
    for c in pres["cuadro_precios_1"]:
        assert letras.get(c["codigo"]), f"sin precio en letra: {c['codigo']}"


def test_resumen_pem_pec(tmp_path):
    pres = _presupuesto()
    ruta = _componer(tmp_path)
    _, tablas = _extraer(ruta)
    concep = {}
    for t in tablas:
        cab = [c.strip() for c in t[0]] if t else []
        if cab[:1] == ["Concepto"]:
            for fila in t[1:]:
                concep[(fila[0] or "").strip()] = _num(fila[-1])
    r = pres["resumen"]
    pem = next(v for k, v in concep.items() if "(PEM)" in k)
    pec = next(v for k, v in concep.items() if "(PEC)" in k)
    assert abs(pem - r["PEM"]) <= 0.01
    assert abs(pec - r["PEC"]) <= 0.01


def test_determinismo_contenido(tmp_path):
    """D3: componer 2x -> mismo CONTENIDO extraído (no bytes)."""
    a = _extraer(_componer(tmp_path, "a.docx"))
    b = _extraer(_componer(tmp_path, "b.docx"))
    assert a == b
