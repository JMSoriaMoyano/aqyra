# -*- coding: utf-8 -*-
"""Formato del despacho — MODULO COMPARTIDO de la capa de documentos (D-EX-1, opcion A).

Extraccion del formato del despacho que hasta ahora se ESPEJABA en cada compositor
(documentos/presupuesto/formato.py y documentos/pliego/formato.py). Al aparecer el TERCER
consumidor —el compositor de la proyeccion de valor en documentos/export— se materializa la
nota forward que dejo la D1 del pliego: «si aparece un tercer documento, extraer a un modulo
comun». Este modulo lo consume SOLO el paquete nuevo documentos/export; presupuesto y pliego
siguen CONGELADOS con su copia espejada (son paquetes firmados/anclados; no se re-ancla nada).

Hermetico (solo python-docx), para que el gate de CI no necesite el skill memoria-calculo-despacho:
- A4, margenes 25 mm; Arial 11 (cuerpo) y Arial 9,5 (tablas).
- Parrafos justificados; titulos a la izquierda, en azul.
- Tablas SIN cebra, cabecera azul claro (#D9E2F3) que SE REPITE en saltos de pagina,
  sin lineas verticales salvo en la cabecera.
- Caratula con la etiqueta destacada (verde).

API: nuevo_doc, caratula, H1, H2, P, tabla, total_destacado, cabecera_pie, cierre, fmt_eur, fmt_num.
"""
from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# ---- paleta del despacho ----
AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x59, 0x59, 0x59)
ROJO = RGBColor(0xB0, 0x00, 0x00)
VERDE = RGBColor(0x2E, 0x7D, 0x32)
HDR_FILL = "D9E2F3"     # cabecera azul claro
VLINE = "9CB7D4"        # vertical de cabecera
HLINE = "BFBFBF"        # horizontales
TBL_PT = 9.5            # Arial dentro de tablas

_SWAP = "\x00"          # marcador temporal para intercambiar separadores


# ============================ numeros (formato espanol) ============================
def fmt_num(x: float, dec: int = 2) -> str:
    """1234.5 -> '1.234,50' (miles con punto, decimales con coma). Determinista."""
    s = f"{float(x):,.{dec}f}"                     # '1,234.50' (miles ',' decimal '.')
    return s.replace(",", _SWAP).replace(".", ",").replace(_SWAP, ".")


def fmt_eur(x: float) -> str:
    return f"{fmt_num(x, 2)} EUR"


# ============================ documento base ============================
def nuevo_doc():
    """Documento con estilos y pagina A4 del despacho. Devuelve (doc, sec)."""
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Arial"; n.font.size = Pt(11)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    n.paragraph_format.line_spacing = 1.15
    n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for lvl, sz in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11.5)]:
        st = doc.styles[lvl]
        st.font.name = "Arial"; st.font.size = Pt(sz); st.font.bold = True
        st.font.color.rgb = AZUL
        st.paragraph_format.space_before = Pt(12); st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sec = doc.sections[0]
    sec.page_height = Mm(297); sec.page_width = Mm(210)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Mm(25))
    sec.different_first_page_header_footer = True
    return doc, sec


# ============================ utilidades XML ============================
def _field(par, instr):
    r = par.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "-"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for e in (fb, it, sep, t, end):
        r._r.append(e)


def _hrule(par, color="1F4E79", size=12):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "4"); b.set(qn("w:color"), color)
    pbdr.append(b); pPr.append(pbdr)


def _header_cell(cell, fill=HDR_FILL, vcolor=VLINE):
    tcPr = cell._tc.get_or_add_tcPr()
    for tag in ("w:tcBorders", "w:shd"):
        for e in tcPr.findall(qn(tag)):
            tcPr.remove(e)
    tb = OxmlElement("w:tcBorders")
    for t in ("left", "right"):
        e = OxmlElement("w:" + t)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), vcolor)
        tb.append(e)
    tcPr.append(tb)
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)
    cant = OxmlElement("w:cantSplit"); cant.set(qn("w:val"), "true"); trPr.append(cant)


def _fmt(run, bold=False):
    run.font.name = "Arial"; run.font.size = Pt(TBL_PT); run.bold = bold


# ============================ bloques de contenido ============================
def H1(doc, t):
    return doc.add_paragraph(t, style="Heading 1")


def H2(doc, t):
    return doc.add_paragraph(t, style="Heading 2")


def P(doc, t):
    par = doc.add_paragraph(); par.add_run(t)
    return par


def tabla(doc, cabecera, filas, anchos_mm=None, alineaciones=None, centrar=True):
    """Tabla del despacho: Arial 9,5, sin cebra, cabecera azul repetida, sin verticales.

    `alineaciones`: lista opcional (por columna) de WD_ALIGN_PARAGRAPH para las filas de datos.
    """
    ncol = len(cabecera)
    t = doc.add_table(rows=1, cols=ncol)
    if centrar:
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = t._tbl.tblPr
    for b in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(b)
    borders = OxmlElement("w:tblBorders")
    for tag, val, col in [("top", "single", HLINE), ("bottom", "single", HLINE),
                          ("insideH", "single", HLINE), ("left", "nil", "auto"),
                          ("right", "nil", "auto"), ("insideV", "nil", "auto")]:
        e = OxmlElement("w:" + tag); e.set(qn("w:val"), val)
        e.set(qn("w:sz"), "4" if val == "single" else "0")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), col)
        borders.append(e)
    tblPr.append(borders)
    hrow = t.rows[0]; _repeat_header(hrow)
    for i, txt in enumerate(cabecera):
        _header_cell(hrow.cells[i])
        _fmt(hrow.cells[i].paragraphs[0].add_run(str(txt)), bold=True)
    for fila in filas:
        cells = t.add_row().cells
        for i, val in enumerate(fila):
            run = cells[i].paragraphs[0].add_run(str(val))
            _fmt(run)
            if alineaciones and i < len(alineaciones) and alineaciones[i] is not None:
                cells[i].paragraphs[0].alignment = alineaciones[i]
    if anchos_mm:
        for r in t.rows:
            for i, w in enumerate(anchos_mm):
                r.cells[i].width = Mm(w)
    return t


def total_destacado(doc, etiqueta, valor_texto):
    """Etiqueta destacada (verde) en la caratula."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r0 = p.add_run(etiqueta + "  ")
    r0.bold = True; r0.font.size = Pt(13); r0.font.color.rgb = AZUL
    r = p.add_run(valor_texto)
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = VERDE
    return p


def caratula(doc, titulo, subtitulo, campos, total_etq=None, total_val=None, nota=None):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); r = p.add_run(titulo)
    r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = AZUL
    _hrule(p)
    s = doc.add_paragraph(); r = s.add_run(subtitulo)
    r.font.size = Pt(16); r.font.color.rgb = GRIS
    if total_etq is not None:
        total_destacado(doc, total_etq, total_val)
    doc.add_paragraph()
    tbl = doc.add_table(rows=0, cols=2)
    for k, v in campos:
        row = tbl.add_row().cells
        rk = row[0].paragraphs[0].add_run(k); rk.bold = True
        row[1].paragraphs[0].add_run(v)
        row[0].width = Mm(45); row[1].width = Mm(115)
    for _ in range(2):
        doc.add_paragraph()
    dp = doc.add_paragraph(); _hrule(dp, color="B00000", size=8)
    rr = dp.add_run(nota or (
        "Documento generado de forma DETERMINISTA a partir de salidas autoritativas de Aqyra. "
        "Debe ser revisado y firmado por tecnico competente antes de su uso."))
    rr.italic = True; rr.font.color.rgb = ROJO; rr.font.size = Pt(9.5)
    doc.add_page_break()


def cabecera_pie(sec, proyecto, autor, doc_label="Pliego"):
    hp = sec.header.paragraphs[0]; hp.text = ""
    r = hp.add_run(proyecto); r.font.size = Pt(8.5); r.font.color.rgb = GRIS
    _hrule(hp, color="1F4E79", size=6)
    fp = sec.footer.paragraphs[0]; fp.text = ""
    fp.paragraph_format.tab_stops.add_tab_stop(Mm(160), WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run(autor + " . " + doc_label + "\t"); r.font.size = Pt(8.5); r.font.color.rgb = GRIS
    r2 = fp.add_run("Pag. "); r2.font.size = Pt(8.5); r2.font.color.rgb = GRIS
    _field(fp, "PAGE")
    r3 = fp.add_run(" de "); r3.font.size = Pt(8.5); r3.font.color.rgb = GRIS
    _field(fp, "NUMPAGES")


def cierre(doc, txt):
    doc.add_paragraph()
    cl = doc.add_paragraph(); _hrule(cl, color="B00000", size=8)
    rr = cl.add_run(txt); rr.italic = True; rr.font.color.rgb = ROJO; rr.font.size = Pt(9.5)
