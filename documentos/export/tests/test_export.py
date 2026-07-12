# -*- coding: utf-8 -*-
"""Tests del export firmable (unidad). Consumidor PRIMARIO (contractual): el presupuesto de obra.

El artefacto es el `salida-presupuesto` ANCLADO de GOL-PRE-01: no se re-ancla, se LEE. La conformidad
profunda vive en GOL-EXP-01 (presupuesto) y GOL-EXP-02 (proyeccion); aqui, cobertura de unidad.
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest

import aqyra_documento_export as X
from aqyra_documento_export import manifiesto as M
from aqyra_documento_comun import formato as F

_REPO = Path(__file__).resolve().parents[3]
_PRES = _REPO / "packages" / "golden" / "C5" / "GOL-PRE-01" / "expected.json"
_PROY = _REPO / "packages" / "golden" / "C5" / "GOL-PRE-03" / "expected.json"
_ESQ = _REPO / "documentos" / "export" / "esquemas"
SELLO = "2026-07-12"


def _presupuesto() -> dict:
    if not _PRES.exists():
        pytest.skip(f"no esta {_PRES}")
    return json.loads(_PRES.read_text(encoding="utf-8"))["presupuesto"]


def _descriptor_pres() -> dict:
    return {"artefacto": {"tipo": "presupuesto-obra", "id": "GOL-PRE-01"},
            "formatos": ["word", "pdf", "bc3", "xlsx"], "plantilla": "presupuesto-contractual",
            "sello_tiempo": SELLO,
            "versiones_ancladas": {"contrato_c5": "C5-presupuesto", "engine": "engines/presupuesto",
                                   "banco": "AQ-DEMO/v1", "criterio": "AQ/v1", "schema": "salida-presupuesto"}}


def _componer(tmp_path, nombre="a") -> Path:
    return X.componer_export(_presupuesto(), _descriptor_pres(), {"salida": tmp_path / nombre})


def test_genera_bundle_contractual(tmp_path):
    b = _componer(tmp_path)
    for n in ("Presupuesto.docx", "Presupuesto-firmable.pdf", "Presupuesto.bc3", "Mediciones.xlsx",
              X.NOMBRE_MANIFIESTO):
        assert (b / n).exists() and (b / n).stat().st_size > 0, n


def test_esquemas(tmp_path):
    jsonschema = pytest.importorskip("jsonschema")
    desc_s = json.loads((_ESQ / "descriptor-export.schema.json").read_text(encoding="utf-8"))
    man_s = json.loads((_ESQ / "manifiesto-export.schema.json").read_text(encoding="utf-8"))
    jsonschema.validate(_descriptor_pres(), desc_s)
    b = _componer(tmp_path)
    jsonschema.validate(json.loads((b / X.NOMBRE_MANIFIESTO).read_text(encoding="utf-8")), man_s)


def test_manifiesto_integro(tmp_path):
    pres = _presupuesto()
    b = X.componer_export(pres, _descriptor_pres(), {"salida": tmp_path / "a"})
    man = json.loads((b / X.NOMBRE_MANIFIESTO).read_text(encoding="utf-8"))
    ok, det = M.integridad(man, pres)
    assert ok, det
    assert man["artefacto"]["content_sha256"] == M.hash_canonico(pres)


def test_pdf_contractual_con_justificacion(tmp_path):
    from pypdf import PdfReader
    pres = _presupuesto()
    b = X.componer_export(pres, _descriptor_pres(), {"salida": tmp_path / "a"})
    pt = "\n".join((pg.extract_text() or "") for pg in PdfReader(str(b / "Presupuesto-firmable.pdf")).pages)
    assert F.fmt_num(pres["resumen"]["PEC"], 2) in pt                 # PEC
    for p in pres["estado_mediciones"]:
        assert p["codigo"] in pt                                     # partidas
    # justificacion de medicion: criterio + GUIDs trazados al modelo
    assert "NetVolume" in pt
    assert pres["estado_mediciones"][0]["trazabilidad"][0] in pt


def test_bc3_licitacion(tmp_path):
    pres = _presupuesto()
    b = X.componer_export(pres, _descriptor_pres(), {"salida": tmp_path / "a"})
    bc3 = (b / "Presupuesto.bc3").read_text(encoding="utf-8")
    assert bc3.startswith("~V") and "~C" in bc3 and "~M" in bc3
    for p in pres["estado_mediciones"]:
        assert p["codigo"] in bc3


def test_iscertified_sin_firma(tmp_path):
    b = _componer(tmp_path)
    assert not X.es_certificado(b)
    assert X.estado_firmable(b) == "computed"


def test_determinismo_contenido(tmp_path):
    from docx import Document
    from openpyxl import load_workbook
    from pypdf import PdfReader

    def docx_c(p):
        d = Document(str(p)); return ([q.text for q in d.paragraphs],
                                      [[[c.text for c in r.cells] for r in t.rows] for t in d.tables])

    def xlsx_c(p):
        wb = load_workbook(str(p), read_only=True)
        return {ws.title: [list(r) for r in ws.iter_rows(values_only=True)] for ws in wb.worksheets}

    def pdf_t(p):
        return "\n".join((pg.extract_text() or "") for pg in PdfReader(str(p)).pages)

    a, b = _componer(tmp_path, "a"), _componer(tmp_path, "b")
    assert docx_c(a / "Presupuesto.docx") == docx_c(b / "Presupuesto.docx")
    assert xlsx_c(a / "Mediciones.xlsx") == xlsx_c(b / "Mediciones.xlsx")
    assert pdf_t(a / "Presupuesto-firmable.pdf") == pdf_t(b / "Presupuesto-firmable.pdf")
    assert (a / "Presupuesto.bc3").read_bytes() == (b / "Presupuesto.bc3").read_bytes()
    assert (a / X.NOMBRE_MANIFIESTO).read_bytes() == (b / X.NOMBRE_MANIFIESTO).read_bytes()


def test_proyeccion_secundario(tmp_path):
    """El consumidor de gestion (proyeccion) sigue vivo como export secundario."""
    if not _PROY.exists():
        pytest.skip("no esta GOL-PRE-03")
    art = json.loads(_PROY.read_text(encoding="utf-8"))
    desc = {"artefacto": {"tipo": "proyeccion-valor", "id": "GOL-PRE-03"},
            "formatos": ["word", "xlsx", "pdf"], "sello_tiempo": SELLO,
            "versiones_ancladas": {"criterio": "AQ/v2"}}
    b = X.componer_export(art, desc, {"salida": tmp_path / "proy"})
    assert (b / "informe-proyeccion.docx").exists()
    assert not X.es_certificado(b)
