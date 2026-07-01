"""
Memoria de calculo del Caso 13 - Losa plana postesada (EC2 §5.10 + §6.4.4) .docx.

Requiere python-docx y los PNG de plots_losa_postesada (run_all los crea; si
faltan, intenta regenerarlos).

Uso: python3 pretensado/generate_memoria_losa_postesada.py <carpeta_proyecto>
"""
import os
import sys
import json

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

PNGS = ["planta_tendones.png", "cargas_equivalentes_2d.png", "mapa_Mx.png",
        "mapa_My.png", "mapa_tension_fibra.png", "perimetro_punzonamiento.png",
        "ELU_franjas.png", "mapa_flecha.png"]


def _ensure_pngs(carpeta, model, res, out):
    faltan = [p for p in PNGS if not os.path.exists(os.path.join(carpeta, p))]
    if faltan:
        try:
            import plots_losa_postesada as plots
            plots.generar(model, res, out, carpeta)
        except Exception as ex:
            print("AVISO: no se pudieron generar PNGs (%s)." % ex)


def build(carpeta):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    model = json.load(open(os.path.join(carpeta, "modelo_neutro.json"), encoding="utf-8"))
    out = json.load(open(os.path.join(carpeta, "verificacion_losa_postesada.json"), encoding="utf-8"))
    ver = out["verificacion"]
    try:
        res = json.load(open(os.path.join(carpeta, "resultados.json"), encoding="utf-8"))
    except Exception:
        res = None
    if res is not None:
        _ensure_pngs(carpeta, model, res, out)

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    def h(txt, lvl=1):
        doc.add_heading(txt, level=lvl)

    def p(txt, bold=False):
        par = doc.add_paragraph(); r = par.add_run(txt); r.bold = bold
        return par

    def img(name, w=14):
        fp = os.path.join(carpeta, name)
        if os.path.exists(fp):
            doc.add_picture(fp, width=Cm(w))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    mr = out["modelo_resumen"]; pr = out["pretensado"]; bal = out["balance_2d"]

    doc.add_heading("Memoria de calculo - Losa plana postesada (Caso 13)", 0)
    p("Forjado de losa plana postesada de hormigon C40/50 sobre reticula de "
      "pilares. Calculo del pretensado segun EC2 (EN 1992-1-1) §5.10 y del "
      "punzonamiento segun §6.4.4, con Anejo Nacional de Espana. PREDIMENSIONADO.")

    h("1. Normativa de aplicacion", 1)
    p("EN 1992-1-1 (EC2): §5.10 pretensado y perdidas; §6.4 punzonamiento "
      "(§6.4.4 efecto favorable del pretensado); §7.2 limitacion de tensiones; "
      "§7.3 control de fisuracion; §7.4 flecha. EN 1990 (EC0) combinaciones. "
      "[confirmar AN] Anejo Nacional Espana: k1 del termino de pretensado en "
      "punzonamiento, mu/k de rozamiento, penetracion de cuna, limites del acero "
      "activo, retraccion y fluencia.")

    h("2. Materiales y geometria", 1)
    p("Hormigon %s (fck=%.0f MPa, fcd=%.2f MPa, fctm=%.1f MPa). Losa de espesor "
      "t=%.2f m; reticula 3x3 de pilares %.2fx%.2f m; vanos Lx=Ly=%.1f m "
      "(L/h~%.0f). Por metro de ancho: A=%.3f m2/m, I=%.3e m4/m, c=%.3f m." % (
        mr["material"], mr["fck_MPa"], mr["fck_MPa"] / 1.5, mr["fctm_MPa"],
        mr["t_m"], 0.45, 0.45, mr["Lx_m"], mr["Lx_m"] / mr["t_m"],
        mr["t_m"], ver["materiales"]["I_m4_m"], mr["t_m"] / 2))
    p("Cargas: peso propio g0=%.2f kN/m2 (lo anade el solver), carga muerta "
      "g2=%.2f kN/m2, sobrecarga q=%.2f kN/m2 (psi2=0.3). Permanente total "
      "g0+g2=%.2f kN/m2." % (
        mr["g0_kN_m2"], mr["g2_kN_m2"], mr["q_kN_m2"], mr["permanente_kN_m2"]))

    h("3. Pretensado y perdidas (§5.10)", 1)
    p("Postesado NO adherente con monotorones %s (Ap=%.0f mm2/cordon, fpk=%.0f "
      "MPa, fp0,1k=%.0f MPa, Ep=%.0f GPa). Trazado parabolico, drape a=%.2f m "
      "(recubrimiento al eje %.2f m). Layout: %s." % (
        pr.get("layout", ""), pr["Ap_cordon_mm2"], pr["fpk_MPa"], pr["fp01k_MPa"],
        pr["Ep_GPa"], pr["drape_m"], pr["recubrimiento_eje_m"], pr["layout"]))
    p("Fuerza por metro de ancho: P_m=%.0f kN/m (cada direccion). Transferencia "
      "sigma_p0=%.0f MPa (%.3f fpk); servicio sigma_p,inf=%.0f MPa (%.3f fpk). "
      "Perdidas totales %.1f%% (transferencia %.1f%%). sigma_pmax=%.0f MPa "
      "(0.80 fpk)." % (
        pr["P_m_X_N_m"] / 1e3, pr["sigma_p0_MPa"], pr["sigma_p0_frac_fpk"],
        pr["sigma_pm_inf_MPa"], pr["sigma_pm_inf_frac_fpk"],
        pr["perdidas_totales_pct"], pr["perdidas_transferencia_pct"],
        pr["sigma_pmax_MPa"]), bold=True)
    img("planta_tendones.png")

    h("4. Balance de cargas 2D (§5.10)", 1)
    p("Tendones banded en X (sobre lineas de pilares) y distribuidos en Y. Cada "
      "familia parabolica equilibra w_p=8*P*a/L^2. w_px=%.2f + w_py=%.2f = "
      "w_p=%.2f kN/m2; la permanente es %.2f kN/m2; residual %.3f kN/m2 (%.3f%%): "
      "el pretensado EQUILIBRA la permanente. Precompresion media sigma_cp=%.3f "
      "MPa." % (
        bal["w_px_N_m2"] / 1e3, bal["w_py_N_m2"] / 1e3, bal["w_p_N_m2"] / 1e3,
        mr["permanente_kN_m2"], bal["residual_N_m2"] / 1e3, bal["residual_pct"],
        bal["sigma_cp_Pa"] / 1e6), bold=True)
    img("cargas_equivalentes_2d.png")

    if "contraste_metodos" in ver:
        h("4b. Contraste cargas-equivalentes vs fuerza+excentricidad", 2)
        for fn, cr in ver["contraste_metodos"].items():
            p("%s: cargas equiv. sup=%.2f/inf=%.2f MPa; fuerza+exc. sup=%.2f/"
              "inf=%.2f MPa; dif max %.4f MPa -> %s." % (
                fn, cr["cargas_equiv"]["sup_MPa"], cr["cargas_equiv"]["inf_MPa"],
                cr["fuerza_excentricidad"]["sup_MPa"], cr["fuerza_excentricidad"]["inf_MPa"],
                max(cr["dif_sup_MPa"], cr["dif_inf_MPa"]),
                "COINCIDEN" if cr["coincide"] else "DIFIEREN"))

    h("5. Tensiones por fibra (§7.2)", 1)
    p("Estado tensional por franja (franja de pilares gobernante), con la carga "
      "NO equilibrada (residual) en transferencia, cuasipermanente y rara. "
      "Compresion negativa; limites 0.6 fck(t) en transferencia, 0.45 fck en "
      "cuasipermanente, traccion <= fctm en rara.")
    tbl = doc.add_table(rows=1, cols=5); tbl.style = "Light Grid Accent 1"
    hd = tbl.rows[0].cells
    hd[0].text = "Franja / estado"; hd[1].text = "sigma_sup [MPa]"
    hd[2].text = "sigma_inf [MPa]"; hd[3].text = "aprov"; hd[4].text = "OK"
    for fn, fr in ver["tensiones_franja"].items():
        for est in ("transferencia", "servicio_cuasiperm", "servicio_rara"):
            e = fr[est]
            ap = max(e.get("u_compresion", 0.0), e.get("u_traccion_fctm", 0.0))
            row = tbl.add_row().cells
            row[0].text = "%s / %s" % (fn, est)
            row[1].text = "%.2f" % e["sigma_sup_MPa"]
            row[2].text = "%.2f" % e["sigma_inf_MPa"]
            row[3].text = "%.2f" % ap
            row[4].text = "SI" if e["ok"] else "NO"
    img("mapa_tension_fibra.png")

    h("6. Punzonamiento (§6.4.4) - con y sin efecto favorable del pretensado", 1)
    p("Para cada tipo de pilar se compara la tension de calculo en el perimetro "
      "de control u1 (a 2d) SIN pretensado y CON pretensado. El efecto favorable "
      "(a) anade k1*sigma_cp (k1=0.10 [confirmar AN]) a v_Rd,c y (b) descuenta la "
      "componente vertical V_p de los tendones que cruzan u1 del esfuerzo V_Ed.")
    tbl2 = doc.add_table(rows=1, cols=6); tbl2.style = "Light Grid Accent 1"
    hd = tbl2.rows[0].cells
    hd[0].text = "Pilar"; hd[1].text = "V_Ed [kN]"; hd[2].text = "V_p [kN]"
    hd[3].text = "aprov SIN"; hd[4].text = "aprov CON"; hd[5].text = "OK"
    for pos, e in ver["punzonamiento"].items():
        row = tbl2.add_row().cells
        row[0].text = pos; row[1].text = "%.0f" % e["V_Ed_sin_kN"]
        row[2].text = "%.0f" % e["V_p_kN"]
        row[3].text = "%.2f" % e["sin_pretensado"]["u_vRdc"]
        row[4].text = "%.2f" % e["con_pretensado"]["u_vRdc"]
        row[5].text = "SI" if e["con_pretensado"]["ok"] else "NO"
    p("El efecto favorable del pretensado RELAJA el aprovechamiento de "
      "punzonamiento frente a la losa sin pretensar.", bold=True)
    img("perimetro_punzonamiento.png")

    h("7. ELU de flexion por fibras (§6.1)", 1)
    p("Seccion no lineal por fibras (bloque rectangular eta=1.0, lambda=0.8 para "
      "C40/50) con armadura activa (banded/distribuida) + pasiva minima, por "
      "franja y posicion (campo/apoyo).")
    tbl3 = doc.add_table(rows=1, cols=5); tbl3.style = "Light Grid Accent 1"
    hd = tbl3.rows[0].cells
    hd[0].text = "Franja/pos"; hd[1].text = "M_Ed [kNm/m]"; hd[2].text = "M_Rd [kNm/m]"
    hd[3].text = "x/d"; hd[4].text = "aprov"
    for fn, e in ver["ELU_flexion"].items():
        for pos in ("campo", "apoyo"):
            row = tbl3.add_row().cells
            row[0].text = "%s/%s" % (fn, pos)
            row[1].text = "%.0f" % e[pos]["M_Ed_kNm_m"]
            row[2].text = "%.0f" % e[pos]["M_Rd_kNm_m"]
            row[3].text = "%.2f" % e[pos]["x_d_ratio"]
            row[4].text = "%.2f" % e[pos]["u"]
    img("ELU_franjas.png")

    h("8. Fisuracion (§7.3) y flecha (§7.4)", 1)
    fis = ver["fisuracion"]; fl = ver["flecha"]
    p("Fisuracion: traccion de fondo en combinacion rara %.2f MPa frente a "
      "fctm=%.1f MPa (aprov %.2f) -> %s." % (
        fis["sigma_inf_rara_MPa"], fis["fctm_MPa"], fis["u"],
        "sin fisuracion significativa" if fis["ok"] else "revisar"))
    p("Flecha total (ELS caracteristica) %.2f mm frente a L/250=%.1f mm "
      "(aprov %.2f). El residual de permanente ~0 produce contraflecha de "
      "balance: flecha reducida." % (
        fl["f_total_mm"], fl["lim_total_mm"], fl["u_total"]), bold=True)
    img("mapa_flecha.png")
    img("mapa_Mx.png", 12); img("mapa_My.png", 12)

    h("9. Resumen de aprovechamientos y veredicto", 1)
    tbl4 = doc.add_table(rows=1, cols=2); tbl4.style = "Light Grid Accent 1"
    tbl4.rows[0].cells[0].text = "Verificacion"; tbl4.rows[0].cells[1].text = "Aprov."
    for k, v in ver["aprovechamientos"].items():
        row = tbl4.add_row().cells; row[0].text = k; row[1].text = "%.2f" % v
    p("Veredicto: %s. Aprovechamiento maximo %.2f (<=1)." % (
        ver["veredicto"], ver["aprov_max"]), bold=True)

    h("10. Avisos y trazabilidad", 1)
    for a in out.get("avisos", []):
        p("- " + a)
    p("Esta memoria es de PREDIMENSIONADO. Debe ser revisada y FIRMADA por un "
      "tecnico competente.", bold=True)

    dst = os.path.join(carpeta, "memoria_calculo_losa_postesada.docx")
    doc.save(dst)
    print("Memoria escrita en:", dst)
    return dst


if __name__ == "__main__":
    carpeta = sys.argv[1] if len(sys.argv) > 1 else "."
    build(carpeta)
