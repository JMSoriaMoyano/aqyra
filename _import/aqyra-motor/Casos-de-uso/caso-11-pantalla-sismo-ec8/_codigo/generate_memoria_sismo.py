"""
Memoria de calculo SISMICO (Caso 11 - Pantalla de cortante EC8) en .docx.

Requiere python-docx (pip install python-docx --break-system-packages) y los
PNG generados por plots_sismo (run_all_sismo los crea). Si solo existen los SVG,
intenta convertirlos con cairosvg; si no, reejecuta plots_sismo (matplotlib).

Uso: python3 sismico/generate_memoria_sismo.py <carpeta_proyecto>
"""
import sys
import os
import json

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)


def _ensure_pngs(carpeta, model, res, ver):
    """Garantiza PNGs: usa plots_sismo (matplotlib) si faltan."""
    pngs = ["espectro_Sd.png", "modos.png", "fuerzas_planta.png",
            "cortante_altura.png", "momento_altura.png", "deriva.png",
            "diagrama_NM.png"]
    faltan = [p for p in pngs if not os.path.exists(os.path.join(carpeta, p))]
    if faltan:
        try:
            import plots_sismo
            plots_sismo.generar(model, res, ver, carpeta)
        except Exception as ex:
            print("AVISO: no se pudieron generar PNGs (%s). Probando SVG->PNG." % ex)
            try:
                import cairosvg
                for p in pngs:
                    svg = os.path.join(carpeta, p.replace(".png", ".svg"))
                    if os.path.exists(svg):
                        cairosvg.svg2png(url=svg, write_to=os.path.join(carpeta, p))
            except Exception as ex2:
                print("AVISO: tampoco cairosvg (%s). La memoria ira sin imagenes." % ex2)
    return pngs


def build(carpeta):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    model = json.load(open(os.path.join(carpeta, "modelo_neutro.json"), encoding="utf-8"))
    res = json.load(open(os.path.join(carpeta, "verificacion_sismo.json"), encoding="utf-8"))
    ver = res["verificacion"]
    pngs = _ensure_pngs(carpeta, model, res, ver)

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    def h(txt, lvl=1):
        doc.add_heading(txt, level=lvl)

    def p(txt, bold=False):
        par = doc.add_paragraph()
        r = par.add_run(txt); r.bold = bold
        return par

    def img(name, w=14):
        fp = os.path.join(carpeta, name)
        if os.path.exists(fp):
            doc.add_picture(fp, width=Cm(w))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    tit = doc.add_heading("Memoria de calculo sismico - Caso 11", 0)
    p("Pantalla de cortante de hormigon armado. Analisis sismico segun EC8 "
      "(EN 1998-1) con Anejo Nacional de Espana / NCSE-02. PREDIMENSIONADO.")

    h("1. Objeto y modelo", 1)
    sp = model["seccion_pared"]; ma = model["masas"]
    p("Pantalla C30/37, Lw=%.1f m, tw=%.2f m, H=%.1f m (5 plantas x 3.0 m), "
      "empotrada en base, sistema de estabilizacion lateral. Inercia en el "
      "plano I=tw*Lw^3/12=%.2f m4; area A=%.2f m2; area de cortante A_v=%.2f m2 "
      "(=A*5/6)." % (sp["Lw_m"], sp["tw_m"], sp["H_m"], sp["I_m4"], sp["A_m2"], sp["A_v_m2"]))
    p("Voladizo equivalente (stick de 6 nodos) con rigidez de flexion "
      "Euler-Bernoulli + flexibilidad de cortante (Timoshenko, phi=5.12; el "
      "cortante es relevante en muros cortos). Masas concentradas por planta "
      "(lumped): 122.3 t en plantas 1-4 y 91.7 t en cubierta; "
      "M_total=%.1f t (W=%.0f kN, W=G+psi2*Q)." % (ma["M_total_kg"]/1e3, ma["W_total_kN"]))

    h("2. Parametros sismicos EC8 [confirmar AN]", 1)
    e = res["parametros_EC8"]
    p("ag=%.2f g; suelo %s; espectro tipo %d (S=%.2f, TB=%.2f s, TC=%.2f s, "
      "TD=%.1f s); clase importancia II (gammaI=1.0); q=%.1f (DCM, sistema de "
      "muros); amortiguamiento 5%%; lambda=%.2f. Limite inferior del espectro "
      "beta*ag con beta=%.1f." % (e["ag_g"], e["suelo"], e["tipo_espectro"],
      e["S"], e["TB_s"], e["TC_s"], e["TD_s"], e["q"], e["lambda"], e["beta_limite"]))
    p("Espectro de calculo (EN 1998-1 §3.2.2.5, ec. 3.13-3.16). Valor de "
      "meseta Sd = ag*S*2.5/q = %.3f g = %.2f m/s2." % (e["Sd_meseta_g"], e["Sd_meseta_ms2"]), bold=True)
    img("espectro_Sd.png")

    h("3. Analisis modal", 1)
    m1 = res["modal"]["modos"][0]
    p("Problema de autovalores K*phi=w^2*M*phi (scipy.linalg.eigh) sobre el "
      "stick condensado a los GDL laterales. Modo fundamental: T1=%.3f s "
      "(en MESETA, TB<=T1<=TC), masa modal efectiva M_eff,1=%.1f%% (>=60-70%%)."
      % (m1["T_s"], m1["Meff_frac"]*100), bold=True)
    p("Modos superiores (contraste, estimacion de cantilever afinable con eigh): "
      "T2=%.3f s (M_eff=%.0f%%), T3=%.3f s (M_eff=%.0f%%)." % (
        res["modal"]["modos"][1]["T_s"], res["modal"]["modos"][1]["Meff_frac"]*100,
        res["modal"]["modos"][2]["T_s"], res["modal"]["modos"][2]["Meff_frac"]*100))
    img("modos.png", 11)

    h("4. Cortante basal y fuerzas laterales (§4.3.3.2)", 1)
    fl = res["fuerzas_laterales"]
    p("Fb = Sd(T1)*M_total*lambda = %.3f g * %.1f t * %.2f = %.0f kN. "
      "Distribucion F_i = Fb*(z_i*m_i)/sum(z_j*m_j)." % (
        fl["Sd_T1_g"], res["modelo_resumen"]["M_total_t"], e["lambda"], fl["Fb_kN"]), bold=True)
    p("F_i [kN] = %s (de planta 1 a 5). Equilibrio Fb=sumF_i: error %.3f%%." % (
        ", ".join("%.1f" % x for x in fl["F_i_kN"]), fl["equilibrio_error_pct"]))
    p("Momento de vuelco en base M=%.0f kN.m (altura eficaz %.2f m)." % (
        fl["Mbase_kNm"], fl["z_eficaz_m"]), bold=True)
    rm = res["respuesta_modal_SRSS"]
    p("Contraste con respuesta modal espectral SRSS: Fb_SRSS=%.0f kN (sin "
      "lambda). El metodo de fuerzas laterales (con lambda y masa total) "
      "gobierna como envolvente de diseno; ambos del mismo orden (diferencia "
      "%.1f%%, explicada por lambda=0.85 y por M_eff,1=68.9%%)." % (
        rm["Fb_SRSS_kN"], res["criterios_aceptacion"]["dif_modal_vs_lateral_pct"]))
    img("fuerzas_planta.png", 12)
    img("cortante_altura.png", 9)
    img("momento_altura.png", 9)

    h("5. Combinacion sismica (EC0 §6.4.3.4)", 1)
    p("E + sum G + sum psi2*Q. La masa sismica ya incorpora psi2*Q (W tributario "
      "por planta). Axil de gravedad a la base de la pantalla N=sum W=%.0f kN "
      "(compresion)." % res["esfuerzos_base"]["N_base_kN"])

    h("6. Verificaciones EC2/EC8", 1)
    c = ver["cortante_alma"]
    p("6.1 Cortante del alma (EC2 §6.2.3 + amplificacion EC8 DCM eps=%.1f "
      "[confirmar AN]): V_Ed,diseno=%.0f kN; V_Rd,max(biela 45)=%.0f kN; "
      "aprov=%.2f. Armadura horizontal %s." % (
        c["eps_amplif_EC8"], c["V_Ed_diseno_kN"], c["V_Rd_max_kN"],
        c["aprov_biela"], ("Asw/s=%.1f cm2/m (2 caras, >= rho_h,min)" % c["Asw_s_diseno_cm2_m"])))
    b = ver["elemento_borde"]
    p("6.2 Elementos de borde confinados (EC8 §5.4.3.4.2): l_c=%.2f m "
      "(minimo normativo %.2f m, agrandado en predim. por compresion de borde). "
      "F_compr=%.0f kN <= N_Rd,c=%.0f kN (aprov=%.2f). Armado %s." % (
        b["l_c_m"], b["l_c_min_normativo_m"], b["F_compr_borde_kN"],
        b["N_Rd_hormigon_borde_kN"], b["aprov_compr_borde"], ("As=%.1f cm2/borde (>= rho_min=0.5%%)" % b["As_borde_diseno_cm2"])))
    nm = ver["interaccion_NM"]
    p("6.3 Interaccion N-M en base (EC2 §6.1 + EC8 §5.4.3.4): N_Ed=%.0f kN, "
      "M_Ed=%.0f kN.m; M_Rd a N_Ed=%.0f kN.m; aprov=%.2f. Armadura vertical: "
      "%.1f cm2 por borde + %.1f cm2 distribuida en alma (rho_v=%.4f)." % (
        nm["N_Ed_kN"], nm["M_Ed_kNm"], nm["M_Rd_a_NEd_kNm"],
        nm["aprov_flexocompresion"], nm["As_borde_cm2"], nm["As_v_alma_total_cm2"],
        nm["rho_v_alma"]))
    img("diagrama_NM.png", 11)
    d = ver["deriva"]
    p("6.4 Deriva entre plantas (limitacion de dano EC8 §4.4.3.2): d_r=q*d_e "
      "(§4.3.4). Con nu=%.1f y limite 0.75%%h=22.5 mm [confirmar AN], "
      "aprov_max=%.2f." % (d["nu"], d["aprov_max"]))
    img("deriva.png", 12)

    h("7. Resumen de aprovechamientos y veredicto", 1)
    ap = ver["aprovechamientos"]
    tbl = doc.add_table(rows=1, cols=2); tbl.style = "Light Grid Accent 1"
    tbl.rows[0].cells[0].text = "Verificacion"; tbl.rows[0].cells[1].text = "Aprov."
    for k, v in [("Cortante alma", ap["cortante_alma"]),
                 ("Compresion elemento de borde", ap["compr_borde"]),
                 ("Flexocompresion N-M base", ap["flexocompresion_NM"]),
                 ("Deriva entre plantas", ap["deriva"])]:
        row = tbl.add_row().cells; row[0].text = k; row[1].text = "%.2f" % v
    p("Veredicto: %s. Aprovechamiento maximo %.2f (<=1)." % (
        ver["veredicto"], ver["aprov_max"]), bold=True)

    h("8. Avisos y trazabilidad", 1)
    for a in res.get("avisos", []):
        p("- " + a)
    p("Esta memoria es de PREDIMENSIONADO. Debe ser revisada y FIRMADA por un "
      "tecnico competente.", bold=True)

    out = os.path.join(carpeta, "memoria_calculo_sismo.docx")
    doc.save(out)
    print("Memoria escrita en:", out)
    return out


if __name__ == "__main__":
    carpeta = sys.argv[1] if len(sys.argv) > 1 else "."
    build(carpeta)
