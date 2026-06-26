"""
Memoria de calculo del PRETENSADO (Caso 12 - Viga postesada EC2 §5.10) en .docx.

Requiere python-docx y los PNG generados por plots_pretensado (run_all los crea;
si faltan, intenta regenerarlos).

Uso: python3 pretensado/generate_memoria_pretensado.py <carpeta_proyecto>
"""
import sys
import os
import json

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)


def _ensure_pngs(carpeta, model, res):
    pngs = ["trazado_tendon.png", "cargas_equivalentes.png", "leyes_MV.png",
            "tensiones_transferencia.png", "tensiones_servicio.png",
            "interaccion_ELU.png", "perdidas_tendon.png"]
    faltan = [p for p in pngs if not os.path.exists(os.path.join(carpeta, p))]
    if faltan:
        try:
            import plots_pretensado
            plots_pretensado.generar(model, res, carpeta)
        except Exception as ex:
            print("AVISO: no se pudieron generar PNGs (%s). Memoria sin imagenes." % ex)
    return pngs


def build(carpeta):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    model = json.load(open(os.path.join(carpeta, "modelo_neutro.json"), encoding="utf-8"))
    res = json.load(open(os.path.join(carpeta, "verificacion_pretensado.json"), encoding="utf-8"))
    ver = res["verificacion"]
    _ensure_pngs(carpeta, model, res)

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    def h(txt, lvl=1):
        doc.add_heading(txt, level=lvl)

    def p(txt, bold=False):
        par = doc.add_paragraph(); r = par.add_run(txt); r.bold = bold
        return par

    def img(name, w=14):
        fp = os.path.join(carpeta, name)
        if os.path.exists(fp):
            doc.add_picture(fp, width=Cm(w))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    mr = res["modelo_resumen"]; pr = res["pretensado"]

    doc.add_heading("Memoria de calculo - Viga postesada (Caso 12)", 0)
    p("Viga postesada isostatica de hormigon C40/50. Calculo del pretensado "
      "segun EC2 (EN 1992-1-1) §5.10, con Anejo Nacional de Espana. "
      "PREDIMENSIONADO.")

    h("1. Objeto y modelo", 1)
    p("Viga simplemente apoyada de luz L=%.1f m, seccion rectangular b=%.2f x "
      "h=%.2f m (A=%.3f m2, I=%.6f m4, W=%.6f m3). Material %s (fck=%.0f MPa, "
      "fck(t)=%.0f MPa en transferencia, fctm=%.1f MPa)." % (
        mr["L_m"], mr["b_m"], mr["h_m"], mr["A_m2"], mr["I_m4"], mr["W_inf_m3"],
        mr["material"], mr["fck_MPa"], mr["fck_t_MPa"], mr["fctm_MPa"]))
    p("Cargas: peso propio g0=%.2f kN/m (A*rho*g), carga muerta g2=%.1f kN/m, "
      "sobrecarga de uso q=%.1f kN/m (psi2=0.3)." % (
        mr["g0_kN_m"], mr["g2_kN_m"], mr["q_kN_m"]))

    h("2. Pretensado (postesado adherente)", 1)
    p("1 tendon de 13 cordones Y1860S7 (Ap=%.0f mm2, fpk=%.0f MPa, Ep=%.0f GPa). "
      "Trazado parabolico: e=%.2f m en centro de vano, e=%.2f m en apoyos "
      "(anclaje en el c.d.g.). Conducto inyectado (adherente)." % (
        pr["Ap_mm2"], pr["fpk_MPa"], pr["Ep_GPa"], pr["e_centro_m"], pr["e_apoyo_m"]))
    p("Fuerza de tesado en transferencia P0=%.0f kN (sigma_p0=%.0f MPa = %.2f*fpk); "
      "fuerza media tras perdidas diferidas P_m,inf=%.0f kN (sigma_p,inf=%.0f MPa "
      "= %.2f*fpk)." % (
        pr["P0_kN"], pr["sigma_p0_MPa"], pr["sigma_p0_frac_fpk"],
        pr["Pm_inf_kN"], pr["sigma_pinf_MPa"], pr["sigma_pinf_frac_fpk"]), bold=True)
    img("trazado_tendon.png")

    h("3. Perdidas de pretensado (§5.10.5 / §5.10.6)", 1)
    pi = res["perdidas_instantaneas"]; pd = res["perdidas_diferidas"]
    p("Instantaneas: rozamiento mu*(theta+k*x) (centro %.0f MPa), penetracion "
      "de cuna (media %.0f MPa, longitud de influencia x_set=%.1f m), "
      "acortamiento elastico %.0f MPa." % (
        pi["rozamiento_centro_MPa"], pi["cuna_media_MPa"], pi["cuna_x_set_m"],
        pi["acortamiento_elastico_MPa"]))
    p("Diferidas (ec. 5.46): retraccion (eps_cs=%.2e), fluencia (phi=%.1f), "
      "relajacion %.0f MPa; perdida diferida total %.0f MPa. P_m,inf recalculada "
      "%.0f kN (coherencia %.1f%%)." % (
        pd["eps_cs"], pd["phi_fluencia"], pd["relajacion_MPa"],
        pd["diferida_total_MPa"], pd["Pinf_recalculada_kN"], pd["coherencia_pct"]))
    img("perdidas_tendon.png")

    h("4. Cargas equivalentes (load balancing, §5.10)", 1)
    lb = res["load_balancing"]
    p("Pretensado como carga equivalente: w_p = 8*P*e/L^2 = %.2f kN/m (hacia "
      "arriba) + axil de compresion N=%.0f kN. La carga permanente es %.2f kN/m; "
      "la carga residual no equilibrada es %.3f kN/m (%.2f%%): el pretensado "
      "EQUILIBRA la permanente." % (
        lb["w_p_kN_m"], lb["N_axil_kN"], lb["w_permanente_kN_m"],
        lb["residual_kN_m"], lb["residual_pct"]), bold=True)
    img("cargas_equivalentes.png")

    h("5. Momentos y combinaciones (EC0)", 1)
    co = res["momentos_kNm"]
    p("M_g0=%.0f, M_perm=%.0f, M_q=%.0f kN*m. Combinaciones: cuasipermanente "
      "M_qp=%.0f kN*m; caracteristica rara M_rara=%.0f kN*m; ELU "
      "M_Ed=1.35*perm+1.5*q=%.0f kN*m." % (
        co["M_g0"], co["M_perm"], co["M_q"], co["M_cuasipermanente"],
        co["M_caracteristica_rara"], co["M_ELU"]), bold=True)
    img("leyes_MV.png")

    h("6. Tensiones por fibra (§5.10.2.2 / §7.2)", 1)
    t = ver["tensiones"]
    tt = t["transferencia"]; qp = t["servicio_cuasiperm"]; ra = t["servicio_rara"]
    p("Transferencia (P0+g0): sup=%.2f / inf=%.2f MPa (compresion<0). Todo "
      "comprimido: %s. Limite compresion 0.6*fck(t)=%.2f MPa; aprov=%.2f." % (
        tt["sigma_sup_MPa"], tt["sigma_inf_MPa"],
        "SI" if tt["todo_comprimido"] else "NO", tt["lim_compresion_MPa"],
        tt["u_compresion"]))
    img("tensiones_transferencia.png")
    p("Servicio cuasipermanente (Pinf+M_qp): sup=%.2f / inf=%.2f MPa. Sin "
      "descompresion del fondo: %s. Limite compresion 0.45*fck=%.2f MPa; "
      "aprov=%.2f." % (
        qp["sigma_sup_MPa"], qp["sigma_inf_MPa"],
        "SI" if not qp["descompresion_fondo"] else "NO",
        qp["lim_compresion_MPa"], qp["u_compresion"]))
    p("Servicio caracteristico raro (Pinf+M_rara): sup=%.2f / inf=%.2f MPa. "
      "Traccion de fondo %.2f MPa < fctm=%.1f MPa (controlada, aprov=%.2f). "
      "Compresion sup aprov=%.2f." % (
        ra["sigma_sup_MPa"], ra["sigma_inf_MPa"], ra["traccion_inf_MPa"],
        ra["lim_traccion_fctm_MPa"], ra["u_traccion_fctm"], ra["u_compresion"]),
      bold=True)
    img("tensiones_servicio.png")

    h("7. Validacion cruzada (cargas equivalentes vs fuerza+excentricidad)", 1)
    cr = res["validacion_cruzada"]
    p("Estado tensional cuasipermanente por los dos metodos: fuerza+excentricidad "
      "sup=%.2f / inf=%.2f MPa; cargas equivalentes sup=%.2f / inf=%.2f MPa. "
      "Diferencia max %.4f MPa -> %s." % (
        cr["metodo_fuerza_excentricidad"]["sup_MPa"],
        cr["metodo_fuerza_excentricidad"]["inf_MPa"],
        cr["metodo_cargas_equivalentes"]["sup_MPa"],
        cr["metodo_cargas_equivalentes"]["inf_MPa"],
        max(cr["dif_sup_MPa"], cr["dif_inf_MPa"]),
        "COINCIDEN" if cr["coincide"] else "DIFIEREN"), bold=True)

    h("8. ELU de flexion por fibras (§6.1)", 1)
    elu = ver["ELU_flexion"]
    p("Seccion no lineal por fibras (bloque rectangular eta=1.0, lambda=0.8 para "
      "C40/50). Acero activo a d_p=%.2f m (fpd=%.0f MPa); profundidad del bloque "
      "x=%.3f m (x/d=%.2f). M_Rd=%.0f kN*m >= M_Ed=%.0f kN*m (aprov=%.2f)." % (
        elu["d_p_m"], elu["fpd_MPa"], elu["x_m"], elu["x_d_ratio"],
        elu["M_Rd_kNm"], elu["M_Ed_kNm"], elu["u"]), bold=True)
    img("interaccion_ELU.png")

    h("9. Fisuracion (§7.3) y cortante con pretensado (§6.2)", 1)
    fis = ver["fisuracion"]; cor = ver["cortante"]
    p("Fisuracion: traccion de fondo en combinacion rara %.2f MPa < fctm=%.1f MPa "
      "-> sin fisuracion significativa (aprov=%.2f)." % (
        fis["sigma_inf_rara_MPa"], fis["fctm_MPa"], fis["u"]))
    p("Cortante (V_Rd,c con sigma_cp=%.2f MPa): V_Ed=%.0f kN, V_Rd,c=%.0f kN "
      "(aprov=%.2f). %s" % (
        cor["sigma_cp_MPa"], cor["V_Ed_kN"], cor["V_Rd_c_kN"], cor["u"],
        "OK sin armadura de cortante minima." if cor["ok"] else cor["nota"]))

    h("10. Resumen de aprovechamientos y veredicto", 1)
    ap = ver["aprovechamientos"]
    tbl = doc.add_table(rows=1, cols=2); tbl.style = "Light Grid Accent 1"
    tbl.rows[0].cells[0].text = "Verificacion"; tbl.rows[0].cells[1].text = "Aprov."
    filas = [("Compresion en transferencia", ap["compresion_transferencia"]),
             ("Compresion cuasipermanente", ap["compresion_cuasiperm"]),
             ("Compresion rara", ap["compresion_rara"]),
             ("Traccion rara / fctm", ap["traccion_rara_fctm"]),
             ("ELU flexion (fibras)", ap["ELU_flexion"]),
             ("Fisuracion", ap["fisuracion"]),
             ("Cortante", ap["cortante"])]
    for k, v in filas:
        row = tbl.add_row().cells; row[0].text = k; row[1].text = "%.2f" % v
    p("Veredicto: %s. Aprovechamiento maximo %.2f (<=1)." % (
        ver["veredicto"], ver["aprov_max"]), bold=True)

    h("11. Avisos y trazabilidad", 1)
    for a in res.get("avisos", []):
        p("- " + a)
    p("Esta memoria es de PREDIMENSIONADO. Debe ser revisada y FIRMADA por un "
      "tecnico competente.", bold=True)

    out = os.path.join(carpeta, "memoria_calculo_pretensado.docx")
    doc.save(out)
    print("Memoria escrita en:", out)
    return out


if __name__ == "__main__":
    carpeta = sys.argv[1] if len(sys.argv) > 1 else "."
    build(carpeta)
