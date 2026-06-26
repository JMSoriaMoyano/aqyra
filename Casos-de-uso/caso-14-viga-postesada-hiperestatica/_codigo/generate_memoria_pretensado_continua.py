"""Memoria de calculo Word del caso 14 (viga postesada continua hiperestatica)."""
import os, json, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _h(doc, txt, lvl=1):
    h = doc.add_heading(txt, level=lvl)
    return h


def _p(doc, txt, bold=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.bold = bold; r.font.size = Pt(size)
    return p


def _kv(doc, rows):
    t = doc.add_table(rows=0, cols=2); t.style = "Light Grid Accent 1"
    for k, v in rows:
        c = t.add_row().cells
        c[0].text = str(k); c[1].text = str(v)
        c[0].paragraphs[0].runs[0].font.size = Pt(9)
        c[1].paragraphs[0].runs[0].font.size = Pt(9)
    return t


def build(veri, arrays, outdir, imgs):
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10)

    ti = doc.add_heading("Memoria de calculo estructural", level=0)
    _p(doc, "Caso 14 - Viga postesada CONTINUA (hiperestatica) - EC2 (EN 1992-1-1) §5.10",
       bold=True, size=13)
    _p(doc, "Momentos hiperestaticos (secundarios) de pretensado, linea de presiones y "
            "concordancia, ELU con el momento secundario como accion (gamma_P=1,0) y "
            "redistribucion §5.5. Anejo Nacional Espana. PREDIMENSIONADO, a revisar y "
            "firmar por tecnico competente. NDP marcados [confirmar AN].", size=9)

    _h(doc, "1. Descripcion y datos", 1)
    g = veri["geometria"]; s = veri["seccion"]; m = veri["material"]; pr = veri["pretensado"]
    _p(doc, "Viga postesada continua de dos vanos de un forjado/estructura de aparcamiento, "
            "apoyada en 3 apoyos (dos extremos articulados y uno central, que introduce la "
            "hiperestaticidad). Postesado adherente con tendon parabolico continuo por vano.")
    _kv(doc, [
        ("Hormigon", "%s (fck=%.0f MPa, fctm=%.1f MPa, Ecm=%.0f GPa, fck(t)=%.0f MPa)" % (
            m["nombre"], m["fck_Pa"]/1e6, m["fctm_Pa"]/1e6, m["E_Pa"]/1e9, m["fck_t_Pa"]/1e6)),
        ("Seccion", "b=%.2f x h=%.2f m  (A=%.4f m2, I=%.5f m4, W=%.5f m3, c=%.2f m)" % (
            s["b_m"], s["h_m"], s["A_m2"], s["I_m4"], s["W_m3"], s["cdg_m"])),
        ("Geometria", "%d vanos de L=%.1f m (total %.0f m, L/h=%.1f), 3 apoyos en x=%s" % (
            g["n_vanos"], g["L_vano_m"], g["n_vanos"]*g["L_vano_m"], g["L_h"], g["apoyos_x_m"])),
        ("Pretensado", "%d cordones Y1860S7 0.6\" (Ap=%.0f mm2), P0=%.0f kN, Pm,inf=%.0f kN" % (
            pr["n_cordones"], pr["Ap_mm2"], pr["P0_kN"], pr["Pm_inf_kN"])),
        ("Trazado", "parabolico por vano: e_vano=+%.2f m, e_apoyo_central=%.2f m, drape a=%.2f m" % (
            pr["e_centro_vano_m"], pr["e_apoyo_central_m"], pr["drape_m"])),
        ("Cargas", "g0=A·25=16,25 kN/m, g2=5,0 kN/m, q=12,0 kN/m (psi2=0,3)"),
    ])

    _h(doc, "2. Balance de cargas y pretensado", 1)
    _p(doc, "El tendon parabolico de drape a=%.2f m con P_m,inf=%.0f kN genera una carga "
            "equivalente uniforme hacia arriba w_p = 8·P·a/L^2 = %.2f kN/m que equilibra la "
            "carga permanente g0+g2 = 21,25 kN/m (residual %.2f%%). sigma_p0/fpk=%.3f, "
            "sigma_pm,inf/fpk=%.3f (perdidas diferidas %.1f%%)." % (
            pr["drape_m"], pr["Pm_inf_kN"], pr["w_p_kN_m"],
            veri["validaciones"]["balance_residual_pct"],
            veri["validaciones"]["sigma_p0_sobre_fpk"], veri["validaciones"]["sigma_pinf_sobre_fpk"],
            veri["validaciones"]["perdidas_diferidas_pct"]))
    doc.add_picture(imgs["alzado_tendon.png"], width=Cm(16))
    doc.add_picture(imgs["cargas_equivalentes.png"], width=Cm(16))

    _h(doc, "3. Momentos hiperestaticos (secundarios) de pretensado", 1)
    mp = veri["momentos_pretensado"]
    _p(doc, "En una estructura HIPERESTATICA el pretensado induce reacciones en los apoyos "
            "redundantes (aqui el apoyo central). Se distinguen:")
    _p(doc, "  - Primario  M1(x) = -P·e(x)  (estructura liberada/isostatica).", size=9)
    _p(doc, "  - Total  M_p,tot(x) = momento de las cargas equivalentes aplicadas a la viga "
            "CONTINUA (FEM).", size=9)
    _p(doc, "  - Secundario/hiperestatico  M_sec = M_p,tot - M1, LINEAL entre apoyos y NULO "
            "en los apoyos extremos (isostaticos), maximo sobre el apoyo central.", size=9)
    ac = mp["apoyo_central"]; cv = mp["centro_vano"]
    _kv(doc, [
        ("Apoyo central", "M1=%.0f, M_p,tot=%.0f, M_sec=%.0f kN·m" % (
            ac["M1_kNm"], ac["M_p_tot_kNm"], ac["M_sec_kNm"])),
        ("Centro de vano", "M1=%.0f, M_p,tot=%.0f, M_sec=%.0f kN·m" % (
            cv["M1_kNm"], cv["M_p_tot_kNm"], cv["M_sec_kNm"])),
        ("M_sec en extremos", "%.2f / %.2f kN·m (nulo)" % (
            veri["validaciones"]["M_sec_extremo0_kNm"], veri["validaciones"]["M_sec_extremo2_kNm"])),
        ("Linealidad M_sec", "R2 = %.6f" % veri["validaciones"]["M_sec_linealidad_R2"]),
        ("FEM vs metodo de las fuerzas", "%.1f vs %.1f kN·m (Delta %.3f%%)" % (
            veri["validaciones"]["FEM_vs_metodo_fuerzas"]["M_sec_FEM_kNm"],
            veri["validaciones"]["FEM_vs_metodo_fuerzas"]["M_sec_fuerzas_kNm"],
            veri["validaciones"]["FEM_vs_metodo_fuerzas"]["delta_pct"])),
        ("Identidad M_p,tot = M1 + M_sec", "max error %.2e N·m" % (
            veri["validaciones"]["identidad_Mptot_M1_Msec_max_Nm"])),
    ])
    doc.add_picture(imgs["momentos_M1_Mptot_Msec.png"], width=Cm(16))

    _h(doc, "4. Linea de presiones y concordancia", 1)
    co = mp["concordancia"]
    _p(doc, "La linea de presiones e_p(x) = M_p,tot/P = e(x) + M_sec/P. El tendon es "
            "CONCORDANTE si M_sec = 0 (la linea de presiones coincide con el tendon). Aqui "
            "M_sec != 0 -> tendon NO concordante; la linea de presiones se separa del tendon "
            "%.3f m en el apoyo central (= M_sec/P)." % co["desviacion_linea_presiones_apoyo_m"])
    doc.add_picture(imgs["linea_presiones.png"], width=Cm(16))

    _h(doc, "5. Esfuerzos externos y leyes de servicio/ELU", 1)
    ee = veri["esfuerzos_externos"]
    _kv(doc, [
        ("Apoyo central (hogging)", "M_perm=%.0f, M_qp=%.0f, M_rara=%.0f, M_ELU=%.0f kN·m" % (
            ee["apoyo_central"]["M_perm_kNm"], ee["apoyo_central"]["M_qp_kNm"],
            ee["apoyo_central"]["M_rara_kNm"], ee["apoyo_central"]["M_ELU_kNm"])),
        ("Centro de vano (sagging, x=%.1f m)" % ee["centro_vano"]["x_m"],
         "M_perm=%.0f, M_rara=%.0f, M_ELU=%.0f (patron %.0f) kN·m" % (
            ee["centro_vano"]["M_perm_kNm"], ee["centro_vano"]["M_rara_kNm"],
            ee["centro_vano"]["M_ELU_kNm"], ee["centro_vano"]["M_ELU_patron_kNm"])),
    ])
    doc.add_picture(imgs["leyes_MV.png"], width=Cm(16))

    _h(doc, "6. Tensiones por fibra (con momento secundario)", 1)
    _p(doc, "Tensiones sigma = -P/A -+ (M_ext - P·e + M_sec)·c/I (compresion negativa). "
            "Limites: transferencia 0,6·fck(t)=19,2 MPa; cuasiperm 0,45·fck=18 MPa; rara "
            "0,6·fck=24 MPa y traccion < fctm=3,5 MPa.")
    for key, ttl in [("tensiones_apoyo_central", "Apoyo central"),
                     ("tensiones_centro_vano", "Centro de vano")]:
        t = veri[key]
        _p(doc, ttl + ":", bold=True, size=9)
        _kv(doc, [(e, "sup=%+.2f / inf=%+.2f MPa  (u_comp=%.2f, ok=%s)" % (
            t[e]["sigma_sup_MPa"], t[e]["sigma_inf_MPa"], t[e]["u_compresion"], t[e]["ok"]))
            for e in ["transferencia", "servicio_cuasiperm", "servicio_rara"]])
    doc.add_picture(imgs["tensiones_fibra.png"], width=Cm(16))

    _h(doc, "7. ELU de flexion con el momento secundario (§5.10.8)", 1)
    _p(doc, "El momento secundario es un efecto hiperestatico de reaccion -> se incluye como "
            "ACCION con gamma_P=1,0:  M_Ed = gamma_G·M_g + gamma_Q·M_q + 1,0·M_sec. La "
            "resistencia M_Rd se obtiene por FIBRAS con el acero activo a f_pd (+ pasivo), sin "
            "doble computo del primario.")
    ea = veri["ELU_apoyo_central"]; ec = veri["ELU_centro_vano"]
    _kv(doc, [
        ("Apoyo central (hogging)", "M_ELU_ext=%.0f + M_sec=%.0f -> M_Ed=%.0f / M_Rd=%.0f kN·m "
         "(u=%.2f, x/d=%.3f). El secundario ALIVIA el hogging." % (
            ea["M_ELU_ext_kNm"], ea["M_sec_kNm"], ea["M_Ed_con_secundario_kNm"],
            ea["M_Rd_kNm"], ea["u"], ea["x_d_ratio"])),
        ("Centro de vano (sagging)", "M_Ed=%.0f / M_Rd=%.0f kN·m (u=%.2f, x/d=%.3f)" % (
            ec["M_Ed_con_secundario_kNm"], ec["M_Rd_kNm"], ec["u"], ec["x_d_ratio"])),
    ])
    rd = veri["redistribucion_apoyo_central"]
    _p(doc, "Redistribucion §5.5: con x/d=%.3f el coeficiente minimo admisible es delta=%.3f "
            "(reduccion maxima del hogging %.0f%%), opcional ya que u<1; no se aplica." % (
            rd["x_d"], rd["delta_min"], rd["reduccion_max_pct"]), size=9)
    doc.add_picture(imgs["interaccion_ELU.png"], width=Cm(13))

    _h(doc, "8. Flecha", 1)
    fl = veri["flecha"]
    _p(doc, "Bajo carga residual neta (perm + psi2·q - balance del pretensado = %.2f kN/m) la "
            "flecha maxima es %.2f mm << L/250 = %.0f mm (u=%.2f): el balance del pretensado "
            "deja un residual permanente pequeno." % (
            fl["residual_neto_kN_m"], fl["flecha_max_mm"], fl["limite_L250_mm"], fl["u"]))
    doc.add_picture(imgs["flecha.png"], width=Cm(16))

    _h(doc, "9. Verificacion y conclusiones", 1)
    ap = veri["aprovechamientos"]
    _kv(doc, [(k, "%.2f" % v) for k, v in ap.items()] + [
        ("Aprovechamiento maximo", "%.2f" % veri["aprov_max"]),
        ("Veredicto", veri["veredicto"])])
    _p(doc, "Validaciones: balance residual %.2f%%; M_sec lineal (R2=%.4f) y nula en extremos; "
            "FEM vs metodo de las fuerzas Delta %.3f%%; identidad M_p,tot=M1+M_sec; "
            "sigma_p0/fpk=%.3f, sigma_pm,inf/fpk=%.3f. Aprovechamientos <= 1, picos como "
            "envolvente." % (
            veri["validaciones"]["balance_residual_pct"], veri["validaciones"]["M_sec_linealidad_R2"],
            veri["validaciones"]["FEM_vs_metodo_fuerzas"]["delta_pct"],
            veri["validaciones"]["sigma_p0_sobre_fpk"], veri["validaciones"]["sigma_pinf_sobre_fpk"]), size=9)
    _p(doc, "PREDIMENSIONADO. Resultado a revisar y firmar por tecnico competente. NDP del "
            "Anejo Nacional Espana marcados [confirmar AN]: coeficientes de perdidas, limites "
            "de tension del acero activo, mu/k de rozamiento, coeficiente de redistribucion delta.",
       size=9)

    out = os.path.join(outdir, "memoria_calculo_pretensado_continua.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    outdir = sys.argv[1] if len(sys.argv) > 1 else "."
    veri = json.load(open(os.path.join(outdir, "verificacion_pretensado_continua.json")))
    arrays = json.load(open(os.path.join(outdir, "_arrays_plot.json")))
    imgs = {n: os.path.join(outdir, n) for n in os.listdir(outdir) if n.endswith(".png")}
    out = build(veri, arrays, outdir, imgs)
    print("Memoria:", out)
